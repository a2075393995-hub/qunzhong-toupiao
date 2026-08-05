from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from vote_core import (
    VoteRecord,
    apply_field_targets,
    blank_mapping,
    infer_field_targets,
    select_preview_record,
    selected_mark_pair_refs,
    split_room_value,
)


class FieldLayoutTests(unittest.TestCase):
    def record(self, room: str = "1-101") -> VoteRecord:
        return VoteRecord(2, room, "张三", "13900000000", [], {})

    def test_split_room_value(self):
        self.assertEqual(("1", "101"), split_room_value("1-101"))
        self.assertEqual(("2", "305"), split_room_value("2栋305室"))

    def test_plain_underlines_keep_units_and_split_values(self):
        document = Document()
        document.add_paragraph("地址：____栋_____室")
        inferred = infer_field_targets(document)
        self.assertIn("building", inferred)
        self.assertIn("roomNo", inferred)

        mapping = blank_mapping()
        mapping["fieldStyles"]["building"].update({"fontName": "微软雅黑", "fontSize": 14, "bold": True, "offsetY": -1})
        mapping["fieldStyles"]["roomNo"].update({"fontSize": 11, "offsetY": 1})
        apply_field_targets(document, mapping, self.record())

        paragraph = document.paragraphs[0]
        self.assertEqual("地址： 1  栋 101 室", paragraph.text)
        self.assertIn("栋", paragraph.text)
        self.assertIn("室", paragraph.text)
        building_run = next(run for run in paragraph.runs if run.text.strip() == "1")
        room_run = next(run for run in paragraph.runs if run.text.strip() == "101")
        self.assertEqual("微软雅黑", building_run.font.name)
        self.assertEqual(14.0, building_run.font.size.pt)
        self.assertTrue(building_run.bold)
        self.assertEqual("2", building_run._r.get_or_add_rPr().find(qn("w:position")).get(qn("w:val")))
        self.assertEqual("-2", room_run._r.get_or_add_rPr().find(qn("w:position")).get(qn("w:val")))

    def test_formatted_underlines_keep_surrounding_runs(self):
        document = Document()
        paragraph = document.add_paragraph("地址：")
        building_slot = paragraph.add_run("    ")
        building_slot.font.underline = True
        paragraph.add_run("幢")
        room_slot = paragraph.add_run("     ")
        room_slot.font.underline = True
        paragraph.add_run("室")

        mapping = blank_mapping()
        apply_field_targets(document, mapping, self.record("3-502"))
        self.assertEqual("地址： 3  幢 502 室", paragraph.text)
        self.assertEqual("幢", paragraph.runs[2].text)
        self.assertEqual("室", paragraph.runs[4].text)

    def test_dedicated_room_cell_keeps_complete_original_value(self):
        document = Document()
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "楼栋房号："
        table.cell(0, 1).text = ""
        mapping = blank_mapping()
        mapping["fieldTargets"]["room"] = {"table": 0, "row": 0, "col": 1}
        apply_field_targets(document, mapping, self.record())
        self.assertEqual("1-101", table.cell(0, 1).text)

    def test_preview_uses_first_valid_record_for_document_and_overlay(self):
        mapping = blank_mapping()
        mapping["validation"].update({"mode": "range", "min": 1, "max": 3})
        invalid = VoteRecord(2, "1-101", "废票行", "13000000000", [], {})
        valid = VoteRecord(3, "2-202", "有效行", "13100000000", ["选项1"], {})
        selected, reasons = select_preview_record([invalid, valid], mapping)
        self.assertIs(selected, valid)
        self.assertEqual([], reasons)

    def test_print_preview_lists_only_marks_visible_for_selected_record(self):
        mapping = blank_mapping()
        mapping["options"]["结果1"] = {
            "pairs": [
                {"judgmentText": "同意", "mark": {"table": 0, "row": 1, "col": 1}},
                {"judgmentText": "反对", "mark": {"table": 0, "row": 1, "col": 2}},
            ]
        }
        record = VoteRecord(2, "1-101", "张三", "13900000000", [], {}, {"结果1": ["反对"]})
        refs = selected_mark_pair_refs(record, mapping)
        self.assertEqual([{"key": "结果1", "pairIndex": 1, "label": "结果1：反对"}], refs)


if __name__ == "__main__":
    unittest.main()
