from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))


def set_run_font(run, name="Microsoft YaHei UI", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(document, text="", size=10.5, bold=False, color=None, align=None):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(13 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(
        run,
        size=15 if level == 1 else 12.5,
        bold=True,
        color="123B5D" if level == 1 else "176B87",
    )
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5)
    return paragraph


def shade_cell(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=9.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)


def add_image(document, path: Path, caption: str, width=6.35):
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(8)
    run = caption_paragraph.add_run(caption)
    set_run_font(run, size=9, color="64748B")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8.5, color="64748B")
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend([field_begin, instruction, field_end])
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=8.5, color="64748B")


def build(output_dir: Path, screenshot_dir: Path) -> tuple[Path, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.76)
    section.right_margin = Inches(0.76)
    add_page_number(section.footer.paragraphs[0])

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei UI"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
    normal.font.size = Pt(10.5)

    document.core_properties.title = "群众选票格式化打印工具 v0.3.6 使用说明"
    document.core_properties.author = "群众选票格式化打印工具"
    document.core_properties.subject = "Gitee 无依赖便携版操作说明"

    add_para(
        document,
        "群众选票格式化打印工具",
        size=23,
        bold=True,
        color="0F172A",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        document,
        "v0.3.6 · Word / WPS / 内置引擎 · 真实 PDF 打印预览",
        size=12,
        color="176B87",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    callout = document.add_table(rows=1, cols=1)
    callout.style = "Table Grid"
    shade_cell(callout.cell(0, 0), "E7F5F7")
    set_cell_text(
        callout.cell(0, 0),
        "用途：读取 Word 表决票模板和 CSV/XLSX 投票数据，通过可视化标注生成 DOCX；导出前按 Microsoft Word、WPS、内置引擎的顺序生成真实 PDF 预览，文字和打勾均可拖拽或用方向键微调。",
        size=10.5,
    )

    add_heading(document, "一、交付包内容")
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ("文件/目录", "用途")):
        shade_cell(cell, "123B5D")
        set_cell_text(cell, text, bold=True, color="FFFFFF")
    for left, right in [
        ("QunzhongVote.exe", "正式 Windows 程序，双击后显示中文界面。"),
        ("runtime\\libreoffice", "程序自带文档转换引擎，请勿删除或单独移动 EXE。"),
        ("PORTABLE_README.txt", "无依赖运行和解压说明。"),
    ]:
        cells = table.add_row().cells
        set_cell_text(cells[0], left)
        set_cell_text(cells[1], right)

    add_heading(document, "二、六步使用流程")
    for step in [
        "1. 上传模板：选择 .docx 或 .doc 表决票模板。",
        "2. 上传数据文件：选择 CSV/XLSX，并设置有效投票结果数量。",
        "3. 设置模板：选择字段或投票结果，按顺序标注判断区和标记区。",
        "4. 确认设置：保存当前模板全部判断区、标记区、字段位置和样式。",
        "5. 导出前预览：查看 Word、WPS 或内置引擎生成的真实 PDF 页面，并调整文字/打勾位置。",
        "6. 开始导出：生成多文件或单文件 DOCX，并输出投票结果汇总表。",
    ]:
        add_para(document, step)

    add_heading(document, "三、模板设置界面")
    add_bullet(document, "主页面只负责标注判断区、标记区和字段填入区，不再提供打勾位置微调。")
    add_bullet(document, "标注会按模板文件内容保存在本机；重新选择同一模板或重启程序后自动恢复。")
    add_bullet(document, "判断区与标记区必须按相同顺序一一对应。Ctrl/Shift 可选择多个共用规则的投票结果。")
    add_image(document, screenshot_dir / "template-profile.png", "图 1：模板设置与数据预览")

    add_heading(document, "四、真实打印预览与位置调整")
    for item in [
        "导出前预览按 Microsoft Word → WPS Writer → 内置文档引擎的顺序生成真实 PDF 页面。",
        "已安装 WPS 但没有 Word 时，预览会优先贴近 WPS 实际打印效果；两者都没有时仍可零依赖运行。",
        "左侧“调整对象”同时列出用户信息文字和当前票面实际出现的打勾。",
        "可直接拖动红框；普通方向键移动 1pt，Shift 移动 5pt，Ctrl 移动 0.1pt。",
        "停手后后台重新生成精确文档→PDF 预览；刷新完成前不能确认导出。",
    ]:
        add_bullet(document, item)
    add_image(document, screenshot_dir / "true-print-preview.png", "图 2：真实 A4 打印预览及打勾位置调整")

    main_image = screenshot_dir / "release-main.png"
    if main_image.exists():
        add_heading(document, "五、正式程序")
        add_image(document, main_image, "图 3：本地构建并启动的 v0.3.6 正式程序")

    add_heading(document, "六、随版本完成的测试数据与结果")
    test_table = document.add_table(rows=1, cols=4)
    test_table.style = "Table Grid"
    for cell, text in zip(test_table.rows[0].cells, ("楼栋房号", "姓名", "电话号码", "投票选项")):
        shade_cell(cell, "123B5D")
        set_cell_text(cell, text, bold=True, color="FFFFFF", size=9)
    for values in [
        ("1-101", "验收用户", "13900000000", "选项1"),
        ("2-202", "方向键测试", "13800000000", "选项2"),
    ]:
        cells = test_table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell_text(cell, value, size=9)
    for result in [
        "29 项单元测试全部通过，覆盖 Word → WPS → 内置引擎的优先级、失败回退和超时进程树清理。",
        "纯净环境自检成功：v0.3.6、PDF 1 页；不需要 Word、WPS 或 Python。",
        "真实预览验收成功：PDF 页面图像 1 张、当前票面打勾 1 个、方向键最终保存为 +1pt。",
        "模板配置重启恢复成功：判断区/标记区由 1/1 恢复为 1/1；4 份批量回归 DOCX，0 条警告。",
    ]:
        add_bullet(document, result)

    add_heading(document, "七、导出与更新")
    for item in [
        "多文件模式：每条有效数据生成一份 DOCX。",
        "单文件模式：全部有效数据合并到同一 DOCX，每条记录从新页开始。",
        "纯净模式：只保留程序写入的文字和打勾，适合套打到预印纸张。",
        "点击右上角“检查更新”会访问 Gitee 的公开 Release；程序不会上传模板、投票数据或导出文件。",
        "更新仓库：https://gitee.com/zhang-jiaxin654/qunzhong-toupiao",
    ]:
        add_bullet(document, item)

    add_heading(document, "八、常见问题")
    for question, answer in [
        ("预览为什么不再是灰色？", "v0.3.6 会在真实 PDF 转换完成后立即绘制页面；Canvas 至少包含一张实际 PDF 页面图像。"),
        ("设置模板后原来的标记会消失吗？", "不会。同一模板按内容哈希恢复全部判断区、标记区、字段位置和样式。"),
        ("为什么确认预览按钮暂时不可点？", "说明程序正在重新生成精确文档→PDF 页面，完成后按钮会自动恢复。"),
        ("只有 WPS 可以吗？", "可以。程序会在没有 Word 时自动尝试 WPS；若 WPS 接口不可用，再自动回退到内置引擎。"),
        ("Word 和 WPS 都没有怎么办？", "无需安装。完整便携版自带文档引擎，会自动完成真实 PDF 预览。"),
    ]:
        add_para(document, question, bold=True, color="123B5D")
        add_para(document, answer)

    docx_path = output_dir / "群众选票格式化打印工具_v0.3.6_使用说明.docx"
    pdf_path = output_dir / "群众选票格式化打印工具_v0.3.6_使用说明.pdf"
    document.save(docx_path)

    from print_preview import docx_to_pdf

    docx_to_pdf(docx_path, pdf_path)
    return docx_path, pdf_path


def main():
    parser = argparse.ArgumentParser(description="Build the illustrated v0.3.6 DOCX/PDF user manual.")
    parser.add_argument("output_dir")
    parser.add_argument("screenshot_dir")
    args = parser.parse_args()
    docx_path, pdf_path = build(
        Path(args.output_dir).resolve(),
        Path(args.screenshot_dir).resolve(),
    )
    print(docx_path)
    print(pdf_path)


if __name__ == "__main__":
    main()
