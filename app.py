from __future__ import annotations

import copy
import os
import sys
import threading
import tkinter as tk
import webbrowser
from pathlib import Path
from tkinter import filedialog, messagebox, ttk
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont, ImageTk

from print_preview import docx_to_pdf, render_pdf_pages, search_pdf_text
from update_service import APP_VERSION, REPOSITORY_URL, ReleaseInfo, fetch_latest_release, is_newer_version

from vote_core import (
    APP_NAME,
    FIELD_LABELS,
    blank_mapping,
    config_pairs,
    configured_option_keys,
    convert_doc_to_docx,
    document_tables,
    docx_page_count,
    generate_all,
    generate_preview_docx,
    make_result_option_key,
    normalize_option,
    normalize_result_name,
    record_parse_options,
    read_vote_records,
    select_preview_record,
    split_room_value,
    target_label,
    validate_vote_record,
    UNDERLINE_PATTERN,
    formatted_underline_spans,
)


def app_base_dir() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent


def resource_path(relative_path: str) -> Path:
    base = Path(getattr(sys, "_MEIPASS", app_base_dir()))
    return base / relative_path


BASE_DIR = app_base_dir()


class ScrollFrame(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        self.canvas = tk.Canvas(self, borderwidth=0, highlightthickness=0, background="#eef2f7")
        self.inner = ttk.Frame(self.canvas)
        self.scrollbar_y = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.scrollbar_x = ttk.Scrollbar(self, orient="horizontal", command=self.canvas.xview)
        self.canvas.configure(yscrollcommand=self.scrollbar_y.set, xscrollcommand=self.scrollbar_x.set)

        self.scrollbar_y.pack(side="right", fill="y")
        self.scrollbar_x.pack(side="bottom", fill="x")
        self.canvas.pack(side="left", fill="both", expand=True)
        self.window_id = self.canvas.create_window((0, 0), window=self.inner, anchor="nw")
        self.inner.bind("<Configure>", self._on_inner_configure)
        self.canvas.bind("<Configure>", self._on_canvas_configure)
        self.bind_mousewheel(self.canvas)
        self.bind_mousewheel(self.inner)

    def _on_inner_configure(self, _event):
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))

    def _on_canvas_configure(self, event):
        self.canvas.itemconfigure(self.window_id, width=max(event.width, self.inner.winfo_reqwidth()))

    def bind_mousewheel(self, widget):
        widget.bind("<MouseWheel>", self._on_mousewheel)
        widget.bind("<Button-4>", self._on_mousewheel)
        widget.bind("<Button-5>", self._on_mousewheel)

    def bind_mousewheel_to_descendants(self, widget=None, exclude=()):
        widget = widget or self.inner
        for child in widget.winfo_children():
            if not isinstance(child, exclude):
                child.bind("<MouseWheel>", self._on_mousewheel, add="+")
                child.bind("<Button-4>", self._on_mousewheel, add="+")
                child.bind("<Button-5>", self._on_mousewheel, add="+")
            self.bind_mousewheel_to_descendants(child, exclude=exclude)

    def bind_outer_mousewheel(self, widget):
        widget.bind("<MouseWheel>", self._on_mousewheel)
        widget.bind("<Button-4>", self._on_mousewheel)
        widget.bind("<Button-5>", self._on_mousewheel)

    def _on_mousewheel(self, event):
        if getattr(event, "num", None) == 4:
            units = -3
        elif getattr(event, "num", None) == 5:
            units = 3
        else:
            units = -3 if event.delta > 0 else 3
        self.canvas.yview_scroll(units, "units")
        return "break"


class VoteDocxApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title(f"{APP_NAME} - 发布版")
        self.geometry("1240x780")
        self.minsize(1100, 700)
        self.configure(background="#eef2f7")
        self.set_window_icon()

        self.template_path = tk.StringVar()
        self.data_path = tk.StringVar()
        self.output_dir = tk.StringVar(value=str(BASE_DIR / "output"))
        self.filename_prefix = tk.StringVar()

        self.brush_mode = tk.StringVar(value="judgment")
        self.custom_option = tk.StringVar()
        self.choice_mode = tk.StringVar(value="single")
        self.field_keys = ["field:building", "field:roomNo", "field:room", "field:name", "field:phone"]
        self.field_list = None
        self.selected_key: Optional[str] = None
        self.option_keys: List[str] = []
        self.option_labels: Dict[str, str] = {}
        self.mapping: Dict[str, Any] = blank_mapping()
        self.pending_pair_index: Dict[str, int] = {}
        self.tables: List[List[List[str]]] = []
        self.records = []
        self.preview_cells: List[Dict[str, Any]] = []
        self.preview_image_tk = None

        self.validation_mode = tk.StringVar(value="range")
        self.min_count = tk.IntVar(value=1)
        self.max_count = tk.IntVar(value=11)
        self.exact_count = tk.IntVar(value=1)
        self.export_mode = tk.StringVar(value="multi")
        self.clean_mode = tk.BooleanVar(value=False)
        self.status_text = tk.StringVar(value="先选择左侧项目，再点击右侧预览图上的区域")
        self.mark_horizontal = tk.StringVar(value="居中")
        self.mark_vertical = tk.StringVar(value="居中")
        self.mark_font_size = tk.IntVar(value=10)
        self.mark_offset_x = 0
        self.mark_offset_y = 0
        self.apply_style_to_all_options = tk.BooleanVar(value=True)
        self.preview_ready = False
        self.preview_path: Optional[Path] = None
        self.file_status_text = tk.StringVar(value="流程：上传模板 -> 上传数据文件 -> 设置模板 -> 导出前预览 -> 开始导出")
        self.drag_adjust_start: Optional[Tuple[int, int, int, int]] = None
        self.drag_adjust_moved = False
        self.mouse_down_cell: Optional[Dict[str, Any]] = None
        self.preview_widget = None
        self.workflow_buttons: Dict[str, tk.Button] = {}
        self.workflow_done = set()
        self.debug_workspace_enabled = False
        self.debug_completed = False
        self.debug_controls: List[Any] = []
        self.select_mark_for_adjust = False
        self.selected_mark_ref: Optional[Tuple[str, int]] = None
        self.undo_stack: List[Dict[str, Any]] = []
        self.result_group_text = tk.StringVar(value="多结果标记：未选择")
        self.last_result_selection: List[str] = []
        self.template_page_count: Optional[int] = None
        self.update_button: Optional[tk.Button] = None

        self._build_style()
        self._build_ui()
        self.set_debug_workspace_enabled(False)

    def set_window_icon(self):
        icon_path = resource_path("assets/app_icon.ico")
        if icon_path.exists():
            try:
                self.iconbitmap(default=str(icon_path))
            except Exception:
                pass

    def _build_style(self):
        style = ttk.Style(self)
        if "vista" in style.theme_names():
            style.theme_use("vista")
        style.configure("TFrame", background="#eef2f7")
        style.configure("TLabel", background="#eef2f7", foreground="#111827")
        style.configure("TLabelframe", background="#eef2f7")
        style.configure("Header.TFrame", background="#1f2937")
        style.configure("Header.TLabel", background="#1f2937", foreground="#ffffff", font=("Microsoft YaHei UI", 16, "bold"))
        style.configure("SubHeader.TLabel", background="#1f2937", foreground="#cbd5e1", font=("Microsoft YaHei UI", 9))
        style.configure("Version.TLabel", background="#1f2937", foreground="#a7f3d0", font=("Microsoft YaHei UI", 10, "bold"))
        style.configure("Panel.TLabelframe", background="#eef2f7")
        style.configure("Panel.TLabelframe.Label", font=("Microsoft YaHei UI", 10, "bold"))
        style.configure("Accent.TButton", font=("Microsoft YaHei UI", 9, "bold"))
        style.configure("Brush.TButton", font=("Microsoft YaHei UI", 10, "bold"))
        style.configure("Cell.TButton", font=("Microsoft YaHei UI", 9))
        style.configure("Mapped.TButton", font=("Microsoft YaHei UI", 9, "bold"))

    def _build_ui(self):
        header = ttk.Frame(self, style="Header.TFrame", padding=(18, 14))
        header.pack(side="top", fill="x")
        header_left = ttk.Frame(header, style="Header.TFrame")
        header_left.pack(side="left", fill="x", expand=True)
        ttk.Label(header_left, text=APP_NAME, style="Header.TLabel").pack(anchor="w")
        ttk.Label(header_left, text="上传模板和投票数据，设置填入位置后批量生成一人一份 DOCX", style="SubHeader.TLabel").pack(anchor="w", pady=(4, 0))
        header_right = ttk.Frame(header, style="Header.TFrame")
        header_right.pack(side="right", padx=(16, 0))
        ttk.Label(header_right, text=f"发布版 v{APP_VERSION}", style="Version.TLabel").pack(side="left", padx=(0, 10))
        self.update_button = tk.Button(
            header_right,
            text="检查更新",
            command=self.check_for_updates,
            relief="flat",
            borderwidth=0,
            padx=12,
            pady=5,
            background="#0f766e",
            foreground="#ffffff",
            activebackground="#115e59",
            activeforeground="#ffffff",
            font=("Microsoft YaHei UI", 9, "bold"),
            cursor="hand2",
        )
        self.update_button.pack(side="left")

        self._build_workflow_bar()

        main_split = tk.PanedWindow(
            self,
            orient=tk.VERTICAL,
            sashwidth=8,
            sashrelief="raised",
            showhandle=True,
            handlesize=16,
            bg="#cbd5e1",
            bd=0,
            opaqueresize=True,
        )
        main_split.pack(fill="both", expand=True, padx=12, pady=12)
        self.tab_mapper = ttk.Frame(main_split, padding=0)
        self.data_panel = ttk.Frame(main_split, padding=0)
        main_split.add(self.tab_mapper, minsize=360, stretch="always")
        main_split.add(self.data_panel, minsize=130)

        self._build_mapper_tab()
        self._build_data_panel()
        self.update_file_status()

    def _build_workflow_bar(self):
        workflow = ttk.LabelFrame(self, text="操作流程", style="Panel.TLabelframe", padding=10)
        workflow.pack(side="top", fill="x", padx=12, pady=(10, 0))
        self.workflow_button(workflow, "template", "1 上传模板", self.choose_template).grid(row=0, column=0, padx=(0, 8), sticky="w")
        self.workflow_button(workflow, "data", "2 上传数据文件", self.choose_data).grid(row=0, column=1, padx=(0, 8), sticky="w")
        self.workflow_button(workflow, "debug", "3 设置模板", self.debug_template).grid(row=0, column=2, padx=(0, 8), sticky="w")
        debug_done_button = self.workflow_button(workflow, "debug_done", "4 确认设置", self.finish_debug)
        debug_done_button.grid(row=0, column=3, padx=(0, 8), sticky="w")
        self.debug_controls.append(debug_done_button)
        self.workflow_button(workflow, "preview", "5 导出前预览", self.preview_docx).grid(row=0, column=4, padx=(0, 8), sticky="w")
        self.workflow_button(workflow, "export", "6 开始导出", self.start_export).grid(row=0, column=5, padx=(0, 14), sticky="w")
        ttk.Button(workflow, text="上一步", command=self.undo_last_action).grid(row=0, column=6, padx=(0, 14), sticky="w")
        ttk.Button(workflow, text="输出目录", command=self.choose_output_dir).grid(row=0, column=7, padx=(0, 8), sticky="w")
        ttk.Button(workflow, text="打开输出目录", command=self.open_output_dir).grid(row=0, column=8, padx=(0, 14), sticky="w")
        tk.Button(
            workflow,
            text="重置",
            command=self.reset_app_state,
            relief="raised",
            borderwidth=1,
            padx=12,
            pady=4,
            background="#fef2f2",
            foreground="#991b1b",
            activebackground="#fee2e2",
            activeforeground="#7f1d1d",
            font=("Microsoft YaHei UI", 9, "bold"),
        ).grid(row=0, column=9, sticky="w")
        ttk.Label(workflow, textvariable=self.file_status_text, foreground="#374151").grid(row=1, column=0, columnspan=10, sticky="w", pady=(8, 0))
        ttk.Label(workflow, text="文件名前缀：").grid(row=2, column=0, sticky="w", pady=(8, 0))
        prefix_entry = ttk.Entry(workflow, textvariable=self.filename_prefix, width=18)
        prefix_entry.grid(row=2, column=1, columnspan=2, sticky="w", pady=(8, 0))
        prefix_entry.bind("<KeyRelease>", lambda _event: self.on_filename_prefix_changed())
        ttk.Label(workflow, text="可为空，导出文件名为“前缀+姓名.docx”。", foreground="#6b7280").grid(row=2, column=3, columnspan=7, sticky="w", pady=(8, 0))
        workflow.columnconfigure(10, weight=1)

    def workflow_button(self, parent, key: str, text: str, command):
        button = tk.Button(
            parent,
            text=text,
            command=command,
            relief="raised",
            borderwidth=1,
            width=12,
            padx=10,
            pady=4,
            background="#ffffff",
            foreground="#111827",
            activebackground="#e0f2fe",
            activeforeground="#111827",
            font=("Microsoft YaHei UI", 9, "bold"),
        )
        self.workflow_buttons[key] = button
        return button

    def mark_workflow_done(self, key: str):
        self.workflow_done.add(key)
        button = self.workflow_buttons.get(key)
        if button:
            button.configure(background="#16a34a", foreground="#ffffff", activebackground="#15803d", activeforeground="#ffffff")

    def reset_workflow_after(self, *keys: str):
        for key in keys:
            self.workflow_done.discard(key)
            button = self.workflow_buttons.get(key)
            if button:
                button.configure(background="#ffffff", foreground="#111827", activebackground="#e0f2fe", activeforeground="#111827")

    def set_debug_workspace_enabled(self, enabled: bool):
        self.debug_workspace_enabled = enabled
        state = "normal" if enabled else "disabled"
        for widget in self.debug_controls:
            try:
                widget.configure(state=state)
            except tk.TclError:
                try:
                    widget.state(["!disabled"] if enabled else ["disabled"])
                except Exception:
                    pass
        if enabled:
            self.status_text.set("模板设置区已启用：先选左侧项目，再点击右侧模板预览。")
        else:
            self.status_text.set("请先点击“设置模板”启用模板设置区。")

    def finish_debug(self):
        if not self.debug_workspace_enabled:
            messagebox.showinfo("尚未开始设置", "请先点击顶部“设置模板”。")
            return
        self.sync_validation_to_mapping(invalidate=False)
        self.debug_completed = True
        self.mark_workflow_done("debug_done")
        if self.records:
            self.load_data_preview()
        self.log("模板设置已确认：可以进入“导出前预览”。")
        messagebox.showinfo("设置完成", "当前模板标注已确认，可以点击“导出前预览”。")

    def mark_debug_dirty(self):
        self.debug_completed = False
        self.preview_ready = False
        self.preview_path = None
        self.reset_workflow_after("debug_done", "preview", "export")

    def push_undo_state(self, label: str):
        self.undo_stack.append(
            {
                "label": label,
                "mapping": copy.deepcopy(self.mapping),
                "pending_pair_index": copy.deepcopy(self.pending_pair_index),
                "mark_offset_x": self.mark_offset_x,
                "mark_offset_y": self.mark_offset_y,
                "selected_mark_ref": copy.deepcopy(self.selected_mark_ref),
                "select_mark_for_adjust": self.select_mark_for_adjust,
            }
        )
        if len(self.undo_stack) > 50:
            self.undo_stack.pop(0)

    def undo_last_action(self):
        if self.undo_stack:
            state = self.undo_stack.pop()
            self.mapping = copy.deepcopy(state["mapping"])
            self.pending_pair_index = copy.deepcopy(state["pending_pair_index"])
            self.mark_offset_x = int(state.get("mark_offset_x") or 0)
            self.mark_offset_y = int(state.get("mark_offset_y") or 0)
            self.selected_mark_ref = copy.deepcopy(state.get("selected_mark_ref"))
            self.select_mark_for_adjust = bool(state.get("select_mark_for_adjust"))
            self.mark_debug_dirty()
            self.sync_validation_to_mapping(invalidate=False)
            self.refresh_mapping_tree()
            self.render_table_preview()
            if self.records:
                self.load_data_preview()
            self.status_text.set(f"已撤销：{state.get('label') or '上一步操作'}")
            return
        if self.preview_ready:
            self.preview_ready = False
            self.preview_path = None
            self.reset_workflow_after("preview", "export")
            self.status_text.set("已回到导出前预览之前。")
            return
        if self.debug_completed:
            self.debug_completed = False
            self.set_debug_workspace_enabled(True)
            self.reset_workflow_after("debug_done", "preview", "export")
            self.status_text.set("已回到模板设置状态。")
            return
        messagebox.showinfo("没有可撤销操作", "当前没有可以撤销的模板设置操作。")

    def reset_app_state(self):
        if not messagebox.askyesno("确认重置", "确定清空当前模板、数据、标注和预览吗？\n\n已经导出的文件不会被删除。"):
            return
        self.template_path.set("")
        self.data_path.set("")
        self.output_dir.set(str(BASE_DIR / "output"))
        self.filename_prefix.set("")
        self.brush_mode.set("judgment")
        self.custom_option.set("")
        self.choice_mode.set("single")
        self.validation_mode.set("range")
        self.min_count.set(1)
        self.max_count.set(11)
        self.exact_count.set(1)
        self.export_mode.set("multi")
        self.clean_mode.set(False)
        self.mark_horizontal.set("居中")
        self.mark_vertical.set("居中")
        self.mark_font_size.set(10)
        self.mark_offset_x = 0
        self.mark_offset_y = 0
        self.mapping = blank_mapping()
        self.pending_pair_index = {}
        self.tables = []
        self.records = []
        self.preview_cells = []
        self.preview_image_tk = None
        self.preview_ready = False
        self.preview_path = None
        self.workflow_done.clear()
        self.debug_completed = False
        self.select_mark_for_adjust = False
        self.selected_mark_ref = None
        self.undo_stack.clear()
        self.selected_key = None
        self.last_result_selection = []
        self.template_page_count = None
        self.result_group_text.set("多结果标记：未选择")
        if self.field_list is not None:
            self.field_list.selection_clear(0, "end")
        if getattr(self, "option_list", None) is not None:
            self.option_list.configure(state="normal")
            self.option_list.delete(0, "end")
            self.option_list.configure(state="disabled")
        self.option_keys = []
        self.option_labels = {}
        if getattr(self, "data_tree", None) is not None:
            for item in self.data_tree.get_children():
                self.data_tree.delete(item)
        if getattr(self, "log_text", None) is not None:
            self.log_text.delete("1.0", "end")
        self.reset_workflow_after("template", "data", "debug", "debug_done", "preview", "export")
        self.set_debug_workspace_enabled(False)
        self.refresh_mapping_tree()
        self.render_table_preview()
        self.update_file_status()
        self.status_text.set("已重置。请重新上传模板和数据文件。")
        self.log("已重置当前配置。")

    @staticmethod
    def bind_listbox_mousewheel(widget):
        def on_mousewheel(event):
            if getattr(event, "num", None) == 4:
                units = -3
            elif getattr(event, "num", None) == 5:
                units = 3
            else:
                units = -3 if event.delta > 0 else 3
            widget.yview_scroll(units, "units")
            return "break"

        widget.bind("<MouseWheel>", on_mousewheel)
        widget.bind("<Button-4>", on_mousewheel)
        widget.bind("<Button-5>", on_mousewheel)

    def _build_data_panel(self):
        split = tk.PanedWindow(
            self.data_panel,
            orient=tk.HORIZONTAL,
            sashwidth=8,
            sashrelief="raised",
            showhandle=True,
            handlesize=16,
            bg="#cbd5e1",
            bd=0,
            opaqueresize=True,
        )
        split.pack(fill="both", expand=True)

        preview = ttk.LabelFrame(split, text="数据预览（拖动中间分隔条调整）", style="Panel.TLabelframe", padding=8)
        log_frame = ttk.LabelFrame(split, text="操作记录（拖动中间分隔条调整）", style="Panel.TLabelframe", padding=8)
        split.add(preview, minsize=420, stretch="always")
        split.add(log_frame, minsize=260)

        columns = ("room", "name", "phone", "options", "status")
        self.data_tree = ttk.Treeview(preview, columns=columns, show="headings", height=8)
        for col, title, width in [
            ("room", "房号/地址", 125),
            ("name", "姓名", 95),
            ("phone", "电话号码", 135),
            ("options", "投票结果", 330),
            ("status", "检查结果", 330),
        ]:
            self.data_tree.heading(col, text=title)
            self.data_tree.column(col, width=width, anchor="center" if col not in ("options", "status") else "w")
        data_scroll = ttk.Scrollbar(preview, orient="vertical", command=self.data_tree.yview)
        self.data_tree.configure(yscrollcommand=data_scroll.set)
        self.data_tree.pack(side="left", fill="both", expand=True)
        data_scroll.pack(side="right", fill="y")
        self.bind_listbox_mousewheel(self.data_tree)

        self.log_text = tk.Text(log_frame, height=8, wrap="word", font=("Consolas", 10), background="#111827", foreground="#e5e7eb", insertbackground="#e5e7eb")
        self.log_text.pack(fill="both", expand=True)
        self.bind_listbox_mousewheel(self.log_text)

    def _build_mapper_tab(self):
        paned = ttk.PanedWindow(self.tab_mapper, orient="horizontal")
        paned.pack(fill="both", expand=True)

        left_scroll = ScrollFrame(paned)
        left = left_scroll.inner
        left.configure(padding=(0, 0, 10, 0))
        right = ttk.Frame(paned)
        left_scroll.configure(width=300)
        paned.add(left_scroll, weight=1)
        paned.add(right, weight=4)
        # Apply the compact width after the window has received its real size.
        # With a zero-weight pane Tk can collapse the entire control column to
        # one pixel during the next geometry pass.
        self.after(100, lambda: paned.sashpos(0, 300))

        field_box = ttk.LabelFrame(left, text="用户信息填入", padding=5)
        field_box.pack(fill="x", expand=False)
        self.field_list = tk.Listbox(field_box, height=5, selectmode=tk.SINGLE, exportselection=False, font=("Microsoft YaHei UI", 9))
        self.field_list.pack(fill="x", expand=False)
        for key in self.field_keys:
            field = key.split(":", 1)[1]
            self.field_list.insert("end", FIELD_LABELS.get(field, field))
        self.field_list.bind("<<ListboxSelect>>", self.on_field_selected)
        self.bind_listbox_mousewheel(self.field_list)

        brush_box = ttk.LabelFrame(left, text="格式刷", padding=5)
        brush_box.pack(fill="x", pady=(6, 0))
        judgment_radio = ttk.Radiobutton(brush_box, text="判断区", value="judgment", variable=self.brush_mode, command=lambda: self.set_brush("judgment"))
        judgment_radio.pack(anchor="w")
        mark_radio = ttk.Radiobutton(brush_box, text="标记区 √", value="mark", variable=self.brush_mode, command=lambda: self.set_brush("mark"))
        mark_radio.pack(anchor="w")
        ttk.Label(brush_box, textvariable=self.status_text, wraplength=255, foreground="#374151", font=("Microsoft YaHei UI", 9)).pack(anchor="w", pady=(6, 0))
        ttk.Label(brush_box, text="先按顺序点击判断区，再切到标记区按同样顺序点击打勾位置，一一对应。", wraplength=255, foreground="#374151", font=("Microsoft YaHei UI", 9)).pack(anchor="w", pady=(4, 0))

        option_box = ttk.LabelFrame(left, text="投票结果", padding=5)
        option_box.pack(fill="both", expand=False, pady=(6, 0))
        ttk.Label(option_box, textvariable=self.result_group_text, foreground="#2563eb").pack(anchor="w", pady=(0, 4))
        self.option_list = tk.Listbox(option_box, height=7, selectmode=tk.EXTENDED, exportselection=False, font=("Microsoft YaHei UI", 9))
        self.option_list.pack(fill="both", expand=True)
        self.option_list.bind("<<ListboxSelect>>", self.on_option_selected)
        left_scroll.bind_outer_mousewheel(self.option_list)
        self.reset_project_list([])

        adjust_box = ttk.LabelFrame(left, text="打勾位置微调", padding=5)
        adjust_box.pack(fill="x", pady=(6, 0))
        pad = 2
        select_mark_button = ttk.Button(adjust_box, text="选择打勾位置", command=self.start_select_mark_for_adjust)
        up_button = ttk.Button(adjust_box, text="↑", width=4, command=lambda: self.nudge_mark(0, -1))
        left_button = ttk.Button(adjust_box, text="←", width=4, command=lambda: self.nudge_mark(-1, 0))
        right_button = ttk.Button(adjust_box, text="→", width=4, command=lambda: self.nudge_mark(1, 0))
        down_button = ttk.Button(adjust_box, text="↓", width=4, command=lambda: self.nudge_mark(0, 1))
        select_mark_button.grid(row=0, column=0, columnspan=3, sticky="ew", padx=pad, pady=(0, 6))
        up_button.grid(row=1, column=1, padx=pad, pady=pad)
        left_button.grid(row=2, column=0, padx=pad, pady=pad)
        right_button.grid(row=2, column=2, padx=pad, pady=pad)
        down_button.grid(row=3, column=1, padx=pad, pady=pad)
        ttk.Label(adjust_box, text="先选一个蓝色打勾格，再用按钮/方向键/拖动微调，只影响该位置。", wraplength=255, foreground="#374151", font=("Microsoft YaHei UI", 9)).grid(row=4, column=0, columnspan=3, sticky="w", pady=(6, 0))

        status_box = ttk.LabelFrame(left, text="当前标注", padding=5)
        status_box.pack(fill="both", expand=True, pady=(6, 0))
        self.mapping_tree = ttk.Treeview(status_box, columns=("name", "judgment", "mark"), show="headings", height=12)
        self.mapping_tree.heading("name", text="项目")
        self.mapping_tree.heading("judgment", text="判断区")
        self.mapping_tree.heading("mark", text="标记区")
        self.mapping_tree.column("name", width=90, minwidth=70)
        self.mapping_tree.column("judgment", width=82, minwidth=65)
        self.mapping_tree.column("mark", width=82, minwidth=65)
        self.mapping_tree.pack(fill="both", expand=True)

        map_actions = ttk.Frame(left)
        map_actions.pack(fill="x", pady=(6, 0))
        clear_selected_button = ttk.Button(map_actions, text="清空当前字段/结果", command=self.clear_selected_mapping)
        clear_selected_button.grid(row=0, column=0, sticky="w", padx=(0, 6))
        clear_button = ttk.Button(map_actions, text="手动清空标注", command=self.clear_all_mappings)
        clear_button.grid(row=0, column=1, sticky="w")
        ttk.Label(
            map_actions,
            text="Ctrl/Shift 可多选，点击预览图会同步标注所选结果。",
            wraplength=255,
            font=("Microsoft YaHei UI", 9),
        ).grid(row=1, column=0, columnspan=2, sticky="w", pady=(4, 0))

        right_header = ttk.Frame(right)
        right_header.pack(fill="x", pady=(0, 8))
        ttk.Label(right_header, text="模板预览图：选左侧结果，判断区和标记区按点击顺序一一对应").pack(side="left")

        self.table_frame = ScrollFrame(right)
        self.table_frame.pack(fill="both", expand=True)
        self.debug_controls = list(self.debug_controls) + [
            self.field_list,
            self.option_list,
            judgment_radio,
            mark_radio,
            select_mark_button,
            up_button,
            left_button,
            right_button,
            down_button,
            self.mapping_tree,
            clear_selected_button,
            clear_button,
        ]
        left_scroll.bind_mousewheel_to_descendants(exclude=(tk.Listbox, ttk.Treeview, tk.Text, tk.Canvas))
        self.refresh_mapping_tree()

    def choose_template(self):
        path = filedialog.askopenfilename(title="选择 Word 模板", filetypes=[("Word 文件", "*.docx *.doc"), ("全部文件", "*.*")])
        if not path:
            return
        try:
            docx_path = convert_doc_to_docx(path, BASE_DIR / "converted")
            self.template_path.set(str(docx_path))
            self.mapping = blank_mapping()
            self.pending_pair_index = {}
            self.mark_offset_x = 0
            self.mark_offset_y = 0
            self.selected_mark_ref = None
            self.select_mark_for_adjust = False
            self.template_page_count = docx_page_count(docx_path)
            self.debug_completed = False
            self.set_debug_workspace_enabled(False)
            self.invalidate_preview()
            self.reset_workflow_after("debug", "debug_done", "preview", "export")
            self.mark_workflow_done("template")
            self.update_file_status()
            self.log(f"已选择模板：{docx_path}")
            self.show_template_page_check()
            self.load_template_tables()
            self.refresh_mapping_tree()
        except Exception as exc:
            messagebox.showerror("模板错误", str(exc))

    def show_template_page_check(self):
        count = self.template_page_count
        if count is None:
            messagebox.showinfo(
                "模板页数检查",
                "未能从模板属性中读取页数。\n\n如果后续选择“单文件导出”，请先在 Word 中确认这个模板尽量只占 1 页；如果超过 1 页，程序会从下一条数据开始另起新页追加。",
            )
            return
        if count > 1:
            messagebox.showwarning(
                "模板页数检查",
                f"检测到模板可能为 {count} 页。\n\n建议把模板调整到 1 页后再使用单文件导出；如果强行单文件导出，下一条数据会从第 {count + 1} 页之后开始追加。",
            )
        else:
            messagebox.showinfo("模板页数检查", "检测到模板属性为 1 页。单文件导出前仍建议在 Word 中快速确认一次。")

    def choose_data(self):
        path = filedialog.askopenfilename(title="选择投票数据", filetypes=[("数据文件", "*.xlsx *.xlsm *.csv *.txt"), ("全部文件", "*.*")])
        if path:
            self.data_path.set(path)
            if not self.prompt_vote_count_rule(path):
                self.data_path.set("")
                self.update_file_status()
                return
            self.debug_completed = False
            self.selected_mark_ref = None
            self.select_mark_for_adjust = False
            self.set_debug_workspace_enabled(False)
            self.invalidate_preview()
            self.reset_workflow_after("debug", "debug_done", "preview", "export")
            self.update_file_status()
            self.load_data_preview()

    def prompt_vote_count_rule(self, data_path: str) -> bool:
        try:
            preview_records = read_vote_records(data_path, result_count=0, normalize_options=False, dedupe_options=False)
            default_count = 1
            if preview_records:
                first = preview_records[0]
                default_count = len(first.result_options) if first.result_options else len(first.options)
                default_count = max(1, default_count)
        except Exception as exc:
            messagebox.showerror("数据错误", str(exc))
            return False

        dialog = tk.Toplevel(self)
        dialog.title("投票结果数量检查")
        dialog.geometry("360x210")
        dialog.resizable(False, False)
        dialog.transient(self)
        dialog.grab_set()

        mode = tk.StringVar(value=self.validation_mode.get() or "exact")
        exact = tk.IntVar(value=self.safe_int_var(self.exact_count) or default_count)
        min_value = tk.IntVar(value=self.safe_int_var(self.min_count) or 1)
        max_value = tk.IntVar(value=self.safe_int_var(self.max_count) or max(default_count, 11))
        accepted = {"value": False}

        frame = ttk.Frame(dialog, padding=14)
        frame.pack(fill="both", expand=True)
        ttk.Label(frame, text="请选择每条数据允许的投票结果数量：").grid(row=0, column=0, columnspan=4, sticky="w", pady=(0, 12))
        ttk.Radiobutton(frame, text="精确型", variable=mode, value="exact").grid(row=1, column=0, sticky="w", pady=5)
        ttk.Spinbox(frame, from_=1, to=99, textvariable=exact, width=6).grid(row=1, column=1, sticky="w")
        ttk.Label(frame, text="个").grid(row=1, column=2, sticky="w")
        ttk.Radiobutton(frame, text="范围型", variable=mode, value="range").grid(row=2, column=0, sticky="w", pady=5)
        ttk.Spinbox(frame, from_=0, to=99, textvariable=min_value, width=6).grid(row=2, column=1, sticky="w")
        ttk.Label(frame, text="-").grid(row=2, column=2, sticky="w", padx=4)
        ttk.Spinbox(frame, from_=1, to=99, textvariable=max_value, width=6).grid(row=2, column=3, sticky="w")
        ttk.Label(frame, text="范围包含两端，例如 1-11 个。").grid(row=3, column=0, columnspan=4, sticky="w", pady=(8, 0))

        footer = ttk.Frame(frame)
        footer.grid(row=4, column=0, columnspan=4, sticky="e", pady=(18, 0))

        def confirm():
            if mode.get() == "exact" and int(exact.get() or 0) <= 0:
                messagebox.showwarning("数量错误", "精确数量必须大于 0。")
                return
            if mode.get() == "range":
                lo = int(min_value.get() or 0)
                hi = int(max_value.get() or 0)
                if hi <= 0 or lo > hi:
                    messagebox.showwarning("范围错误", "范围最大值必须大于 0，且最小值不能大于最大值。")
                    return
            self.validation_mode.set(mode.get())
            self.exact_count.set(int(exact.get() or 0))
            self.min_count.set(int(min_value.get() or 0))
            self.max_count.set(int(max_value.get() or 0))
            self.sync_validation_to_mapping(invalidate=False)
            accepted["value"] = True
            dialog.destroy()

        def cancel():
            accepted["value"] = False
            dialog.destroy()

        ttk.Button(footer, text="取消", command=cancel).pack(side="right", padx=(8, 0))
        ttk.Button(footer, text="确定", command=confirm).pack(side="right")
        dialog.protocol("WM_DELETE_WINDOW", cancel)
        dialog.wait_window()
        return bool(accepted["value"])

    def choose_output_dir(self):
        path = filedialog.askdirectory(title="选择输出目录")
        if path:
            self.output_dir.set(path)
            self.invalidate_preview()
            self.update_file_status()

    def open_output_dir(self):
        path = Path(self.output_dir.get() or BASE_DIR / "output")
        try:
            path.mkdir(parents=True, exist_ok=True)
            os.startfile(str(path))
        except Exception as exc:
            messagebox.showerror("打开失败", f"无法打开输出目录：\n{exc}")

    def debug_template(self):
        if not self.template_path.get():
            self.choose_template()
            if not self.template_path.get():
                return
        if not self.data_path.get():
            self.choose_data()
            if not self.data_path.get():
                return
        self.load_template_tables()
        self.refresh_mapping_tree()
        self.render_table_preview()
        self.set_debug_workspace_enabled(True)
        self.debug_completed = False
        self.preview_ready = False
        self.preview_path = None
        self.mark_workflow_done("debug")
        self.reset_workflow_after("debug_done", "preview", "export")
        self.log("模板设置已打开：完成后请点击顶部“4 确认设置”。")

    def start_export(self):
        if not self.preview_ready or not self.preview_path:
            messagebox.showinfo("请先确认预览", "请先点击“导出前预览”，在程序内部预览窗口确认无误后再开始导出。")
            return
        if not self.prompt_export_settings():
            return
        self.sync_validation_to_mapping(invalidate=False)
        self.export_all()

    def prompt_export_settings(self) -> bool:
        dialog = tk.Toplevel(self)
        dialog.title("导出设置")
        dialog.geometry("420x270")
        dialog.resizable(False, False)
        dialog.transient(self)
        dialog.grab_set()

        mode = tk.StringVar(value=self.export_mode.get() or "multi")
        clean = tk.BooleanVar(value=bool(self.clean_mode.get()))
        accepted = {"value": False}

        frame = ttk.Frame(dialog, padding=16)
        frame.pack(fill="both", expand=True)
        ttk.Label(frame, text="请选择本次导出方式：", font=("Microsoft YaHei UI", 10, "bold")).pack(anchor="w", pady=(0, 10))
        ttk.Radiobutton(frame, text="多文件：一条数据生成一份 DOCX", variable=mode, value="multi").pack(anchor="w", pady=4)
        ttk.Radiobutton(frame, text="单文件：所有数据合并到一个 DOCX，每条数据从新页开始", variable=mode, value="single").pack(anchor="w", pady=4)
        ttk.Checkbutton(frame, text="纯净模式：只打印用户信息和打勾，隐藏模板原有内容", variable=clean).pack(anchor="w", pady=(12, 6))

        page_count = self.template_page_count
        if page_count is None:
            page_text = "模板页数：未读取到，请确认模板最好只占 1 页。"
        else:
            page_text = f"模板页数：检测到 {page_count} 页。"
        ttk.Label(frame, text=page_text, foreground="#6b7280", wraplength=370).pack(anchor="w", pady=(4, 0))

        footer = ttk.Frame(frame)
        footer.pack(side="bottom", fill="x", pady=(18, 0))

        def confirm():
            if mode.get() == "single":
                if page_count is None:
                    if not messagebox.askyesno(
                        "确认单文件导出",
                        "未能读取模板页数。\n\n请确认模板在 Word 中尽量只占 1 页。仍然使用单文件导出吗？",
                        parent=dialog,
                    ):
                        return
                elif page_count > 1:
                    if not messagebox.askyesno(
                        "确认单文件导出",
                        f"检测到模板可能为 {page_count} 页。\n\n单文件导出会让下一条数据从上一条完整内容后另起新页追加。仍然继续吗？",
                        parent=dialog,
                    ):
                        return
            self.export_mode.set(mode.get())
            self.clean_mode.set(bool(clean.get()))
            accepted["value"] = True
            dialog.destroy()

        def cancel():
            accepted["value"] = False
            dialog.destroy()

        ttk.Button(footer, text="取消", command=cancel).pack(side="right", padx=(8, 0))
        ttk.Button(footer, text="开始导出", command=confirm).pack(side="right")
        dialog.protocol("WM_DELETE_WINDOW", cancel)
        dialog.wait_window()
        return bool(accepted["value"])

    def load_data_preview(self):
        if not self.data_path.get():
            return
        try:
            self.sync_validation_to_mapping(invalidate=False)
            self.records = read_vote_records(self.data_path.get(), result_count=0, **record_parse_options(self.mapping))
            self.reset_project_list(self.unique_vote_options())
            self.refresh_mapping_tree()
            for item in self.data_tree.get_children():
                self.data_tree.delete(item)
            for record in self.records[:200]:
                reasons = validate_vote_record(record, self.mapping)
                if reasons:
                    status = "异常：" + "；".join(reasons)
                else:
                    status = "正常"
                option_text = "、".join(self.display_option_key(option) for option in record.options)
                self.data_tree.insert("", "end", values=(record.room, record.name, record.phone, option_text, status))
            self.log(f"已读取数据：{len(self.records)} 行。")
            self.mark_workflow_done("data")
        except Exception as exc:
            messagebox.showerror("数据错误", str(exc))

    def update_file_status(self):
        template = Path(self.template_path.get()).name if self.template_path.get() else "未上传模板"
        data = Path(self.data_path.get()).name if self.data_path.get() else "未上传数据"
        output = self.output_dir.get() or "未选择输出目录"
        prefix = self.filename_prefix.get().strip()
        prefix_text = prefix if prefix else "无"
        self.file_status_text.set(f"模板：{template}    数据：{data}    输出：{output}    文件名前缀：{prefix_text}")

    def on_filename_prefix_changed(self):
        self.mapping["filenamePrefix"] = self.filename_prefix.get().strip()
        self.invalidate_preview()
        self.update_file_status()

    def unique_vote_options(self) -> List[str]:
        if not self.records:
            return []
        if self.records[0].result_options:
            result_keys = sorted(self.records[0].result_options.keys(), key=self.option_sort_key)
        else:
            result_keys = sorted([option for option in self.records[0].options if option], key=self.option_sort_key)
        return result_keys

    def load_template_tables(self):
        if not self.template_path.get():
            return
        try:
            self.tables = document_tables(self.template_path.get())
            self.render_table_preview()
            self.log(f"已读取模板表格：{len(self.tables)} 个。")
        except Exception as exc:
            messagebox.showerror("模板读取失败", str(exc))

    def render_table_preview(self):
        for child in self.table_frame.inner.winfo_children():
            child.destroy()
        self.preview_cells = []

        if not self.template_path.get():
            ttk.Label(self.table_frame.inner, text="请先选择 Word 模板。").grid(row=0, column=0, sticky="w", padx=8, pady=8)
            return

        try:
            image, cells = self.build_print_preview_image()
        except Exception as exc:
            ttk.Label(self.table_frame.inner, text=f"模板预览生成失败：{exc}").grid(row=0, column=0, sticky="w", padx=8, pady=8)
            return

        self.preview_cells = cells
        self.preview_image_tk = ImageTk.PhotoImage(image)
        preview = tk.Label(
            self.table_frame.inner,
            image=self.preview_image_tk,
            background="#f7f8fa",
            cursor="crosshair",
            borderwidth=0,
            highlightthickness=0,
            takefocus=1,
        )
        self.preview_widget = preview
        preview.grid(row=0, column=0, sticky="nw", padx=14, pady=14)
        preview.bind("<ButtonPress-1>", self.on_preview_mouse_down)
        preview.bind("<B1-Motion>", self.on_preview_mouse_drag)
        preview.bind("<ButtonRelease-1>", self.on_preview_mouse_up)
        self.table_frame.bind_mousewheel(preview)
        preview.bind("<Left>", lambda _event: self.nudge_mark(-1, 0))
        preview.bind("<Right>", lambda _event: self.nudge_mark(1, 0))
        preview.bind("<Up>", lambda _event: self.nudge_mark(0, -1))
        preview.bind("<Down>", lambda _event: self.nudge_mark(0, 1))

    def build_print_preview_image(self) -> Tuple[Image.Image, List[Dict[str, Any]]]:
        return self.build_print_preview_image_for_path(self.template_path.get(), highlight=True)

    def build_print_preview_image_for_path(self, docx_path: str | Path, highlight: bool = True) -> Tuple[Image.Image, List[Dict[str, Any]]]:
        document = Document(docx_path)
        page_width = 900
        page_height = 6000
        margin_x = 64
        margin_y = 56
        usable_width = page_width - margin_x * 2
        image = Image.new("RGB", (page_width, page_height), "#ffffff")
        draw = ImageDraw.Draw(image)
        font_body = self.preview_font(17)
        font_small = self.preview_font(15)
        font_title = self.preview_font(24, bold=True)
        font_mark = self.preview_font(15, bold=True)

        def draw_field_preview_value(target: Dict[str, Any], x1: int, y1: int, x2: int, y2: int, font=None) -> None:
            value = self.field_preview_value_for_target(target)
            if not value:
                return
            value_font = font or font_mark
            value_width = self.text_width(draw, value, value_font)
            value_height = self.text_height(draw, value, value_font)
            value_x = x1 + max(2, (x2 - x1 - value_width) // 2)
            value_y = y1 + max(1, (y2 - y1 - value_height) // 2)
            draw.text((value_x, value_y), value, fill="#166534", font=value_font)

        def draw_role_badge(target: Dict[str, Any], role: str, x: int, y: int, prefix: str = "") -> None:
            label = self.target_role_labels(target, role)
            if not label:
                return
            text = f"{prefix}{label}" if prefix else label
            fill = "#92400e" if role == "judgment" else "#1d4ed8"
            draw.text((x, y), text, fill=fill, font=font_mark)

        def draw_target_overlay(target: Dict[str, Any], x1: int, y1: int, x2: int, y2: int, draw_line: bool = False) -> None:
            roles = self.target_roles(target) if highlight else []
            mapped = bool(roles)
            if "mark" in roles:
                fill = "#dbeafe"
                outline = "#2563eb"
            elif "judgment" in roles:
                fill = "#fef3c7"
                outline = "#d97706"
            elif "field" in roles:
                fill = "#dcfce7"
                outline = "#16a34a"
            else:
                fill = "#ffffff"
                outline = "#9ca3af"
            selected_mark = "mark" in roles and self.selected_mark_target() == target
            if draw_line:
                line_color = "#dc2626" if selected_mark else (outline if mapped else "#111827")
                draw.line((x1, y2 - 4, x2, y2 - 4), fill=line_color, width=3 if mapped else 1)
                if "judgment" in roles:
                    draw_role_badge(target, "judgment", x1 + 4, y1 + 1, "判 ")
                if "field" in roles:
                    draw_field_preview_value(target, x1, y1, x2, y2)
                if "mark" in roles:
                    draw_role_badge(target, "mark", x1 + 4, y1 + 1)
                    style = self.mark_style_for_target(target)
                    try:
                        mark_font = self.preview_font(int(style.get("fontSize") or 10), bold=True)
                        offset_x = int(style.get("offsetX") or 0)
                        offset_y = int(style.get("offsetY") or 0)
                    except Exception:
                        mark_font = font_title
                        offset_x = 0
                        offset_y = 0
                    check = "√"
                    check_width = self.text_width(draw, check, mark_font)
                    check_height = self.text_height(draw, check, mark_font)
                    check_x = x1 + (x2 - x1 - check_width) // 2 + offset_x
                    check_y = y1 + (y2 - y1 - check_height) // 2 + offset_y
                    draw.text((check_x, check_y), check, fill="#dc2626" if selected_mark else "#1d4ed8", font=mark_font)
                return
            if mapped:
                draw.rectangle((x1, y1, x2, y2), fill=fill, outline="#dc2626" if selected_mark else outline, width=3 if selected_mark else 2)
            elif highlight:
                draw.rectangle((x1, y1, x2, y2), outline=outline, width=1)
            if "judgment" in roles:
                draw_role_badge(target, "judgment", x1 + 4, y1 + 1, "判 ")
            if "field" in roles:
                draw_field_preview_value(target, x1, y1, x2, y2)
            if "mark" in roles:
                draw_role_badge(target, "mark", x1 + 4, y1 + 1)
                style = self.mark_style_for_target(target)
                try:
                    mark_font = self.preview_font(int(style.get("fontSize") or 10), bold=True)
                    offset_x = int(style.get("offsetX") or 0)
                    offset_y = int(style.get("offsetY") or 0)
                except Exception:
                    mark_font = font_title
                    offset_x = 0
                    offset_y = 0
                check = "√"
                check_width = self.text_width(draw, check, mark_font)
                check_height = self.text_height(draw, check, mark_font)
                check_x = x1 + (x2 - x1 - check_width) // 2 + offset_x
                check_y = y1 + (y2 - y1 - check_height) // 2 + offset_y
                draw.text((check_x, check_y), check, fill="#dc2626" if selected_mark else "#1d4ed8", font=mark_font)

        draw.rectangle((0, 0, page_width - 1, page_height - 1), outline="#d1d5db", width=2)
        y = margin_y
        table_index = 0
        paragraph_index = 0
        cells: List[Dict[str, Any]] = []

        for block in self.iter_doc_blocks(document):
            if isinstance(block, Paragraph):
                raw_text = block.text.replace("\r", "")
                text = raw_text.strip()
                run_spans = formatted_underline_spans(block, {"kind": "underline", "paragraph": paragraph_index})
                if not text and not run_spans:
                    y += 9
                    paragraph_index += 1
                    continue
                is_title = y < 150 and len(text) <= 48
                font = font_title if is_title else font_body
                preview_text = raw_text if run_spans else text
                lines = self.wrap_preview_text_with_offsets(draw, preview_text, font, usable_width)
                underline_index = 0
                for line_info in lines:
                    line = line_info["text"]
                    line_start = int(line_info["start"])
                    line_end = line_start + len(line)
                    line_width = self.text_width(draw, line, font)
                    x = margin_x + (usable_width - line_width) // 2 if is_title else margin_x
                    line_height = self.text_height(draw, line, font)
                    for match in UNDERLINE_PATTERN.finditer(line):
                        target = {"kind": "underline", "paragraph": paragraph_index, "underline": underline_index}
                        prefix = line[:match.start()]
                        matched = line[match.start():match.end()]
                        ux1 = x + self.text_width(draw, prefix, font)
                        ux2 = ux1 + max(36, self.text_width(draw, matched, font))
                        uy1 = y
                        uy2 = y + line_height + 6
                        roles = self.target_roles(target) if highlight else []
                        mapped = bool(roles)
                        if "mark" in roles:
                            fill = "#dbeafe"
                            outline = "#2563eb"
                        elif "judgment" in roles:
                            fill = "#fef3c7"
                            outline = "#d97706"
                        elif "field" in roles:
                            fill = "#dcfce7"
                            outline = "#16a34a"
                        else:
                            fill = "#ffffff"
                            outline = "#9ca3af"
                        selected_mark = "mark" in roles and self.selected_mark_target() == target
                        if mapped:
                            draw.rectangle((ux1, uy1, ux2, uy2), fill=fill, outline="#dc2626" if selected_mark else outline, width=3 if selected_mark else 2)
                        elif highlight:
                            draw.rectangle((ux1, uy1, ux2, uy2), outline=outline, width=1)
                        if "judgment" in roles:
                            draw_role_badge(target, "judgment", ux1 + 4, uy1 + 1, "判 ")
                        if "field" in roles:
                            draw_field_preview_value(target, ux1, uy1, ux2, uy2)
                        if "mark" in roles:
                            draw_role_badge(target, "mark", ux1 + 4, uy1 + 1)
                            style = self.mark_style_for_target(target)
                            try:
                                mark_font = self.preview_font(int(style.get("fontSize") or 10), bold=True)
                                offset_x = int(style.get("offsetX") or 0)
                                offset_y = int(style.get("offsetY") or 0)
                            except Exception:
                                mark_font = font_title
                                offset_x = 0
                                offset_y = 0
                            check = "√"
                            check_width = self.text_width(draw, check, mark_font)
                            check_height = self.text_height(draw, check, mark_font)
                            check_x = ux1 + (ux2 - ux1 - check_width) // 2 + offset_x
                            check_y = uy1 + (uy2 - uy1 - check_height) // 2 + offset_y
                            draw.text((check_x, check_y), check, fill="#dc2626" if selected_mark else "#1d4ed8", font=mark_font)
                        context = UNDERLINE_PATTERN.sub("", line).strip() or text
                        cells.append(
                            {
                                "target": target,
                                "table": -1,
                                "row": paragraph_index,
                                "col": underline_index,
                                "text": context,
                                "x1": ux1,
                                "y1": uy1,
                                "x2": ux2,
                                "y2": uy2,
                            }
                        )
                        underline_index += 1
                    for span in run_spans:
                        span_start = int(span["start"])
                        span_end = int(span["end"])
                        if span_end < line_start or span_start > line_end:
                            continue
                        local_start = max(0, span_start - line_start)
                        local_end = min(len(line), span_end - line_start)
                        prefix = line[:local_start]
                        matched = line[local_start:local_end] or str(span.get("text") or "")
                        ux1 = x + self.text_width(draw, prefix, font)
                        ux2 = ux1 + max(52, self.text_width(draw, matched, font))
                        uy1 = y
                        uy2 = y + line_height + 8
                        target = dict(span["target"])
                        draw_target_overlay(target, ux1, uy1, ux2, uy2, draw_line=True)
                        cells.append(
                            {
                                "target": target,
                                "table": -1,
                                "row": paragraph_index,
                                "col": int(target.get("underline", 0)),
                                "text": text or raw_text,
                                "x1": ux1,
                                "y1": uy1,
                                "x2": ux2,
                                "y2": uy2,
                            }
                        )
                    draw.text((x, y), line, fill="#111827", font=font)
                    y += line_height + (7 if is_title else 5)
                y += 3
                paragraph_index += 1
                continue

            if isinstance(block, Table):
                rows_text = [[cell.text.strip() for cell in row.cells] for row in block.rows]
                if not rows_text:
                    continue
                max_cols = max((len(row) for row in rows_text), default=0)
                if max_cols <= 0:
                    continue
                y += 8
                col_width = max(72, usable_width // max_cols)
                table_width = col_width * max_cols
                table_x = margin_x + max(0, (usable_width - table_width) // 2)

                for row_index, row in enumerate(rows_text):
                    wrapped_by_col: List[List[Dict[str, Any]]] = []
                    for col_index in range(max_cols):
                        text = row[col_index] if col_index < len(row) else ""
                        lines = self.wrap_preview_text_with_offsets(draw, text or " ", font_small, col_width - 16, max_lines=4)
                        wrapped_by_col.append(lines)
                    row_height = max(38, max(len(lines) for lines in wrapped_by_col) * 20 + 18)

                    for col_index in range(max_cols):
                        x1 = table_x + col_index * col_width
                        y1 = y
                        x2 = x1 + col_width
                        y2 = y + row_height
                        text = row[col_index] if col_index < len(row) else ""
                        target = {"table": table_index, "row": row_index, "col": col_index}
                        roles = self.target_roles(target) if highlight else []
                        mapped = bool(roles)
                        if "mark" in roles:
                            fill = "#dbeafe"
                            outline = "#2563eb"
                        elif "judgment" in roles:
                            fill = "#fef3c7"
                            outline = "#d97706"
                        elif "field" in roles:
                            fill = "#dcfce7"
                            outline = "#16a34a"
                        else:
                            fill = "#ffffff"
                            outline = "#111827"
                        selected_mark = "mark" in roles and self.selected_mark_target() == target
                        draw.rectangle((x1, y1, x2, y2), fill=fill, outline="#dc2626" if selected_mark else outline, width=3 if selected_mark else (2 if mapped else 1))
                        if "judgment" in roles:
                            draw_role_badge(target, "judgment", x1 + 6, y1 + 4, "判 ")
                        if "field" in roles:
                            draw_field_preview_value(target, x1, y1, x2, y2)
                        if "mark" in roles:
                            draw_role_badge(target, "mark", x1 + 6, y1 + 4)
                            style = self.mark_style_for_target(target)
                            try:
                                mark_font = self.preview_font(int(style.get("fontSize") or 10), bold=True)
                                offset_x = int(style.get("offsetX") or 0)
                                offset_y = int(style.get("offsetY") or 0)
                            except Exception:
                                mark_font = font_title
                                offset_x = 0
                                offset_y = 0
                            check = "√"
                            check_width = self.text_width(draw, check, mark_font)
                            check_height = self.text_height(draw, check, mark_font)
                            check_x = x1 + (col_width - check_width) // 2 + offset_x
                            check_y = y1 + (row_height - check_height) // 2 + offset_y
                            draw.text((check_x, check_y), check, fill="#dc2626" if selected_mark else "#1d4ed8", font=mark_font)

                        lines = wrapped_by_col[col_index]
                        text_height = sum(self.text_height(draw, line["text"], font_small) for line in lines) + max(0, len(lines) - 1) * 4
                        text_y = y1 + max(6, (row_height - text_height) // 2)
                        line_positions: List[Tuple[str, int, int, int, int]] = []
                        for line_info in lines:
                            line = line_info["text"]
                            line_width = self.text_width(draw, line, font_small)
                            text_x = x1 + max(8, (col_width - line_width) // 2)
                            line_height = self.text_height(draw, line, font_small)
                            line_positions.append((line, text_x, text_y, line_height, int(line_info["start"])))
                            draw.text((text_x, text_y), line, fill="#111827", font=font_small)
                            text_y += line_height + 4

                        cells.append(
                            {
                                "target": target,
                                "table": table_index,
                                "row": row_index,
                                "col": col_index,
                                "text": text,
                                "x1": x1,
                                "y1": y1,
                                "x2": x2,
                                "y2": y2,
                            }
                        )
                        if highlight and UNDERLINE_PATTERN.search(text):
                            underline_index = 0
                            underline_context = UNDERLINE_PATTERN.sub("", text).strip() or text
                            for line, text_x, line_y, line_height, _line_start in line_positions:
                                for match in UNDERLINE_PATTERN.finditer(line):
                                    underline_target = {
                                        "kind": "cellUnderline",
                                        "table": table_index,
                                        "row": row_index,
                                        "col": col_index,
                                        "paragraph": 0,
                                        "underline": underline_index,
                                    }
                                    prefix = line[: match.start()]
                                    matched = line[match.start() : match.end()]
                                    ux1 = text_x + self.text_width(draw, prefix, font_small)
                                    ux2 = ux1 + max(28, self.text_width(draw, matched, font_small))
                                    uy1 = line_y - 3
                                    uy2 = line_y + line_height + 8
                                    underline_roles = self.target_roles(underline_target)
                                    underline_mapped = bool(underline_roles)
                                    if "mark" in underline_roles:
                                        underline_fill = "#dbeafe"
                                        underline_outline = "#2563eb"
                                    elif "judgment" in underline_roles:
                                        underline_fill = "#fef3c7"
                                        underline_outline = "#d97706"
                                    elif "field" in underline_roles:
                                        underline_fill = "#dcfce7"
                                        underline_outline = "#16a34a"
                                    else:
                                        underline_fill = "#ffffff"
                                        underline_outline = "#9ca3af"
                                    underline_selected = "mark" in underline_roles and self.selected_mark_target() == underline_target
                                    if underline_mapped:
                                        draw.rectangle(
                                            (ux1, uy1, ux2, uy2),
                                            fill=underline_fill,
                                            outline="#dc2626" if underline_selected else underline_outline,
                                            width=3 if underline_selected else 2,
                                        )
                                    else:
                                        draw.rectangle((ux1, uy1, ux2, uy2), outline=underline_outline, width=1)
                                    if "judgment" in underline_roles:
                                        draw_role_badge(underline_target, "judgment", ux1 + 4, uy1 + 1, "判 ")
                                    if "field" in underline_roles:
                                        draw_field_preview_value(underline_target, ux1, uy1, ux2, uy2)
                                    if "mark" in underline_roles:
                                        draw_role_badge(underline_target, "mark", ux1 + 4, uy1 + 1)
                                        style = self.mark_style_for_target(underline_target)
                                        try:
                                            mark_font = self.preview_font(int(style.get("fontSize") or 10), bold=True)
                                            offset_x = int(style.get("offsetX") or 0)
                                            offset_y = int(style.get("offsetY") or 0)
                                        except Exception:
                                            mark_font = font_title
                                            offset_x = 0
                                            offset_y = 0
                                        check = "√"
                                        check_width = self.text_width(draw, check, mark_font)
                                        check_height = self.text_height(draw, check, mark_font)
                                        check_x = ux1 + (ux2 - ux1 - check_width) // 2 + offset_x
                                        check_y = uy1 + (uy2 - uy1 - check_height) // 2 + offset_y
                                        draw.text((check_x, check_y), check, fill="#dc2626" if underline_selected else "#1d4ed8", font=mark_font)
                                    cells.append(
                                        {
                                            "target": underline_target,
                                            "table": table_index,
                                            "row": row_index,
                                            "col": col_index,
                                            "text": underline_context,
                                            "x1": ux1,
                                            "y1": uy1,
                                            "x2": ux2,
                                            "y2": uy2,
                                        }
                                    )
                                    underline_index += 1
                        if highlight and col_index < len(block.rows[row_index].cells):
                            table_cell = block.rows[row_index].cells[col_index]
                            formatted_spans: List[Dict[str, Any]] = []
                            paragraph_offset = 0
                            for paragraph_offset_index, paragraph in enumerate(table_cell.paragraphs):
                                base_target = {
                                    "kind": "cellUnderline",
                                    "table": table_index,
                                    "row": row_index,
                                    "col": col_index,
                                    "paragraph": paragraph_offset_index,
                                }
                                for span in formatted_underline_spans(paragraph, base_target):
                                    if UNDERLINE_PATTERN.search(str(span.get("text") or "")):
                                        continue
                                    copied = dict(span)
                                    copied["start"] = int(copied["start"]) + paragraph_offset
                                    copied["end"] = int(copied["end"]) + paragraph_offset
                                    formatted_spans.append(copied)
                                paragraph_offset += len(paragraph.text) + 1
                            for span in formatted_spans:
                                span_start = int(span["start"])
                                span_end = int(span["end"])
                                for line, text_x, line_y, line_height, line_start in line_positions:
                                    line_end = line_start + len(line)
                                    if span_end < line_start or span_start > line_end:
                                        continue
                                    local_start = max(0, span_start - line_start)
                                    local_end = min(len(line), span_end - line_start)
                                    prefix = line[:local_start]
                                    matched = line[local_start:local_end] or str(span.get("text") or "")
                                    ux1 = text_x + self.text_width(draw, prefix, font_small)
                                    ux2 = ux1 + max(44, self.text_width(draw, matched, font_small))
                                    uy1 = line_y - 3
                                    uy2 = line_y + line_height + 8
                                    underline_target = dict(span["target"])
                                    draw_target_overlay(underline_target, ux1, uy1, ux2, uy2, draw_line=True)
                                    cells.append(
                                        {
                                            "target": underline_target,
                                            "table": table_index,
                                            "row": row_index,
                                            "col": col_index,
                                            "text": UNDERLINE_PATTERN.sub("", text).strip() or text,
                                            "x1": ux1,
                                            "y1": uy1,
                                            "x2": ux2,
                                            "y2": uy2,
                                        }
                                    )
                    y += row_height
                table_index += 1
                y += 18

        final_height = min(page_height, max(1100, y + margin_y))
        cropped = image.crop((0, 0, page_width, final_height))
        return cropped, cells

    @staticmethod
    def iter_doc_blocks(document):
        for child in document.element.body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, document)
            elif isinstance(child, CT_Tbl):
                yield Table(child, document)

    @staticmethod
    def preview_font(size: int, bold: bool = False):
        font_names = [
            "msyhbd.ttc" if bold else "msyh.ttc",
            "simhei.ttf" if bold else "simsun.ttc",
            "arialbd.ttf" if bold else "arial.ttf",
        ]
        for font_name in font_names:
            font_path = Path(r"C:\Windows\Fonts") / font_name
            if font_path.exists():
                try:
                    return ImageFont.truetype(str(font_path), size)
                except Exception:
                    continue
        return ImageFont.load_default()

    @staticmethod
    def text_width(draw: ImageDraw.ImageDraw, text: str, font) -> int:
        box = draw.textbbox((0, 0), text, font=font)
        return box[2] - box[0]

    @staticmethod
    def text_height(draw: ImageDraw.ImageDraw, text: str, font) -> int:
        box = draw.textbbox((0, 0), text or " ", font=font)
        return box[3] - box[1]

    def wrap_preview_text(self, draw: ImageDraw.ImageDraw, text: str, font, max_width: int, max_lines: Optional[int] = None) -> List[str]:
        return [item["text"] for item in self.wrap_preview_text_with_offsets(draw, text, font, max_width, max_lines=max_lines)]

    def wrap_preview_text_with_offsets(self, draw: ImageDraw.ImageDraw, text: str, font, max_width: int, max_lines: Optional[int] = None) -> List[Dict[str, Any]]:
        text = text.replace("\r", "").replace("\t", " ")
        lines: List[Dict[str, Any]] = []
        offset = 0
        for raw_line in text.split("\n"):
            current = ""
            current_start = offset
            for local_index, char in enumerate(raw_line):
                char_offset = offset + local_index
                if not current or self.text_width(draw, current + char, font) <= max_width:
                    current += char
                    continue
                lines.append({"text": current, "start": current_start})
                current = char
                current_start = char_offset
            lines.append({"text": current, "start": current_start})
            offset += len(raw_line) + 1
        lines = [line for line in lines if line["text"].strip()] or [{"text": "", "start": 0}]
        if max_lines and len(lines) > max_lines:
            lines = lines[:max_lines]
            while lines[-1]["text"] and self.text_width(draw, lines[-1]["text"] + "...", font) > max_width:
                lines[-1]["text"] = lines[-1]["text"][:-1]
            lines[-1]["text"] = lines[-1]["text"] + "..."
        return lines

    def find_preview_cell(self, x: int, y: int) -> Optional[Dict[str, Any]]:
        for cell in reversed(self.preview_cells):
            if cell["x1"] <= x <= cell["x2"] and cell["y1"] <= y <= cell["y2"]:
                return cell
        return None

    @staticmethod
    def target_from_preview_cell(cell: Dict[str, Any]) -> Dict[str, Any]:
        return dict(cell.get("target") or {"table": cell["table"], "row": cell["row"], "col": cell["col"]})

    def on_preview_mouse_down(self, event):
        if not self.debug_workspace_enabled:
            self.status_text.set("请先点击顶部“设置模板”启用模板设置区。")
            return
        if self.preview_widget is not None:
            self.preview_widget.focus_set()
        cell = self.find_preview_cell(event.x, event.y)
        if not cell:
            self.status_text.set("未点中可标注区域，请点击预览图里的表格或下划线单元格。")
            return
        self.mouse_down_cell = cell
        self.drag_adjust_moved = False
        target = self.target_from_preview_cell(cell)
        roles = self.target_roles(target)
        if self.select_mark_for_adjust:
            if self.select_mark_at_target(target):
                return
            self.select_mark_for_adjust = False
            return
        if "mark" in roles and self.selected_mark_target() == target:
            style = self.mark_style_for_target(target)
            self.push_undo_state("拖动微调打勾位置")
            self.drag_adjust_start = (event.x, event.y, int(style.get("offsetX") or 0), int(style.get("offsetY") or 0))
            return
        else:
            self.drag_adjust_start = None
        self.on_cell_clicked(target, cell["text"])

    def on_preview_mouse_drag(self, event):
        if not self.drag_adjust_start:
            return
        start_x, start_y, base_x, base_y = self.drag_adjust_start
        self.drag_adjust_moved = True
        self.mark_offset_x = base_x + event.x - start_x
        self.mark_offset_y = base_y + event.y - start_y
        self.on_mark_style_changed()
        self.status_text.set("正在单独拖动微调当前打勾位置；松开鼠标后刷新预览。")

    def on_preview_mouse_up(self, _event):
        if self.drag_adjust_start:
            cell = self.mouse_down_cell
            moved = self.drag_adjust_moved
            self.drag_adjust_start = None
            self.drag_adjust_moved = False
            self.mouse_down_cell = None
            if moved:
                self.render_table_preview()
                self.status_text.set("已微调当前打勾位置。")
            elif cell:
                self.select_mark_at_target(self.target_from_preview_cell(cell))
        else:
            self.mouse_down_cell = None

    def on_preview_image_clicked(self, event):
        if not self.debug_workspace_enabled:
            self.status_text.set("请先点击顶部“设置模板”启用模板设置区。")
            return
        cell = self.find_preview_cell(event.x, event.y)
        if cell:
            if self.select_mark_for_adjust:
                self.select_mark_at_target(self.target_from_preview_cell(cell))
                return
            self.on_cell_clicked(self.target_from_preview_cell(cell), cell["text"])
            return
        self.status_text.set("未点中可标注区域，请点击预览图里的表格或下划线单元格。")

    def cell_roles(self, table, row, col) -> List[str]:
        return self.target_roles({"table": table, "row": row, "col": col})

    def target_roles(self, target: Dict[str, Any]) -> List[str]:
        roles: List[str] = []
        for config in self.mapping.get("options", {}).values():
            for pair in config_pairs(config):
                if pair.get("judgment") == target and "judgment" not in roles:
                    roles.append("judgment")
                if pair.get("mark") == target and "mark" not in roles:
                    roles.append("mark")
        for field_target in self.mapping.get("fieldTargets", {}).values():
            if field_target == target and "field" not in roles:
                roles.append("field")
        return roles

    def field_for_target(self, target: Dict[str, Any]) -> Optional[str]:
        for field, field_target in self.mapping.get("fieldTargets", {}).items():
            if field_target == target:
                return field
        return None

    def field_preview_value_for_target(self, target: Dict[str, Any]) -> str:
        field = self.field_for_target(target)
        if not field:
            return ""
        if self.records:
            record = self.records[0]
            if field in {"building", "roomNo"}:
                parts = split_room_value(record.room)
                if parts:
                    return str(parts[0] if field == "building" else parts[1])
            return str(getattr(record, field, "") or FIELD_LABELS.get(field, field))
        return FIELD_LABELS.get(field, field)

    def mapping_label_for_key(self, key: str) -> str:
        result_name = self.result_name_for_key(key)
        if result_name:
            return result_name
        return self.option_labels.get(key, key)

    def target_role_labels(self, target: Dict[str, Any], role: str) -> str:
        labels: List[str] = []
        for key, config in self.mapping.get("options", {}).items():
            for pair in config_pairs(config):
                if pair.get(role) != target:
                    continue
                label = pair.get("groupLabel") or self.mapping_label_for_key(key)
                if label not in labels:
                    labels.append(label)
        if not labels:
            return ""
        if len(labels) <= 2:
            return "、".join(labels)
        return f"{labels[0]}等{len(labels)}项"

    def mark_pair_refs_for_target(self, target: Dict[str, int]) -> List[Tuple[str, int, Dict[str, Any], Dict[str, Any]]]:
        refs: List[Tuple[str, int, Dict[str, Any], Dict[str, Any]]] = []
        for key, config in self.mapping.get("options", {}).items():
            for index, pair in enumerate(config_pairs(config)):
                if pair.get("mark") == target:
                    refs.append((key, index, pair, config))
        return refs

    def mark_pair_for_ref(self, ref: Optional[Tuple[str, int]]) -> Optional[Tuple[str, int, Dict[str, Any], Dict[str, Any]]]:
        if not ref:
            return None
        key, pair_index = ref
        config = self.mapping.get("options", {}).get(key)
        if not config:
            return None
        pairs = config_pairs(config)
        if pair_index < 0 or pair_index >= len(pairs):
            return None
        return key, pair_index, pairs[pair_index], config

    def selected_mark_target(self) -> Optional[Dict[str, int]]:
        ref = self.mark_pair_for_ref(self.selected_mark_ref)
        if not ref:
            return None
        return ref[2].get("mark")

    def mark_style_for_target(self, target: Dict[str, int]) -> Dict[str, Any]:
        selected = self.mark_pair_for_ref(self.selected_mark_ref)
        if selected and selected[2].get("mark") == target:
            _key, _index, pair, config = selected
            return dict(pair.get("markStyle") or config.get("markStyle") or self.mapping.get("markStyle") or {})
        refs = self.mark_pair_refs_for_target(target)
        if refs:
            _key, _index, pair, config = refs[0]
            return dict(pair.get("markStyle") or config.get("markStyle") or self.mapping.get("markStyle") or {})
        return dict(self.mapping.get("markStyle") or {})

    def start_select_mark_for_adjust(self):
        if not self.debug_workspace_enabled:
            self.status_text.set("请先点击顶部“设置模板”启用模板设置区。")
            return
        self.select_mark_for_adjust = True
        self.status_text.set("请选择一个蓝色打勾位置；选中后再用方向键或按钮微调。")

    def select_mark_at_target(self, target: Dict[str, int]) -> bool:
        refs = self.mark_pair_refs_for_target(target)
        if not refs:
            messagebox.showinfo("未选中打勾位置", "请点击已经标成蓝色的打勾区域。")
            return False
        selected_keys = [key for key in self.selected_option_keys() if not key.startswith("field:")]
        ref = next((item for item in refs if item[0] in selected_keys), refs[0])
        key, pair_index, pair, config = ref
        self.selected_mark_ref = (key, pair_index)
        self.select_mark_for_adjust = False
        self.set_mark_style_controls(pair.get("markStyle") or config.get("markStyle") or self.mapping.get("markStyle"))
        self.status_text.set(f"已选择 {self.option_labels.get(key, key)} 第 {pair_index + 1} 个打勾位置；现在可单独微调。")
        self.render_table_preview()
        return True

    def is_cell_mapped(self, table, row, col) -> bool:
        return bool(self.cell_roles(table, row, col))

    def on_option_selected(self, _event=None):
        selection = self.option_list.curselection()
        if not selection:
            if self.selected_key and self.selected_key.startswith("field:"):
                return
            self.selected_key = None
            self.last_result_selection = []
            self.update_result_group_text()
            return
        previous_result_keys = list(self.last_result_selection)
        if self.field_list is not None:
            self.field_list.selection_clear(0, "end")
        self.selected_key = self.option_keys[selection[0]]
        selected_keys = self.selected_option_keys()
        result_keys = [key for key in selected_keys if not key.startswith("field:")]
        if self.selected_mark_ref and self.selected_mark_ref[0] not in selected_keys:
            self.selected_mark_ref = None
        self.load_selected_mark_style()
        self.load_selected_choice_mode()
        current_mode = self.brush_mode.get()
        if len(result_keys) == 1 and result_keys != previous_result_keys:
            self.set_brush("judgment")
        else:
            self.set_brush(current_mode if current_mode in {"judgment", "mark"} else "judgment")
        self.last_result_selection = result_keys
        self.update_result_group_text()

    def on_field_selected(self, _event=None):
        if self.field_list is None:
            return
        selection = self.field_list.curselection()
        if not selection:
            if self.selected_key and self.selected_key.startswith("field:"):
                self.selected_key = None
            return
        self.option_list.selection_clear(0, "end")
        self.last_result_selection = []
        index = selection[0]
        if index >= len(self.field_keys):
            return
        self.selected_key = self.field_keys[index]
        self.selected_mark_ref = None
        self.select_mark_for_adjust = False
        self.load_selected_mark_style()
        self.set_brush(self.brush_mode.get() if self.brush_mode.get() in {"judgment", "mark"} else "judgment")

    def selected_option_keys(self) -> List[str]:
        if self.field_list is not None:
            field_selection = self.field_list.curselection()
            if field_selection:
                return [self.field_keys[index] for index in field_selection if index < len(self.field_keys)]
        selection = self.option_list.curselection()
        if selection:
            return [self.option_keys[index] for index in selection if index < len(self.option_keys)]
        return [self.selected_key] if self.selected_key else []

    def selected_result_keys(self) -> List[str]:
        return [key for key in self.selected_option_keys() if key and not key.startswith("field:")]

    def group_label_for_keys(self, keys: List[str], create: bool = True) -> Optional[str]:
        keys = sorted(dict.fromkeys(keys), key=self.option_sort_key)
        if len(keys) <= 1:
            return None
        groups = self.mapping.setdefault("groups", [])
        key_tuple = tuple(keys)
        for group in groups:
            if tuple(group.get("keys") or []) == key_tuple:
                return str(group.get("label") or "")
        if not create:
            return None
        label = f"集合{len(groups) + 1}"
        groups.append({"label": label, "keys": list(key_tuple)})
        return label

    def update_result_group_text(self):
        keys = self.selected_result_keys()
        if not keys:
            self.result_group_text.set("多结果标记：未选择")
            return
        if len(keys) == 1:
            self.result_group_text.set(f"当前：{self.option_labels.get(keys[0], keys[0])}")
            return
        label = self.group_label_for_keys(keys, create=True)
        names = "、".join(self.option_labels.get(key, key) for key in keys)
        self.result_group_text.set(f"多结果标记：{label}（{names}）")

    def add_custom_option(self):
        value = self.custom_option.get().strip()
        if not value:
            return
        normalized = self.normalize_option_key(value)
        for index, key in enumerate(self.option_keys):
            if key == normalized:
                self.option_list.selection_clear(0, "end")
                self.option_list.selection_set(index)
                self.option_list.see(index)
                self.selected_key = normalized
                self.load_selected_choice_mode()
                return
        self.add_option_list_item(normalized, self.result_display_label(normalized, len(self.option_keys) + 1))
        self.custom_option.set("")

    def reset_project_list(self, option_keys: List[str]):
        selected = self.selected_key
        selected_keys = [key for key in self.selected_option_keys() if key and not key.startswith("field:")]
        previous_state = str(self.option_list.cget("state"))
        self.option_list.configure(state="normal")
        self.option_list.delete(0, "end")
        self.option_keys = []
        self.option_labels = {}

        keys = list(dict.fromkeys(option_keys))
        for index, key in enumerate(sorted(keys, key=self.option_sort_key), start=1):
            self.add_option_list_item(key, self.result_display_label(key, index))

        restored_any = False
        for key in selected_keys:
            if key in self.option_keys:
                index = self.option_keys.index(key)
                self.option_list.selection_set(index)
                self.option_list.see(index)
                restored_any = True
        if restored_any:
            self.selected_key = selected_keys[0]
            self.load_selected_choice_mode()
        elif selected in self.option_keys:
            index = self.option_keys.index(selected)
            self.option_list.selection_set(index)
            self.option_list.see(index)
        elif selected and selected.startswith("field:"):
            self.selected_key = selected
        elif self.option_keys:
            self.option_list.selection_set(0)
            self.selected_key = self.option_keys[0]
            self.load_selected_choice_mode()
        else:
            self.selected_key = None
        if previous_state == "disabled":
            self.option_list.configure(state="disabled")
        self.last_result_selection = self.selected_result_keys()
        self.update_result_group_text()

    @staticmethod
    def result_display_label(key: str, index: int) -> str:
        if key.startswith("field:"):
            field = key.split(":", 1)[1]
            return f"字段：{FIELD_LABELS.get(field, field)}"
        if normalize_result_name(key):
            return normalize_result_name(key) or key
        if ":" in key:
            result_name, option = key.split(":", 1)
            if normalize_result_name(result_name):
                return f"{result_name}：{option}"
        return f"结果{index}：{key}"

    @staticmethod
    def display_option_key(key: str) -> str:
        if ":" in key:
            result_name, option = key.split(":", 1)
            if normalize_result_name(result_name):
                return f"{result_name}：{option}"
        return key

    @staticmethod
    def normalize_option_key(value: str) -> str:
        text = value.strip().replace("：", ":")
        if ":" in text:
            result_name, option = text.split(":", 1)
            normalized_result = normalize_result_name(result_name)
            normalized_option = normalize_option(option) or option.strip()
            if normalized_result and normalized_option:
                return make_result_option_key(normalized_result, normalized_option)
        return normalize_option(text) or text

    def add_option_list_item(self, key: str, label: str):
        self.option_keys.append(key)
        self.option_labels[key] = label
        self.option_list.insert("end", label)

    def set_brush(self, mode: str):
        self.brush_mode.set(mode)
        mode_label = {"judgment": "判断区", "mark": "标记区", "field": "字段填入区"}[mode]
        selected_keys = self.selected_option_keys()
        if len(selected_keys) == 1 and selected_keys[0].startswith("field:"):
            selected = self.option_labels.get(selected_keys[0], selected_keys[0])
            self.status_text.set(f"当前项目：{selected}；请直接点击模板中的填入位置。")
            return
        if len(selected_keys) > 1:
            group_label = self.group_label_for_keys([key for key in selected_keys if not key.startswith("field:")], create=True)
            selected = f"{group_label}（{len(selected_keys)} 个结果）" if group_label else f"已选择 {len(selected_keys)} 个结果"
        else:
            selected = self.option_labels.get(self.selected_key or "", self.selected_key or "未选择项目")
        self.status_text.set(f"当前格式刷：{mode_label}；当前项目：{selected}")

    def current_mark_style(self) -> Dict[str, Any]:
        horizontal = {"靠左": "left", "居中": "center", "靠右": "right"}.get(self.mark_horizontal.get(), "center")
        vertical = {"靠上": "top", "居中": "middle", "靠下": "bottom"}.get(self.mark_vertical.get(), "middle")
        try:
            font_size = int(self.mark_font_size.get() or 10)
        except Exception:
            font_size = 10
        return {"horizontal": horizontal, "vertical": vertical, "fontSize": font_size, "offsetX": self.mark_offset_x, "offsetY": self.mark_offset_y}

    def set_mark_style_controls(self, style: Optional[Dict[str, Any]]):
        style = style or {}
        horizontal = {"left": "靠左", "center": "居中", "right": "靠右"}.get(style.get("horizontal"), "居中")
        vertical = {"top": "靠上", "middle": "居中", "bottom": "靠下"}.get(style.get("vertical"), "居中")
        self.mark_horizontal.set(horizontal)
        self.mark_vertical.set(vertical)
        self.mark_font_size.set(int(style.get("fontSize") or 10))
        self.mark_offset_x = int(style.get("offsetX") or 0)
        self.mark_offset_y = int(style.get("offsetY") or 0)

    def load_selected_mark_style(self):
        if not self.selected_key or self.selected_key.startswith("field:"):
            self.set_mark_style_controls(self.mapping.get("markStyle"))
            return
        config = self.mapping.get("options", {}).get(self.selected_key, {})
        self.set_mark_style_controls(config.get("markStyle") or self.mapping.get("markStyle"))

    def load_selected_choice_mode(self):
        if not self.selected_key or self.selected_key.startswith("field:"):
            return
        config = self.mapping.get("options", {}).get(self.selected_key, {})
        result_name = self.result_name_for_key(self.selected_key)
        result_mode = self.mapping.get("resultModes", {}).get(result_name or "")
        self.choice_mode.set(config.get("choiceMode") or result_mode or "single")

    def on_choice_mode_changed(self):
        if not self.selected_key or self.selected_key.startswith("field:"):
            return
        self.push_undo_state("修改选项类型")
        config = self.mapping.setdefault("options", {}).setdefault(
            self.selected_key,
            {"label": self.option_labels.get(self.selected_key, self.selected_key)},
        )
        config["choiceMode"] = self.choice_mode.get()
        result_name = self.result_name_for_key(self.selected_key)
        if result_name:
            self.mapping.setdefault("resultModes", {})[result_name] = self.choice_mode.get()
        self.invalidate_preview()
        self.refresh_mapping_tree()

    def apply_mark_style_to_selected(self):
        if not self.selected_key or self.selected_key.startswith("field:"):
            messagebox.showinfo("请选择选项", "请先选择一个投票选项。")
            return
        self.push_undo_state("调整打勾样式")
        config = self.mapping.setdefault("options", {}).setdefault(self.selected_key, {"label": self.selected_key})
        config["markStyle"] = self.current_mark_style()
        self.invalidate_preview()
        self.log(f"{self.selected_key} 打勾位置已调整。")

    def apply_mark_style_to_all(self):
        if not self.debug_workspace_enabled:
            return
        self.push_undo_state("同步打勾样式")
        style = self.current_mark_style()
        self.mapping["markStyle"] = style
        for config in self.mapping.get("options", {}).values():
            config["markStyle"] = dict(style)
            for pair in config.get("pairs", []) or []:
                if pair.get("mark"):
                    pair["markStyle"] = dict(style)
        self.invalidate_preview()
        self.log("已将打勾位置应用到全部选项。")

    def on_mark_style_changed(self):
        if not self.debug_workspace_enabled:
            return
        style = self.current_mark_style()
        selected = self.mark_pair_for_ref(self.selected_mark_ref)
        if selected:
            _key, _pair_index, pair, _config = selected
            pair["markStyle"] = dict(style)
            self.mark_debug_dirty()
            return
        self.mapping["markStyle"] = dict(style)
        self.invalidate_preview()

    def nudge_mark(self, dx: int, dy: int):
        if not self.debug_workspace_enabled:
            self.status_text.set("请先点击顶部“设置模板”启用模板设置区。")
            return
        if not self.mark_pair_for_ref(self.selected_mark_ref):
            self.status_text.set("请先点击“选择打勾位置”，再选择一个蓝色打勾格。")
            return
        self.push_undo_state("微调打勾位置")
        self.mark_offset_x += dx
        self.mark_offset_y += dy
        self.on_mark_style_changed()
        self.render_table_preview()
        self.status_text.set("已微调当前打勾位置。")

    def on_cell_clicked(self, target: Dict[str, Any], text: str):
        if not self.debug_workspace_enabled:
            messagebox.showinfo("模板设置区未启用", "请先点击顶部“设置模板”。")
            return
        selected_keys = self.selected_option_keys()
        if not selected_keys:
            messagebox.showinfo("先选择结果", "请先在左侧选择投票结果。")
            return
        mode = self.brush_mode.get()
        self.push_undo_state("标注区域")

        field_keys = [key for key in selected_keys if key.startswith("field:")]
        result_keys = [key for key in selected_keys if not key.startswith("field:")]
        if field_keys:
            if len(field_keys) != 1 or result_keys:
                messagebox.showinfo("字段只能单选", "房号、姓名、电话等字段请一次只选择一个，再点击模板填入位置。")
                return
            selected_keys = field_keys
            field = selected_keys[0].split(":", 1)[1]
            self.mapping.setdefault("fieldTargets", {})[field] = target
            self.mark_debug_dirty()
            self.log(f"{FIELD_LABELS.get(field, field)} 填入区 = {target_label(target)}")
        else:
            if mode not in {"judgment", "mark"}:
                messagebox.showinfo("模式不匹配", "投票结果请使用“判断区”或“标记区”。")
                return
            group_label = self.group_label_for_keys(result_keys, create=True) if len(result_keys) > 1 else None
            mark_pair_indexes: Dict[str, int] = {}
            if mode == "mark":
                for key in selected_keys:
                    pair_index, reason = self.resolve_mark_pair_index(key, target)
                    if reason:
                        messagebox.showwarning("标记区未配对", reason)
                        return
                    if pair_index is not None:
                        mark_pair_indexes[key] = pair_index
            for key in selected_keys:
                config = self.mapping.setdefault("options", {}).setdefault(
                    key,
                    {"label": self.option_labels.get(key, key)},
                )
                config["choiceMode"] = self.choice_mode.get()
                result_name = self.result_name_for_key(key)
                if result_name:
                    self.mapping.setdefault("resultModes", {})[result_name] = self.choice_mode.get()
                pairs = config.setdefault("pairs", [])
                if mode == "judgment":
                    pair = {"judgment": target, "judgmentText": text.strip(), "markStyle": self.current_mark_style()}
                    if group_label:
                        pair["groupLabel"] = group_label
                    pairs.append(pair)
                    self.pending_pair_index[key] = len(pairs) - 1
                    self.selected_mark_ref = None
                else:
                    pair_index = mark_pair_indexes.get(key)
                    if pair_index is None:
                        pair = {"mark": target, "markStyle": self.current_mark_style()}
                        if group_label:
                            pair["groupLabel"] = group_label
                        pairs.append(pair)
                        pair_index = len(pairs) - 1
                    else:
                        pairs[pair_index]["mark"] = target
                        pairs[pair_index]["markStyle"] = self.current_mark_style()
                        if group_label:
                            pairs[pair_index]["groupLabel"] = group_label
                    self.pending_pair_index.pop(key, None)
                    if len(selected_keys) == 1:
                        self.selected_mark_ref = (key, pair_index)
            self.mark_debug_dirty()
            target_name = "判断区" if mode == "judgment" else "标记区"
            self.log(f"{len(selected_keys)} 个结果 {target_name} = {target_label(target)}")
        self.sync_validation_to_mapping(invalidate=False)
        self.refresh_mapping_tree()
        self.render_table_preview()
        if self.records:
            self.load_data_preview()

    def resolve_mark_pair_index(self, key: str, target: Dict[str, int]) -> Tuple[Optional[int], str]:
        config = self.mapping.get("options", {}).get(key, {})
        pairs = config.get("pairs", []) or []
        label = self.option_labels.get(key, key)
        if not pairs:
            return None, f"{label} 还没有判断区；请先在“判断区”模式点击该结果的判断条件，再手动切到“标记区”。"

        for index, pair in enumerate(pairs):
            if pair.get("judgment") and not pair.get("mark"):
                return index, ""

        return None, f"{label} 当前没有等待配对的判断区；请先切到“判断区”新增判断条件，再按相同顺序切回“标记区”。"

    def clear_selected_mapping(self):
        selected_keys = self.selected_option_keys()
        if not selected_keys:
            return
        self.push_undo_state("清空当前字段/结果")
        for key in selected_keys:
            if key.startswith("field:"):
                field = key.split(":", 1)[1]
                self.mapping.get("fieldTargets", {}).pop(field, None)
            else:
                self.mapping.get("options", {}).pop(key, None)
                self.pending_pair_index.pop(key, None)
                if self.selected_mark_ref and self.selected_mark_ref[0] == key:
                    self.selected_mark_ref = None
        self.mark_debug_dirty()
        self.refresh_mapping_tree()
        self.render_table_preview()
        self.load_data_preview()

    def clear_all_mappings(self):
        if not messagebox.askyesno("确认清空", "确定清空当前所有标注吗？"):
            return
        self.push_undo_state("清空所有标注")
        self.mapping = blank_mapping()
        self.pending_pair_index = {}
        self.selected_mark_ref = None
        self.select_mark_for_adjust = False
        self.mark_offset_x = 0
        self.mark_offset_y = 0
        self.mark_debug_dirty()
        self.sync_validation_to_mapping(invalidate=False)
        self.refresh_mapping_tree()
        self.render_table_preview()
        self.load_data_preview()

    def refresh_mapping_tree(self):
        for item in self.mapping_tree.get_children():
            self.mapping_tree.delete(item)
        for field in ("room", "name", "phone"):
            target = self.mapping.get("fieldTargets", {}).get(field)
            self.mapping_tree.insert("", "end", values=(f"字段：{FIELD_LABELS.get(field, field)}", target_label(target), "字段填入"))
        for option, config in sorted(self.mapping.get("options", {}).items(), key=lambda item: self.option_sort_key(item[0])):
            label = self.option_labels.get(option, config.get("label") or option)
            pairs = config_pairs(config)
            self.mapping_tree.insert("", "end", values=(label, self.targets_summary([pair.get("judgment") for pair in pairs]), self.targets_summary([pair.get("mark") for pair in pairs])))

    @staticmethod
    def targets_summary(targets: List[Optional[Dict[str, int]]]) -> str:
        configured = [target for target in targets if target]
        if not configured:
            return "未设置"
        if len(configured) == 1:
            return target_label(configured[0])
        return f"{len(configured)} 个区域"

    @staticmethod
    def option_sort_key(value):
        import re

        if isinstance(value, str) and value.startswith("field:"):
            order = {"field:building": 0, "field:roomNo": 1, "field:room": 2, "field:name": 3, "field:phone": 4, "field:area": 5}
            return (-1, order.get(value, 99))
        match = re.search(r"(\d+)", value)
        return (0, int(match.group(1))) if match else (1, value)

    @staticmethod
    def result_name_for_key(key: str) -> Optional[str]:
        direct = normalize_result_name(key)
        if direct:
            return direct
        if ":" not in key:
            return None
        result_name = key.split(":", 1)[0]
        return normalize_result_name(result_name)

    def sync_validation_to_mapping(self, invalidate: bool = True):
        self.mapping["resultSlotCount"] = 0
        self.mapping["filenamePrefix"] = self.filename_prefix.get().strip()
        self.mapping["exportMode"] = self.export_mode.get()
        self.mapping["cleanMode"] = bool(self.clean_mode.get())
        self.mapping.setdefault("validation", {})
        self.mapping["validation"].update(
            {
                "mode": self.validation_mode.get(),
                "min": self.safe_int_var(self.min_count),
                "max": self.safe_int_var(self.max_count),
                "exact": self.safe_int_var(self.exact_count),
                "skipInvalid": True,
            }
        )
        if invalidate:
            self.invalidate_preview()

    @staticmethod
    def safe_int_var(variable) -> int:
        try:
            return int(variable.get() or 0)
        except Exception:
            return 0

    def show_internal_preview(self, preview_path: Path, warnings: List[str]) -> None:
        try:
            initial_pdf = docx_to_pdf(preview_path)
            initial_images, initial_page_info = render_pdf_pages(initial_pdf, zoom=1.25)
        except Exception as exc:
            messagebox.showerror("真实打印预览失败", str(exc))
            return

        dialog = tk.Toplevel(self)
        dialog.title("真实打印预览与用户信息调整")
        dialog.geometry("1220x820")
        dialog.minsize(980, 650)
        dialog.transient(self)
        dialog.grab_set()

        current_docx = Path(preview_path)
        initial_docx = current_docx.resolve()
        current_pdf = Path(initial_pdf)
        rendered_images = initial_images
        page_info = initial_page_info
        render_zoom = 1.25
        page_origins: List[Tuple[int, int]] = []
        overlay_items: List[Dict[str, Any]] = []
        drag_state: Dict[str, Any] = {}
        pdf_field_styles = copy.deepcopy(self.mapping.get("fieldStyles", {}) or {})
        match_cache: Dict[Tuple[str, str], List[Dict[str, float]]] = {}
        refresh_state: Dict[str, Any] = {
            "afterId": None,
            "running": False,
            "requested": 0,
            "closed": False,
            "changeGroupOpen": False,
        }
        controls_loading = False
        confirm_button: Optional[ttk.Button] = None

        def remove_intermediate_docx(path: Optional[Path]) -> None:
            if path is None:
                return
            candidate = Path(path)
            try:
                if candidate.resolve() != initial_docx and candidate.parent.name == "预览":
                    candidate.unlink(missing_ok=True)
            except Exception:
                pass

        selected_field = tk.StringVar(value="room")
        font_name = tk.StringVar(value="宋体")
        font_size = tk.StringVar(value="10")
        font_bold = tk.BooleanVar(value=False)
        offset_x = tk.DoubleVar(value=0)
        offset_y = tk.DoubleVar(value=0)
        zoom_percent = tk.StringVar(value="100")
        paper_text = tk.StringVar(value=page_info[0]["label"] if page_info else "未识别纸张")
        preview_status = tk.StringVar(value="由 Word 直接导出 PDF，纸张、分页和打印版一致。")
        preview_record, _preview_reasons = select_preview_record(self.records, self.mapping)

        header = ttk.Frame(dialog, padding=(12, 10))
        header.pack(side="top", fill="x")
        ttk.Label(header, text="真实打印预览", font=("Microsoft YaHei UI", 12, "bold")).pack(side="left")
        ttk.Label(header, textvariable=paper_text, foreground="#2563eb").pack(side="left", padx=(16, 0))
        if warnings:
            ttk.Label(header, text=f"有 {len(warnings)} 条提示，详情见处理日志。", foreground="#b45309").pack(side="left", padx=(12, 0))

        body = ttk.PanedWindow(dialog, orient="horizontal")
        body.pack(side="top", fill="both", expand=True, padx=12)
        controls = ttk.Frame(body, padding=(0, 0, 12, 0), width=265)
        canvas_box = ttk.Frame(body)
        body.add(controls, weight=0)
        body.add(canvas_box, weight=1)

        field_choices = [
            ("building", "楼栋（1-101 中的 1）"),
            ("roomNo", "房号（1-101 中的 101）"),
            ("room", "完整地址/原始值"),
            ("name", "姓名"),
            ("phone", "电话号码"),
        ]
        field_box = ttk.LabelFrame(controls, text="用户信息字段", padding=8)
        field_box.pack(fill="x")
        field_list = tk.Listbox(field_box, height=5, exportselection=False, font=("Microsoft YaHei UI", 9))
        field_list.pack(fill="x")
        for _field, label in field_choices:
            field_list.insert("end", label)
        field_list.selection_set(2)

        style_box = ttk.LabelFrame(controls, text="字体与位置", padding=8)
        style_box.pack(fill="x", pady=(8, 0))
        ttk.Label(style_box, text="字体").grid(row=0, column=0, sticky="w", pady=3)
        font_combo = ttk.Combobox(
            style_box,
            textvariable=font_name,
            values=("宋体", "微软雅黑", "黑体", "仿宋", "楷体", "Arial", "Times New Roman"),
            width=16,
        )
        font_combo.grid(row=0, column=1, columnspan=2, sticky="ew", pady=3)
        ttk.Label(style_box, text="字号").grid(row=1, column=0, sticky="w", pady=3)
        font_size_spin = ttk.Spinbox(style_box, from_=5, to=72, increment=0.5, textvariable=font_size, width=8)
        font_size_spin.grid(row=1, column=1, sticky="w", pady=3)
        bold_check = ttk.Checkbutton(style_box, text="粗体", variable=font_bold)
        bold_check.grid(row=1, column=2, sticky="w", pady=3)
        ttk.Label(style_box, text="横向偏移(pt)").grid(row=2, column=0, sticky="w", pady=3)
        offset_x_spin = ttk.Spinbox(style_box, from_=-100, to=100, increment=1, textvariable=offset_x, width=8)
        offset_x_spin.grid(row=2, column=1, columnspan=2, sticky="w", pady=3)
        ttk.Label(style_box, text="纵向偏移(pt)").grid(row=3, column=0, sticky="w", pady=3)
        offset_y_spin = ttk.Spinbox(style_box, from_=-100, to=100, increment=1, textvariable=offset_y, width=8)
        offset_y_spin.grid(row=3, column=1, columnspan=2, sticky="w", pady=3)
        nudge_box = ttk.Frame(style_box)
        nudge_box.grid(row=4, column=0, columnspan=3, pady=(6, 2))
        style_box.columnconfigure(1, weight=1)

        apply_button = ttk.Button(controls, text="应用并刷新打印预览")
        apply_button.pack(fill="x", pady=(8, 0))
        ttk.Label(
            controls,
            text="方向按钮每次移动 1pt；也可直接拖动红框。文字和红框会即时移动，停手后自动生成精确的 Word→PDF 打印预览。",
            wraplength=245,
            foreground="#4b5563",
        ).pack(anchor="w", pady=(8, 0))
        zoom_box = ttk.LabelFrame(controls, text="页面显示", padding=8)
        zoom_box.pack(fill="x", pady=(8, 0))
        ttk.Label(zoom_box, text="缩放").pack(side="left")
        zoom_combo = ttk.Combobox(zoom_box, textvariable=zoom_percent, values=("60", "75", "90", "100", "125", "150", "200"), width=6, state="readonly")
        zoom_combo.pack(side="left", padx=(6, 4))
        ttk.Label(zoom_box, text="%").pack(side="left")
        fit_button = ttk.Button(zoom_box, text="适合宽度")
        fit_button.pack(side="right")

        canvas = tk.Canvas(canvas_box, background="#d1d5db", highlightthickness=0)
        scroll_y = ttk.Scrollbar(canvas_box, orient="vertical", command=canvas.yview)
        scroll_x = ttk.Scrollbar(canvas_box, orient="horizontal", command=canvas.xview)
        canvas.configure(yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)
        canvas.grid(row=0, column=0, sticky="nsew")
        scroll_y.grid(row=0, column=1, sticky="ns")
        scroll_x.grid(row=1, column=0, sticky="ew")
        canvas_box.rowconfigure(0, weight=1)
        canvas_box.columnconfigure(0, weight=1)

        def preview_value(field: str) -> str:
            if preview_record is None:
                return ""
            record = preview_record
            if field in {"building", "roomNo"}:
                parts = split_room_value(record.room)
                if parts:
                    return str(parts[0] if field == "building" else parts[1])
                return ""
            return str(getattr(record, field, "") or "")

        def style_for_field(field: str) -> Dict[str, Any]:
            style = {"fontName": "宋体", "fontSize": 10, "bold": False, "offsetX": 0, "offsetY": 0}
            style.update(self.mapping.get("fieldStyles", {}).get(field, {}) or {})
            return style

        def current_style() -> Dict[str, Any]:
            try:
                size = float(font_size.get() or 10)
            except Exception:
                size = 10.0
            try:
                current_offset_x = float(offset_x.get() or 0)
            except Exception:
                current_offset_x = 0.0
            try:
                current_offset_y = float(offset_y.get() or 0)
            except Exception:
                current_offset_y = 0.0
            return {
                "fontName": font_name.get().strip() or "宋体",
                "fontSize": max(5.0, min(72.0, size)),
                "bold": bool(font_bold.get()),
                "offsetX": current_offset_x,
                "offsetY": current_offset_y,
            }

        def field_pdf_matches(field: str, value: str) -> List[Dict[str, float]]:
            cache_key = (field, value)
            if cache_key in match_cache:
                return match_cache[cache_key]
            matches = search_pdf_text(current_pdf, value)
            if field not in {"building", "roomNo"} or not matches:
                match_cache[cache_key] = matches
                return matches
            counterpart_field = "roomNo" if field == "building" else "building"
            counterpart_value = preview_value(counterpart_field)
            counterpart_matches = search_pdf_text(current_pdf, counterpart_value) if counterpart_value else []
            if not counterpart_matches:
                match_cache[cache_key] = matches
                return matches

            def score(match: Dict[str, float]) -> float:
                best = float("inf")
                match_center_y = (match["y0"] + match["y1"]) / 2
                for other in counterpart_matches:
                    if int(other["page"]) != int(match["page"]):
                        continue
                    other_center_y = (other["y0"] + other["y1"]) / 2
                    direction_penalty = 0.0
                    if field == "building" and match["x0"] > other["x0"]:
                        direction_penalty = 500.0
                    if field == "roomNo" and match["x0"] < other["x0"]:
                        direction_penalty = 500.0
                    distance = abs(match_center_y - other_center_y) * 8 + abs(match["x0"] - other["x0"])
                    best = min(best, distance + direction_penalty)
                return best

            selected_matches = [min(matches, key=score)]
            match_cache[cache_key] = selected_matches
            return selected_matches

        def pdf_style_for_field(field: str) -> Dict[str, Any]:
            style = {"fontName": "宋体", "fontSize": 10, "bold": False, "offsetX": 0, "offsetY": 0}
            style.update(pdf_field_styles.get(field, {}) or {})
            return style

        def styles_match(left: Dict[str, Any], right: Dict[str, Any]) -> bool:
            return (
                str(left.get("fontName") or "宋体") == str(right.get("fontName") or "宋体")
                and abs(float(left.get("fontSize") or 10) - float(right.get("fontSize") or 10)) < 0.01
                and bool(left.get("bold", False)) == bool(right.get("bold", False))
                and abs(float(left.get("offsetX") or 0) - float(right.get("offsetX") or 0)) < 0.01
                and abs(float(left.get("offsetY") or 0) - float(right.get("offsetY") or 0)) < 0.01
            )

        def page_background_color(page_index: int, match: Dict[str, float]) -> str:
            try:
                image = rendered_images[page_index]
                candidates = []
                for point_x, point_y in (
                    (match["x0"] - 2, match["y0"] - 2),
                    (match["x1"] + 2, match["y0"] - 2),
                    (match["x0"] - 2, match["y1"] + 2),
                    (match["x1"] + 2, match["y1"] + 2),
                ):
                    sample_x = max(0, min(image.width - 1, int(point_x * render_zoom)))
                    sample_y = max(0, min(image.height - 1, int(point_y * render_zoom)))
                    candidates.append(image.getpixel((sample_x, sample_y)))
                pixel = max(
                    candidates,
                    key=lambda item: int(item) * 3 if isinstance(item, int) else sum(item[:3]),
                )
                if isinstance(pixel, int):
                    return f"#{pixel:02x}{pixel:02x}{pixel:02x}"
                red, green, blue = pixel[:3]
                return f"#{red:02x}{green:02x}{blue:02x}"
            except Exception:
                return "#ffffff"

        def paint_selected_field(force_live: bool = False):
            nonlocal overlay_items
            canvas.delete("field_overlay")
            overlay_items = []
            field = selected_field.get()
            value = preview_value(field)
            if not value or not current_pdf.exists():
                return
            try:
                displayed_style = current_style()
                exact_style = pdf_style_for_field(field)
                show_live_value = force_live or not styles_match(displayed_style, exact_style)
                delta_x = float(displayed_style.get("offsetX") or 0) - float(exact_style.get("offsetX") or 0)
                delta_y = float(displayed_style.get("offsetY") or 0) - float(exact_style.get("offsetY") or 0)
                for match in field_pdf_matches(field, value):
                    page_index = int(match["page"])
                    if page_index >= len(page_origins):
                        continue
                    origin_x, origin_y = page_origins[page_index]
                    exact_x1 = origin_x + match["x0"] * render_zoom
                    exact_y1 = origin_y + match["y0"] * render_zoom
                    exact_x2 = origin_x + match["x1"] * render_zoom
                    exact_y2 = origin_y + match["y1"] * render_zoom
                    if show_live_value:
                        background = page_background_color(page_index, match)
                        canvas.create_rectangle(
                            exact_x1 - 3,
                            exact_y1 - 3,
                            exact_x2 + 3,
                            exact_y2 + 3,
                            fill=background,
                            outline=background,
                            tags=("field_overlay",),
                        )
                        live_x = exact_x1 + delta_x * render_zoom
                        live_y = exact_y1 + delta_y * render_zoom
                        font_pixels = max(6, int(round(float(displayed_style.get("fontSize") or 10) * render_zoom)))
                        font_weight = "bold" if bool(displayed_style.get("bold", False)) else "normal"
                        text_id = canvas.create_text(
                            live_x,
                            live_y,
                            anchor="nw",
                            text=value,
                            fill="#111827",
                            font=(str(displayed_style.get("fontName") or "宋体"), -font_pixels, font_weight),
                            tags=("field_overlay",),
                        )
                        text_box = canvas.bbox(text_id)
                        if text_box:
                            x1, y1, x2, y2 = text_box
                        else:
                            x1 = exact_x1 + delta_x * render_zoom
                            y1 = exact_y1 + delta_y * render_zoom
                            x2 = exact_x2 + delta_x * render_zoom
                            y2 = exact_y2 + delta_y * render_zoom
                    else:
                        x1, y1, x2, y2 = exact_x1, exact_y1, exact_x2, exact_y2
                    border_id = canvas.create_rectangle(
                        x1 - 3,
                        y1 - 3,
                        x2 + 3,
                        y2 + 3,
                        outline="#dc2626",
                        width=2,
                        tags=("field_overlay",),
                    )
                    overlay_items.append({"id": border_id, "x1": x1 - 3, "y1": y1 - 3, "x2": x2 + 3, "y2": y2 + 3})
            except Exception:
                return

        def paint_canvas():
            nonlocal page_origins
            canvas.delete("all")
            page_origins = []
            photos = []
            y = 24
            max_width = 0
            for page_index, image in enumerate(rendered_images):
                photo = ImageTk.PhotoImage(image)
                photos.append(photo)
                x = 24
                page_origins.append((x, y))
                canvas.create_rectangle(x - 2, y - 2, x + image.width + 2, y + image.height + 2, fill="#ffffff", outline="#9ca3af")
                canvas.create_image(x, y, anchor="nw", image=photo)
                canvas.create_text(x, y - 8, anchor="sw", text=f"第 {page_index + 1} 页", fill="#374151")
                y += image.height + 34
                max_width = max(max_width, image.width)
            dialog.preview_image_tk = photos
            paint_selected_field()
            canvas.configure(scrollregion=(0, 0, max_width + 48, max(200, y)))

        def selected_zoom() -> Tuple[int, float]:
            try:
                percent = max(40, min(250, int(float(zoom_percent.get() or 100))))
            except Exception:
                percent = 100
            return percent, 1.25 * percent / 100.0

        def render_current_pdf():
            nonlocal rendered_images, page_info, render_zoom
            _percent, render_zoom = selected_zoom()
            rendered_images, page_info = render_pdf_pages(current_pdf, zoom=render_zoom)
            paper_text.set(page_info[0]["label"] if page_info else "未识别纸张")
            paint_canvas()

        def set_confirm_enabled(enabled: bool):
            if confirm_button is not None and confirm_button.winfo_exists():
                confirm_button.configure(state="normal" if enabled else "disabled")

        def begin_change_group():
            if not refresh_state["changeGroupOpen"]:
                self.push_undo_state("调整用户信息字体/位置")
                refresh_state["changeGroupOpen"] = True

        def store_current_style():
            self.mapping.setdefault("fieldStyles", {})[selected_field.get()] = current_style()

        def finish_exact_refresh(
            token: int,
            mapping_snapshot: Dict[str, Any],
            generated_docx: Optional[Path],
            generated_pdf: Optional[Path],
            generated_images: Optional[List[Image.Image]],
            generated_page_info: Optional[List[Dict[str, Any]]],
            generated_zoom: float,
            refreshed_warnings: List[str],
            error: Optional[Exception],
        ):
            nonlocal current_docx, current_pdf, rendered_images, page_info, render_zoom, pdf_field_styles
            refresh_state["running"] = False
            if refresh_state["closed"] or not dialog.winfo_exists():
                remove_intermediate_docx(generated_docx)
                if generated_pdf is not None and generated_pdf.name.startswith(".live-preview-"):
                    generated_pdf.unlink(missing_ok=True)
                return
            if token != refresh_state["requested"]:
                remove_intermediate_docx(generated_docx)
                if generated_pdf is not None and generated_pdf.name.startswith(".live-preview-"):
                    generated_pdf.unlink(missing_ok=True)
                dialog.after(0, start_exact_refresh)
                return
            if error is not None or generated_docx is None or generated_pdf is None or generated_images is None or generated_page_info is None:
                remove_intermediate_docx(generated_docx)
                refresh_state["changeGroupOpen"] = False
                preview_status.set("精确打印预览生成失败，请重试。")
                set_confirm_enabled(False)
                messagebox.showerror("打印预览刷新失败", str(error or "未知错误"), parent=dialog)
                return
            previous_docx = current_docx
            previous_pdf = current_pdf
            current_docx = Path(generated_docx)
            current_pdf = Path(generated_pdf)
            rendered_images = generated_images
            page_info = generated_page_info
            render_zoom = generated_zoom
            pdf_field_styles = copy.deepcopy(mapping_snapshot.get("fieldStyles", {}) or {})
            match_cache.clear()
            if previous_pdf != current_pdf and previous_pdf.name.startswith(".live-preview-"):
                previous_pdf.unlink(missing_ok=True)
            if previous_docx != current_docx:
                remove_intermediate_docx(previous_docx)
            _current_percent, current_zoom = selected_zoom()
            if abs(current_zoom - render_zoom) > 0.001:
                render_zoom = current_zoom
                rendered_images, page_info = render_pdf_pages(current_pdf, zoom=render_zoom)
            for item in refreshed_warnings:
                if item not in warnings:
                    warnings.append(item)
            paper_text.set(page_info[0]["label"] if page_info else "未识别纸张")
            paint_canvas()
            refresh_state["changeGroupOpen"] = False
            preview_status.set("精确打印预览已更新，当前画面与 Word 打印结果一致。")
            set_confirm_enabled(True)

        def start_exact_refresh():
            if refresh_state["closed"] or refresh_state["running"]:
                return
            refresh_state["afterId"] = None
            token = int(refresh_state["requested"])
            if token <= 0:
                return
            refresh_state["running"] = True
            mapping_snapshot = copy.deepcopy(self.mapping)
            _percent, worker_zoom = selected_zoom()
            template_path = self.template_path.get()
            data_path = self.data_path.get()
            output_dir = self.output_dir.get()
            preview_status.set("正在后台生成精确的 Word→PDF 打印预览……")

            def worker():
                generated_docx: Optional[Path] = None
                generated_pdf: Optional[Path] = None
                generated_images: Optional[List[Image.Image]] = None
                generated_page_info: Optional[List[Dict[str, Any]]] = None
                refreshed_warnings: List[str] = []
                error: Optional[Exception] = None
                try:
                    generated_docx, refreshed_warnings = generate_preview_docx(
                        template_path, data_path, mapping_snapshot, output_dir
                    )
                    generated_pdf = docx_to_pdf(
                        generated_docx,
                        Path(output_dir) / "预览" / f".live-preview-{token}.pdf",
                    )
                    generated_images, generated_page_info = render_pdf_pages(generated_pdf, zoom=worker_zoom)
                except Exception as exc:
                    error = exc
                self.after(
                    0,
                    lambda: finish_exact_refresh(
                        token,
                        mapping_snapshot,
                        generated_docx,
                        generated_pdf,
                        generated_images,
                        generated_page_info,
                        worker_zoom,
                        refreshed_warnings,
                        error,
                    ),
                )

            threading.Thread(target=worker, daemon=True).start()

        def queue_exact_refresh(delay_ms: int = 450, push_undo: bool = True):
            if refresh_state["closed"] or controls_loading:
                return
            if push_undo:
                begin_change_group()
            store_current_style()
            refresh_state["requested"] += 1
            set_confirm_enabled(False)
            paint_selected_field(force_live=True)
            preview_status.set("位置已实时显示；停手后将自动校准为精确打印效果。")
            after_id = refresh_state.get("afterId")
            if after_id is not None:
                try:
                    dialog.after_cancel(after_id)
                except Exception:
                    pass
            refresh_state["afterId"] = dialog.after(max(0, delay_ms), start_exact_refresh)

        def load_field_controls(*_args):
            nonlocal controls_loading
            controls_loading = True
            try:
                selection = field_list.curselection()
                if selection:
                    selected_field.set(field_choices[int(selection[0])][0])
                style = style_for_field(selected_field.get())
                font_name.set(str(style.get("fontName") or "宋体"))
                font_size.set(str(style.get("fontSize") or 10))
                font_bold.set(bool(style.get("bold", False)))
                offset_x.set(float(style.get("offsetX") or 0))
                offset_y.set(float(style.get("offsetY") or 0))
            finally:
                controls_loading = False
            paint_selected_field()

        def on_style_control_changed(*_args):
            queue_exact_refresh()

        def nudge(dx: int, dy: int):
            try:
                next_x = float(offset_x.get() or 0) + dx
            except Exception:
                next_x = float(dx)
            try:
                next_y = float(offset_y.get() or 0) + dy
            except Exception:
                next_y = float(dy)
            offset_x.set(next_x)
            offset_y.set(next_y)
            queue_exact_refresh(delay_ms=300)

        ttk.Button(nudge_box, text="↑", width=4, command=lambda: nudge(0, -1)).grid(row=0, column=1, padx=2, pady=2)
        ttk.Button(nudge_box, text="←", width=4, command=lambda: nudge(-1, 0)).grid(row=1, column=0, padx=2, pady=2)
        ttk.Button(nudge_box, text="→", width=4, command=lambda: nudge(1, 0)).grid(row=1, column=2, padx=2, pady=2)
        ttk.Button(nudge_box, text="↓", width=4, command=lambda: nudge(0, 1)).grid(row=2, column=1, padx=2, pady=2)
        apply_button.configure(command=lambda: queue_exact_refresh(delay_ms=0))
        field_list.bind("<<ListboxSelect>>", load_field_controls)
        font_combo.bind("<<ComboboxSelected>>", on_style_control_changed)
        font_combo.bind("<KeyRelease>", on_style_control_changed)
        bold_check.configure(command=on_style_control_changed)
        for spinbox in (font_size_spin, offset_x_spin, offset_y_spin):
            spinbox.configure(command=on_style_control_changed)
            spinbox.bind("<KeyRelease>", on_style_control_changed)
            spinbox.bind("<Return>", on_style_control_changed)
            spinbox.bind("<FocusOut>", on_style_control_changed)
        zoom_combo.bind("<<ComboboxSelected>>", lambda _event: render_current_pdf())

        def fit_width():
            if not page_info:
                return
            dialog.update_idletasks()
            available = max(300, canvas.winfo_width() - 70)
            factor = available / max(1.0, float(page_info[0]["widthPoints"]))
            zoom_percent.set(str(max(40, min(250, int(round(factor / 1.25 * 100))))))
            render_current_pdf()

        fit_button.configure(command=fit_width)

        def on_canvas_press(event):
            mouse_x, mouse_y = canvas.canvasx(event.x), canvas.canvasy(event.y)
            drag_state.clear()
            for overlay in reversed(overlay_items):
                if overlay["x1"] <= mouse_x <= overlay["x2"] and overlay["y1"] <= mouse_y <= overlay["y2"]:
                    drag_state.update(
                        {
                            "startX": mouse_x,
                            "startY": mouse_y,
                            "baseX": float(offset_x.get() or 0),
                            "baseY": float(offset_y.get() or 0),
                            "moved": False,
                        }
                    )
                    return

        def on_canvas_drag(event):
            if not drag_state:
                return
            dx = canvas.canvasx(event.x) - drag_state["startX"]
            dy = canvas.canvasy(event.y) - drag_state["startY"]
            if not drag_state["moved"]:
                begin_change_group()
                drag_state["moved"] = True
                set_confirm_enabled(False)
            offset_x.set(round(drag_state["baseX"] + dx / render_zoom, 1))
            offset_y.set(round(drag_state["baseY"] + dy / render_zoom, 1))
            store_current_style()
            paint_selected_field(force_live=True)
            preview_status.set("正在实时移动字段；松开鼠标后自动生成精确打印预览。")

        def on_canvas_release(_event):
            if drag_state:
                moved = bool(drag_state.get("moved"))
                drag_state.clear()
                if moved:
                    queue_exact_refresh(delay_ms=120, push_undo=False)

        canvas.bind("<ButtonPress-1>", on_canvas_press)
        canvas.bind("<B1-Motion>", on_canvas_drag)
        canvas.bind("<ButtonRelease-1>", on_canvas_release)

        def on_mousewheel(event):
            if getattr(event, "num", None) == 4:
                units = -3
            elif getattr(event, "num", None) == 5:
                units = 3
            else:
                units = -3 if event.delta > 0 else 3
            canvas.yview_scroll(units, "units")
            return "break"

        canvas.bind("<MouseWheel>", on_mousewheel)
        canvas.bind("<Button-4>", on_mousewheel)
        canvas.bind("<Button-5>", on_mousewheel)

        footer = ttk.Frame(dialog, padding=(12, 10))
        footer.pack(side="bottom", fill="x")
        ttk.Label(footer, textvariable=preview_status, foreground="#374151").pack(side="left")

        def close_preview_dialog(keep_current_docx: bool = False):
            refresh_state["closed"] = True
            after_id = refresh_state.get("afterId")
            if after_id is not None:
                try:
                    dialog.after_cancel(after_id)
                except Exception:
                    pass
            if current_pdf.name.startswith(".live-preview-"):
                current_pdf.unlink(missing_ok=True)
            if not keep_current_docx:
                remove_intermediate_docx(current_docx)
            dialog.destroy()

        def reject_preview():
            self.preview_ready = False
            self.preview_path = None
            self.debug_completed = False
            self.set_debug_workspace_enabled(True)
            self.mark_workflow_done("debug")
            self.reset_workflow_after("debug_done", "preview", "export")
            self.status_text.set("预览未确认，可以继续调整模板标注。")
            close_preview_dialog()

        def confirm_preview():
            if refresh_state["running"] or refresh_state.get("afterId") is not None:
                messagebox.showinfo("请稍候", "精确打印预览仍在生成，请完成后再确认。", parent=dialog)
                return
            self.preview_ready = True
            self.preview_path = current_docx
            self.mark_workflow_done("preview")
            self.status_text.set("预览已确认，可以点击“开始导出”。")
            self.log(f"真实打印预览已确认：{current_docx}；{paper_text.get()}")
            close_preview_dialog(keep_current_docx=True)

        ttk.Button(footer, text="返回修改", command=reject_preview).pack(side="right", padx=(8, 0))
        confirm_button = ttk.Button(footer, text="确认预览，允许导出", command=confirm_preview)
        confirm_button.pack(side="right")
        dialog.protocol("WM_DELETE_WINDOW", reject_preview)
        load_field_controls()
        dialog.focus_set()

    def on_validation_changed(self):
        self.sync_validation_to_mapping()
        if self.records:
            self.load_data_preview()

    def preview_docx(self, show_message: bool = True) -> Optional[Path]:
        if not self.template_path.get() or not self.data_path.get():
            messagebox.showwarning("缺少文件", "请先选择模板和数据源。")
            return None
        if not self.debug_completed:
            messagebox.showwarning("模板设置未完成", "请先点击“设置模板”完成标注，然后点击顶部“4 确认设置”。")
            return None
        self.sync_validation_to_mapping(invalidate=False)
        if not configured_option_keys(self.mapping):
            messagebox.showwarning("缺少标注", "请先在“模板格式刷”里配置投票选项的标记区。")
            return None
        try:
            preview_path, warnings = generate_preview_docx(
                self.template_path.get(),
                self.data_path.get(),
                self.mapping,
                self.output_dir.get(),
            )
            self.log(f"已生成预览：{preview_path}")
            self.preview_ready = False
            self.preview_path = preview_path
            self.reset_workflow_after("preview", "export")
            for item in warnings:
                self.log(f"  - {item}")
            self.show_internal_preview(preview_path, warnings)
            return preview_path
        except Exception as exc:
            messagebox.showerror("预览失败", str(exc))
            return None

    def export_all(self):
        if not self.template_path.get() or not self.data_path.get():
            messagebox.showwarning("缺少文件", "请先选择模板和数据源。")
            return
        if not self.debug_completed:
            messagebox.showwarning("模板设置未完成", "请先点击“设置模板”完成标注，然后点击顶部“4 确认设置”。")
            return
        self.sync_validation_to_mapping(invalidate=False)
        if not configured_option_keys(self.mapping):
            messagebox.showwarning("缺少标注", "请先在“模板格式刷”里配置投票选项的标记区。")
            return

        thread = threading.Thread(target=self._export_worker, daemon=True)
        thread.start()

    def invalidate_preview(self):
        self.preview_ready = False
        self.preview_path = None
        if self.workflow_buttons:
            self.reset_workflow_after("preview", "export")

    def _export_worker(self):
        try:
            export_mode = self.mapping.get("exportMode") or "multi"
            mode_label = "单文件" if export_mode == "single" else "多文件"
            self.log(f"开始批量导出 DOCX... 模式：{mode_label}")
            outputs, warnings, exception_path, failed_count, summary_path, _summary_failed_count, run_output_dir = generate_all(
                self.template_path.get(),
                self.data_path.get(),
                self.mapping,
                self.output_dir.get(),
            )
            self.log(f"生成文件：{len(outputs)} 个 DOCX。输出目录：{run_output_dir}")
            for output in outputs:
                self.log(f"  - {output.name}")
            if failed_count:
                self.log(f"未导出：{failed_count} 行，已写入投票结果汇总.xlsx：{summary_path}")
            if summary_path:
                self.log(f"投票结果汇总：{summary_path}")
            if warnings:
                self.log("警告：")
                for item in warnings:
                    self.log(f"  - {item}")
            def done_message():
                self.mark_workflow_done("export")
                messagebox.showinfo(
                    "导出完成",
                    f"导出模式：{mode_label}\n生成 DOCX：{len(outputs)} 个。\n未导出 {failed_count} 行。\n输出目录：{run_output_dir}\n汇总文件：{summary_path or '无'}",
                )

            self.after(0, done_message)
        except Exception as exc:
            error_message = str(exc)
            self.log(f"导出失败：{error_message}")
            self.after(0, lambda message=error_message: messagebox.showerror("导出失败", message))

    def check_for_updates(self):
        if self.update_button is not None:
            self.update_button.configure(state="disabled", text="检查中...")
        self.log(f"正在从 GitHub 检查更新，当前版本：v{APP_VERSION}")
        threading.Thread(target=self._check_for_updates_worker, daemon=True).start()

    def _check_for_updates_worker(self):
        release: Optional[ReleaseInfo] = None
        error_message = ""
        try:
            release = fetch_latest_release()
        except Exception as exc:
            error_message = str(exc)
        self.after(0, lambda: self._finish_update_check(release, error_message))

    def _finish_update_check(self, release: Optional[ReleaseInfo], error_message: str):
        if self.update_button is not None:
            self.update_button.configure(state="normal", text="检查更新")

        if error_message:
            self.log(f"GitHub 更新检查失败：{error_message}")
            if messagebox.askyesno("检查更新失败", f"无法读取 GitHub 最新版本：\n{error_message}\n\n是否打开项目源码仓库？"):
                webbrowser.open(REPOSITORY_URL)
            return

        if release is None:
            self.log("GitHub 仓库还没有 Release，已提供源码仓库入口。")
            if messagebox.askyesno("暂无发布版本", "GitHub 仓库尚未发布 Release。\n\n是否打开项目源码仓库？"):
                webbrowser.open(REPOSITORY_URL)
            return

        latest_version = release.version
        if is_newer_version(latest_version, APP_VERSION):
            notes = release.body.strip() or "本次发布未填写更新说明。"
            if len(notes) > 600:
                notes = notes[:600].rstrip() + "..."
            self.log(f"发现新版本：v{latest_version}（当前 v{APP_VERSION}）")
            if messagebox.askyesno(
                "发现新版本",
                f"当前版本：v{APP_VERSION}\n最新版本：v{latest_version}\n\n{notes}\n\n是否打开 GitHub 下载页？",
            ):
                webbrowser.open(release.html_url)
            return

        self.log(f"当前已是最新版本：v{APP_VERSION}")
        if messagebox.askyesno("已是最新版本", f"当前版本 v{APP_VERSION} 已是最新版本。\n\n是否打开 GitHub 源码仓库？"):
            webbrowser.open(REPOSITORY_URL)

    def log(self, message: str):
        def append():
            self.log_text.insert("end", message + "\n")
            self.log_text.see("end")

        if threading.current_thread() is threading.main_thread():
            append()
        else:
            self.after(0, append)


if __name__ == "__main__":
    app = VoteDocxApp()
    app.mainloop()
