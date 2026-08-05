from __future__ import annotations

import csv
import re
import shutil
import subprocess
import zipfile
import xml.etree.ElementTree as ET
from copy import deepcopy
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

import openpyxl
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips


APP_NAME = "群众选票格式化打印工具"
CHECK_MARK = "√"
BLACK = RGBColor(0, 0, 0)
WHITE = RGBColor(255, 255, 255)

FIELD_ALIASES = {
    "room": ["房号", "楼栋房号", "楼栋", "地址", "小区楼栋", "房屋"],
    "name": ["姓名", "业主姓名", "选民姓名", "投票人"],
    "phone": ["电话号码", "电话", "手机号", "手机号码", "联系电话"],
    "options": ["投票选项", "选项", "表决意见", "投票结果", "短信结果"],
}

FIELD_LABELS = {
    "building": "楼栋",
    "roomNo": "房号",
    "room": "房号/地址",
    "name": "姓名",
    "phone": "电话号码",
    "area": "房屋面积",
}

OPTION_DIGITS = {
    "①": 1,
    "②": 2,
    "③": 3,
    "④": 4,
    "⑤": 5,
    "⑥": 6,
    "⑦": 7,
    "⑧": 8,
    "⑨": 9,
    "⑩": 10,
    "⑪": 11,
    "⑫": 12,
    "⑬": 13,
    "⑭": 14,
    "⑮": 15,
    "⑯": 16,
    "⑰": 17,
    "⑱": 18,
    "⑲": 19,
    "⑳": 20,
}
OPTION_DIGITS_REVERSE = {number: char for char, number in OPTION_DIGITS.items()}

CN_DIGITS = {
    "零": 0,
    "一": 1,
    "二": 2,
    "两": 2,
    "三": 3,
    "四": 4,
    "五": 5,
    "六": 6,
    "七": 7,
    "八": 8,
    "九": 9,
    "十": 10,
}


@dataclass
class VoteRecord:
    row_no: int
    room: str
    name: str
    phone: str
    options: List[str]
    raw: Dict[str, Any]
    result_options: Dict[str, List[str]] = field(default_factory=dict)


def normalize_text(value: Any) -> str:
    return re.sub(r"\s+", "", "" if value is None else str(value)).strip()


def cn_number_to_int(text: str) -> Optional[int]:
    text = normalize_text(text)
    if not text:
        return None
    if text in CN_DIGITS:
        return CN_DIGITS[text]
    if "十" in text:
        left, _, right = text.partition("十")
        tens = CN_DIGITS.get(left, 1) if left else 1
        ones = CN_DIGITS.get(right, 0) if right else 0
        return tens * 10 + ones
    return None


def normalize_option(value: Any) -> Optional[str]:
    text = normalize_text(value)
    if not text:
        return None

    result_match = re.search(r"结果0*([0-9]{1,3})", text, re.IGNORECASE)
    if result_match:
        return f"结果{int(result_match.group(1))}"

    if text.startswith("结果"):
        number = cn_number_to_int(text.replace("结果", "", 1))
        if number is not None:
            return f"结果{number}"

    for char, number in OPTION_DIGITS.items():
        if char in text:
            return f"选项{number}"

    match = re.search(r"(?:选项|選項|项|項|option)?0*([0-9]{1,3})", text, re.IGNORECASE)
    if match:
        return f"选项{int(match.group(1))}"

    stripped = re.sub(r"^(选项|選項|项|項|第)", "", text)
    stripped = re.sub(r"(项|項|号)$", "", stripped)
    number = cn_number_to_int(stripped)
    if number is not None:
        return f"选项{number}"

    return text


def normalize_result_name(value: Any) -> Optional[str]:
    text = normalize_text(value)
    if not text:
        return None
    match = re.search(r"结果0*([0-9]{1,3})", text, re.IGNORECASE)
    if match:
        return f"结果{int(match.group(1))}"
    if text.startswith("结果"):
        number = cn_number_to_int(text.replace("结果", "", 1))
        if number is not None:
            return f"结果{number}"
    return None


def make_result_option_key(result_name: str, option: str) -> str:
    return f"{result_name}:{option}"


def split_result_option_key(key: str) -> Tuple[Optional[str], str]:
    if ":" in key:
        result_name, option = key.split(":", 1)
        if normalize_result_name(result_name):
            return result_name, option
    return None, key


def display_option_key(key: str) -> str:
    result_name, option = split_result_option_key(key)
    if result_name:
        return f"{result_name}：{option}"
    return key


def option_matches_judgment(selected_option: str, judgment_text: str) -> bool:
    selected = normalize_option(selected_option) or normalize_text(selected_option)
    judgment = normalize_option(judgment_text) or normalize_text(judgment_text)
    if selected and judgment and selected == judgment:
        return True
    selected_raw = normalize_text(selected_option)
    judgment_raw = normalize_text(judgment_text)
    if not selected_raw or not judgment_raw:
        return False
    return selected_raw in judgment_raw or judgment_raw in selected_raw


def matching_pairs_for_selection(config: Dict[str, Any], selected_option: str) -> List[Dict[str, Any]]:
    pairs = config_pairs(config)
    matches = [pair for pair in pairs if option_matches_judgment(selected_option, pair.get("judgmentText", ""))]
    if not matches:
        matches = [pair for pair in pairs if not normalize_text(pair.get("judgmentText", ""))]
    return matches


def result_option_config(
    options: Dict[str, Any],
    result_name: str,
    selected_option: str,
) -> Tuple[Optional[str], Optional[Dict[str, Any]]]:
    for key in (result_name, make_result_option_key(result_name, selected_option), selected_option):
        config = options.get(key)
        if isinstance(config, dict):
            return key, config
    return None, None


def pairs_for_result_selection(
    config_key: str,
    result_name: str,
    config: Dict[str, Any],
    selected_option: str,
) -> List[Dict[str, Any]]:
    if config_key == result_name:
        return matching_pairs_for_selection(config, selected_option)
    return config_pairs(config)


def selected_mark_pair_refs(record: VoteRecord, mapping: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Return the configured mark pairs that are visible for this record's print preview."""
    refs: List[Dict[str, Any]] = []
    seen = set()
    options = mapping.get("options", {})

    def append_pairs(key: str, label: str, config: Dict[str, Any], selected_pairs: List[Dict[str, Any]]) -> None:
        all_pairs = config_pairs(config)
        for pair in selected_pairs:
            if not pair.get("mark"):
                continue
            pair_index = next(
                (index for index, candidate in enumerate(all_pairs) if candidate is pair or candidate == pair),
                None,
            )
            if pair_index is None:
                continue
            identity = (key, pair_index)
            if identity in seen:
                continue
            seen.add(identity)
            pair_label = label if len(selected_pairs) == 1 else f"{label}（第 {pair_index + 1} 个）"
            refs.append({"key": key, "pairIndex": pair_index, "label": pair_label})

    if record.result_options:
        for result_name, selected_options in record.result_options.items():
            for selected_option in selected_options:
                key, config = result_option_config(options, result_name, selected_option)
                if key is None or config is None:
                    continue
                append_pairs(
                    key,
                    f"{result_name}：{selected_option}",
                    config,
                    pairs_for_result_selection(key, result_name, config, selected_option),
                )
        return refs

    for option in record.options:
        config = options.get(option)
        if not isinstance(config, dict):
            continue
        append_pairs(option, display_option_key(option), config, config_pairs(config))
    return refs


def split_options(value: Any, normalize: bool = True, dedupe: bool = True) -> List[str]:
    if value is None:
        return []
    text = str(value).strip()
    if not text:
        return []
    parts = re.split(r"[、,，;；\s]+", text)
    result: List[str] = []
    for part in parts:
        option = normalize_option(part) if normalize else str(part).strip()
        if not option:
            continue
        if dedupe and option in result:
            continue
        result.append(option)
    return result


def record_parse_options(mapping: Dict[str, Any]) -> Dict[str, bool]:
    validation = mapping.get("validation", {}) if isinstance(mapping, dict) else {}
    mode = validation.get("mode", "range")
    if mode == "exact":
        return {"normalize_options": False, "dedupe_options": False}
    return {"normalize_options": False, "dedupe_options": False}


def find_header_index(headers: List[str], field: str) -> Optional[int]:
    aliases = FIELD_ALIASES[field]
    normalized_headers = [normalize_text(header) for header in headers]
    for alias in aliases:
        alias_norm = normalize_text(alias)
        for index, header in enumerate(normalized_headers):
            if alias_norm == header:
                return index
    for alias in aliases:
        alias_norm = normalize_text(alias)
        for index, header in enumerate(normalized_headers):
            if alias_norm and alias_norm in header:
                return index
    return None


def rows_to_records(
    rows: List[List[Any]],
    result_count: int = 0,
    normalize_options: bool = True,
    dedupe_options: bool = True,
) -> List[VoteRecord]:
    if not rows:
        return []
    headers = [str(item).strip() if item is not None else "" for item in rows[0]]
    column_map = {
        "room": find_header_index(headers, "room"),
        "name": find_header_index(headers, "name"),
        "phone": find_header_index(headers, "phone"),
        "options": find_header_index(headers, "options"),
    }
    metadata_columns = {index for key, index in column_map.items() if key != "options" and index is not None}
    result_columns = [
        (index, result_name)
        for index, header in enumerate(headers)
        for result_name in [normalize_result_name(header)]
        if result_name and index not in metadata_columns
    ]

    # Figure 3 style fallback: 房号、姓名、电话号码、投票选项
    if column_map["room"] is None and len(headers) >= 1:
        column_map["room"] = 0
    if column_map["name"] is None and len(headers) >= 2:
        column_map["name"] = 1
    if column_map["phone"] is None and len(headers) >= 3:
        column_map["phone"] = 2
    if column_map["options"] is None and len(headers) >= 4:
        column_map["options"] = 3

    required = ["room", "name", "phone"]
    if not result_columns:
        required.append("options")
    missing = [field for field in required if column_map[field] is None]
    if missing:
        raise ValueError("数据源缺少字段：" + "、".join(FIELD_LABELS.get(field, field) for field in missing))

    records: List[VoteRecord] = []
    for row_no, row in enumerate(rows[1:], start=2):
        if not any(str(cell).strip() for cell in row if cell is not None):
            continue
        raw = {headers[i] if i < len(headers) else f"列{i + 1}": row[i] if i < len(row) else "" for i in range(len(headers))}
        if result_columns:
            options: List[str] = []
            result_options: Dict[str, List[str]] = {}
            for col_index, result_name in result_columns:
                values = split_options(
                    row[col_index] if col_index < len(row) else "",
                    normalize=normalize_options,
                    dedupe=dedupe_options,
                )
                result_options[result_name] = values
                for option in values:
                    key = make_result_option_key(result_name, option)
                    if key not in options:
                        options.append(key)
        else:
            values = split_options(
                row[column_map["options"]] if column_map["options"] < len(row) else "",
                normalize=normalize_options,
                dedupe=dedupe_options,
            )
            options = []
            result_options = {}
            if result_count > 0 or values:
                for index, option in enumerate(values, start=1):
                    result_name = f"结果{index}"
                    result_options.setdefault(result_name, []).append(option)
                    options.append(make_result_option_key(result_name, option))
            else:
                result_options = {}
                options = values
        record = VoteRecord(
            row_no=row_no,
            room=str(row[column_map["room"]] if column_map["room"] < len(row) and row[column_map["room"]] is not None else "").strip(),
            name=str(row[column_map["name"]] if column_map["name"] < len(row) and row[column_map["name"]] is not None else "").strip(),
            phone=str(row[column_map["phone"]] if column_map["phone"] < len(row) and row[column_map["phone"]] is not None else "").strip(),
            options=options,
            raw=raw,
            result_options=result_options,
        )
        records.append(record)
    return records


def read_vote_records(
    path: str | Path,
    result_count: int = 0,
    normalize_options: bool = True,
    dedupe_options: bool = True,
) -> List[VoteRecord]:
    source = Path(path)
    suffix = source.suffix.lower()
    if suffix in {".xlsx", ".xlsm"}:
        workbook = openpyxl.load_workbook(source, data_only=True)
        sheet = workbook.active
        rows = [[cell for cell in row] for row in sheet.iter_rows(values_only=True)]
        return rows_to_records(
            rows,
            result_count=result_count,
            normalize_options=normalize_options,
            dedupe_options=dedupe_options,
        )

    if suffix in {".csv", ".txt"}:
        last_error: Optional[Exception] = None
        for encoding in ("utf-8-sig", "gbk", "gb18030"):
            try:
                with source.open("r", encoding=encoding, newline="") as handle:
                    reader = csv.reader(handle)
                    return rows_to_records(
                        [list(row) for row in reader],
                        result_count=result_count,
                        normalize_options=normalize_options,
                        dedupe_options=dedupe_options,
                    )
            except UnicodeDecodeError as exc:
                last_error = exc
        if last_error:
            raise last_error

    raise ValueError(f"暂不支持的数据源格式：{source.suffix}。请使用 .xlsx、.csv 或 .txt。")


def read_vote_records_for_mapping(path: str | Path, mapping: Dict[str, Any]) -> List[VoteRecord]:
    return read_vote_records(
        path,
        result_count=int(mapping.get("resultSlotCount") or 0),
        **record_parse_options(mapping),
    )


def docx_page_count(path: str | Path) -> Optional[int]:
    try:
        with zipfile.ZipFile(path) as archive:
            data = archive.read("docProps/app.xml")
        root = ET.fromstring(data)
        for element in root.iter():
            if element.tag.rsplit("}", 1)[-1] == "Pages":
                text = (element.text or "").strip()
                return int(text) if text.isdigit() else None
    except Exception:
        return None
    return None


def set_generated_run_visible(run) -> None:
    try:
        run.font.color.rgb = BLACK
    except Exception:
        pass


def set_template_run_hidden(run) -> None:
    try:
        run.font.color.rgb = WHITE
    except Exception:
        pass


def copy_run_color(source_run, target_run) -> None:
    try:
        rgb = source_run.font.color.rgb
        if rgb is not None:
            target_run.font.color.rgb = rgb
    except Exception:
        pass


def set_east_asia_font(run, font_name: str = "宋体", size_pt: Optional[float] = None, bold: Optional[bool] = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    set_generated_run_visible(run)


def set_cell_text(cell, text: str, size_pt: float = 12, bold: bool = False, align: int = WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    set_east_asia_font(run, size_pt=size_pt, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def default_field_style() -> Dict[str, Any]:
    return {"fontName": "宋体", "fontSize": 10, "bold": False, "offsetX": 0, "offsetY": 0}


def normalized_field_style(style: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    result = default_field_style()
    result.update(style or {})
    try:
        result["fontSize"] = max(5.0, min(72.0, float(result.get("fontSize") or 10)))
    except Exception:
        result["fontSize"] = 10.0
    for key in ("offsetX", "offsetY"):
        try:
            result[key] = max(-100.0, min(100.0, float(result.get(key) or 0)))
        except Exception:
            result[key] = 0.0
    result["fontName"] = str(result.get("fontName") or "宋体").strip() or "宋体"
    result["bold"] = bool(result.get("bold", False))
    return result


def set_run_vertical_offset(run, offset_y: float) -> None:
    run_properties = run._r.get_or_add_rPr()
    position = run_properties.find(qn("w:position"))
    if abs(float(offset_y or 0)) < 0.01:
        if position is not None:
            run_properties.remove(position)
        return
    if position is None:
        position = OxmlElement("w:position")
        run_properties.append(position)
    # Word stores character position in half-points. Positive UI Y means down.
    position.set(qn("w:val"), str(int(round(-float(offset_y) * 2))))


def apply_field_run_style(run, style: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    normalized = normalized_field_style(style)
    set_east_asia_font(
        run,
        font_name=normalized["fontName"],
        size_pt=normalized["fontSize"],
        bold=normalized["bold"],
    )
    set_run_vertical_offset(run, normalized["offsetY"])
    return normalized


def set_field_cell_text(cell, text: str, style: Optional[Dict[str, Any]] = None) -> None:
    normalized = normalized_field_style(style)
    # The chosen cell is the field's exact insertion area. Only that cell is
    # changed; labels, units, borders, row heights and neighboring cells stay.
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.left_indent = Pt(normalized["offsetX"]) if normalized["offsetX"] else None
    run = paragraph.add_run(str(text or ""))
    apply_field_run_style(run, normalized)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def paragraph_alignment(value: str):
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(value, WD_ALIGN_PARAGRAPH.CENTER)


def cell_vertical_alignment(value: str):
    return {
        "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
        "middle": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
        "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
    }.get(value, WD_CELL_VERTICAL_ALIGNMENT.CENTER)


def set_mark_text(cell, text: str, style: Optional[Dict[str, Any]] = None) -> None:
    style = style or {}
    # Keep the original paragraph *and its empty formatting runs*. Word uses
    # those runs when it calculates the height of an otherwise blank table
    # row. Deleting them and replacing the cell with a new bold glyph changes
    # the row height by fractions of a point, which accumulates down the page.
    paragraph = cell.paragraphs[0]
    for cell_paragraph in cell.paragraphs:
        for existing_run in cell_paragraph.runs:
            existing_run.text = ""
    paragraph.alignment = paragraph_alignment(style.get("horizontal", "center"))
    try:
        offset_x = float(style.get("offsetX") or 0)
        offset_y = float(style.get("offsetY") or 0)
    except Exception:
        offset_x = 0
        offset_y = 0
    point_units = style.get("offsetUnits") == "pt"
    offset_x_points = offset_x if point_units else offset_x * 0.35
    offset_y_points = offset_y if point_units else offset_y * 0.25
    if offset_x_points:
        paragraph.paragraph_format.left_indent = Pt(offset_x_points)
    run = paragraph.add_run(text)
    set_east_asia_font(
        run,
        # SimSun treats the radical glyph as an unusually tall mathematical
        # character and makes Word grow the otherwise blank mark row. Arial
        # renders the same glyph within the template's existing line box, so
        # adding a mark does not alter any table or page geometry.
        font_name=str(style.get("fontName") or "Arial"),
        size_pt=float(style.get("fontSize") or 10),
        bold=bool(style.get("bold", False)),
    )
    # Character positioning moves the glyph without adding paragraph spacing,
    # so vertical nudging does not change the row height or document layout.
    if offset_y_points:
        run_properties = run._r.get_or_add_rPr()
        position = run_properties.find(qn("w:position"))
        if position is None:
            position = OxmlElement("w:position")
            run_properties.append(position)
        position.set(qn("w:val"), str(int(round(-offset_y_points * 2))))
    cell.vertical_alignment = cell_vertical_alignment(style.get("vertical", "middle"))


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in kwargs[edge].items():
            element.set(qn(f"w:{key}"), str(value))


def no_border_attrs() -> Dict[str, str]:
    return {"val": "nil", "sz": "0", "space": "0", "color": "FFFFFF"}


def set_no_border_element(parent, tag_name: str) -> None:
    element = parent.find(qn(tag_name))
    if element is None:
        element = OxmlElement(tag_name)
        parent.append(element)
    for key, value in no_border_attrs().items():
        element.set(qn(f"w:{key}"), value)


def hide_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    # Keep table style/look metadata because it can control cell margins and
    # row geometry. Pure mode should hide borders, not rebuild the layout.
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    style_name = ""
    try:
        style_name = table.style.name or ""
    except Exception:
        pass
    is_grid_table = "grid" in style_name.lower()

    if borders is not None:
        for edge in borders:
            edge.set(qn("w:color"), "FFFFFF")
    elif is_grid_table:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
        for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            edge = OxmlElement(f"w:{edge_name}")
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), "4")
            edge.set(qn("w:space"), "0")
            edge.set(qn("w:color"), "FFFFFF")
            borders.append(edge)

    # Iterate the raw ``w:tc`` elements. ``row.cells`` resolves vertically
    # merged continuation cells to their master cell and therefore skips the
    # continuation XML, leaving a few original borders visible as L-shapes.
    for table_row in table._tbl.tr_lst:
        for table_cell in table_row.tc_lst:
            tc_pr = table_cell.get_or_add_tcPr()
            cell_borders = tc_pr.first_child_found_in("w:tcBorders")
            if cell_borders is not None:
                for edge in cell_borders:
                    edge.set(qn("w:color"), "FFFFFF")
            elif is_grid_table:
                cell_borders = OxmlElement("w:tcBorders")
                tc_pr.append(cell_borders)
                for edge_name in ("top", "left", "bottom", "right"):
                    edge = OxmlElement(f"w:{edge_name}")
                    edge.set(qn("w:val"), "single")
                    edge.set(qn("w:sz"), "4")
                    edge.set(qn("w:space"), "0")
                    edge.set(qn("w:color"), "FFFFFF")
                    cell_borders.append(edge)


def hide_template_content(document) -> None:
    for paragraph in iter_paragraphs(document):
        for run in paragraph.runs:
            set_template_run_hidden(run)
    for table in document.tables:
        hide_table_borders(table)
        for table_row in table._tbl.tr_lst:
            for table_cell in table_row.tc_lst:
                set_cell_shading(table_cell, "FFFFFF")


def remove_generated_underlines(document) -> None:
    for paragraph in iter_paragraphs(document):
        for run in paragraph.runs:
            try:
                if run.font.color.rgb == BLACK:
                    run.font.underline = False
            except Exception:
                continue


def run_is_generated_visible(run) -> bool:
    try:
        return run.font.color.rgb == BLACK
    except Exception:
        return False


def visible_text_from_paragraph(paragraph) -> str:
    return "".join(run.text or "" for run in paragraph.runs if run_is_generated_visible(run))


def visible_text_from_cell(cell) -> str:
    parts = [visible_text_from_paragraph(paragraph) for paragraph in cell.paragraphs]
    return "\n".join(part for part in parts if part).strip()


def cell_width_twips(cell) -> Optional[int]:
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return None
    tc_w = tc_pr.tcW
    if tc_w is None:
        return None
    try:
        return int(tc_w.w)
    except Exception:
        return None


def row_height_points(row) -> float:
    tr_pr = row._tr.trPr
    if tr_pr is not None:
        for child in tr_pr:
            if child.tag == qn("w:trHeight"):
                value = child.get(qn("w:val"))
                if value and str(value).isdigit():
                    return max(12.0, int(value) / 20.0)
    return 22.0


def table_width_twips(document, column_count: int) -> int:
    try:
        section = document.sections[0]
        usable = int(section.page_width.twips - section.left_margin.twips - section.right_margin.twips)
        if usable > 0:
            return usable
    except Exception:
        pass
    return max(1, column_count) * 1800


def table_column_widths_twips(document, table) -> List[int]:
    rows = list(table.rows)
    if not rows:
        return []
    cells = list(rows[0].cells)
    widths = [cell_width_twips(cell) for cell in cells]
    if widths and all(width and width > 0 for width in widths):
        return [int(width) for width in widths if width]
    total = table_width_twips(document, len(cells))
    width = max(1, total // max(1, len(cells)))
    return [width for _cell in cells]


def flatten_clean_tables(document) -> None:
    for table in list(document.tables):
        widths = table_column_widths_twips(document, table)
        column_count = max(1, len(widths))
        for row in table.rows:
            cells = list(row.cells)
            texts = [visible_text_from_cell(cell) for cell in cells]
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = Pt(row_height_points(row))
            stops = paragraph.paragraph_format.tab_stops
            cumulative = 0
            for width in widths[: max(0, column_count - 1)]:
                cumulative += int(width)
                stops.add_tab_stop(Twips(cumulative))
            run = paragraph.add_run("\t".join(texts))
            set_east_asia_font(run, size_pt=12)
            table._tbl.addprevious(paragraph._p)
        parent = table._tbl.getparent()
        if parent is not None:
            parent.remove(table._tbl)


def set_cell_shading(cell, fill: str) -> None:
    table_cell = getattr(cell, "_tc", cell)
    tc_pr = table_cell.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_table_grid(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": "6", "color": "666666"},
                bottom={"val": "single", "sz": "6", "color": "666666"},
                left={"val": "single", "sz": "6", "color": "666666"},
                right={"val": "single", "sz": "6", "color": "666666"},
            )


def iter_paragraphs(document) -> Iterable[Any]:
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def replace_in_paragraph(paragraph, replacements: Dict[str, str]) -> None:
    full_text = "".join(run.text for run in paragraph.runs)
    if not any(key in full_text for key in replacements):
        return
    for key, value in replacements.items():
        full_text = full_text.replace(key, value)
    first_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    for run in paragraph.runs:
        run.text = ""
    first_run.text = full_text
    set_generated_run_visible(first_run)


def replace_placeholders_in_paragraph_clean(paragraph, replacements: Dict[str, str]) -> None:
    for run in list(paragraph.runs):
        current_run = run
        text = current_run.text or ""
        while text:
            found_key = None
            found_pos = -1
            for key in replacements:
                pos = text.find(key)
                if pos >= 0 and (found_pos < 0 or pos < found_pos):
                    found_key = key
                    found_pos = pos
            if found_key is None:
                break
            before = text[:found_pos]
            after = text[found_pos + len(found_key) :]
            current_run.text = before
            value_run = paragraph.add_run(str(replacements.get(found_key) or ""))
            set_east_asia_font(value_run, size_pt=12)
            current_run._r.addnext(value_run._r)
            after_run = paragraph.add_run(after)
            copy_run_color(current_run, after_run)
            value_run._r.addnext(after_run._r)
            current_run = after_run
            text = after


def replace_placeholders(document, values: Dict[str, str], clean_mode: bool = False) -> None:
    replacements = {
        "{{ROOM}}": values.get("room", ""),
        "{{NAME}}": values.get("name", ""),
        "{{PHONE}}": values.get("phone", ""),
        "{{AREA}}": "",
    }
    for paragraph in iter_paragraphs(document):
        if clean_mode:
            replace_placeholders_in_paragraph_clean(paragraph, replacements)
        else:
            replace_in_paragraph(paragraph, replacements)


AUTO_FIELD_ALIASES = {
    "room": ["楼栋房号", "房号/地址", "房号", "地址", "小区楼栋", "楼栋"],
    "name": ["姓名", "业主姓名", "选民姓名", "投票人"],
    "phone": ["电话号码", "电话", "手机号", "手机号码", "联系电话"],
}


def text_has_field_label(text: str) -> bool:
    normalized = normalize_text(text)
    return any(normalize_text(alias) in normalized for aliases in AUTO_FIELD_ALIASES.values() for alias in aliases)


def find_alias_position(text: str, alias: str) -> int:
    exact = text.find(alias)
    if exact >= 0:
        return exact
    compact_chars: List[str] = []
    original_indexes: List[int] = []
    for index, char in enumerate(text):
        if char.isspace():
            continue
        compact_chars.append(char)
        original_indexes.append(index)
    compact = "".join(compact_chars)
    normalized_alias = normalize_text(alias)
    pos = compact.find(normalized_alias)
    if pos < 0 or pos >= len(original_indexes):
        return -1
    return original_indexes[pos]


def infer_field_targets(document) -> Dict[str, Tuple[Dict[str, int], bool, str]]:
    targets: Dict[str, Tuple[Dict[str, int], bool, str]] = {}
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            cells = list(row.cells)
            for col_index, cell in enumerate(cells):
                text = cell.text
                normalized = normalize_text(text)
                if not normalized:
                    continue
                for field, aliases in AUTO_FIELD_ALIASES.items():
                    if field in targets:
                        continue
                    if not any(normalize_text(alias) in normalized for alias in aliases):
                        continue
                    alias_pos = -1
                    for alias in sorted(aliases, key=len, reverse=True):
                        pos = find_alias_position(text, alias)
                        if pos >= 0:
                            alias_pos = pos
                            break
                    target: Optional[Dict[str, Any]] = None
                    for item in cell_underline_targets(cell, table_index, row_index, col_index):
                        if int(item["start"]) >= max(0, alias_pos):
                            target = item["target"]
                            break
                    if target:
                        targets[field] = (target, False, text)
                        continue
                    for next_col in range(col_index + 1, len(cells)):
                        next_cell = cells[next_col]
                        if next_cell._tc is cell._tc:
                            continue
                        next_text = normalize_text(next_cell.text)
                        if text_has_field_label(next_text):
                            break
                        next_underlines = cell_underline_targets(next_cell, table_index, row_index, next_col)
                        target = next_underlines[0]["target"] if next_underlines else {"table": table_index, "row": row_index, "col": next_col}
                        if not next_text or "{{" in next_cell.text or "___" in next_cell.text or "____" in next_cell.text:
                            break
                    if target:
                        targets[field] = (target, False, text)
                    else:
                        targets[field] = ({"table": table_index, "row": row_index, "col": col_index}, True, text)
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text
        matches = list(UNDERLINE_PATTERN.finditer(text))
        run_matches = formatted_underline_spans(paragraph, {"kind": "underline", "paragraph": paragraph_index})
        all_matches = [
            {"start": match.start(), "target": {"kind": "underline", "paragraph": paragraph_index, "underline": index}}
            for index, match in enumerate(matches)
        ] + [{"start": item["start"], "target": item["target"]} for item in run_matches]
        all_matches.sort(key=lambda item: item["start"])
        if not text or not all_matches:
            continue
        normalized = normalize_text(text)
        for field, aliases in AUTO_FIELD_ALIASES.items():
            if field in targets:
                continue
            alias_hit = ""
            alias_pos = -1
            for alias in sorted(aliases, key=len, reverse=True):
                pos = find_alias_position(text, alias)
                if pos >= 0:
                    alias_hit = alias
                    alias_pos = pos
                    break
            if alias_pos < 0:
                continue
            target = all_matches[0]["target"]
            for item in all_matches:
                if int(item["start"]) >= alias_pos:
                    target = item["target"]
                    break
            targets[field] = (
                target,
                False,
                alias_hit or text,
            )
    for field, target_info in infer_address_component_targets(document).items():
        targets.setdefault(field, target_info)
    return targets


def set_labeled_field_text(
    cell,
    field: str,
    value: str,
    label_text: str,
    style: Optional[Dict[str, Any]] = None,
) -> None:
    aliases = sorted(AUTO_FIELD_ALIASES.get(field, []), key=len, reverse=True)
    text = cell.text.strip() or label_text.strip()
    label = next((alias for alias in aliases if normalize_text(alias) in normalize_text(text)), FIELD_LABELS.get(field, field))
    suffix = "：" if "：" in text else ":"
    normalized = normalized_field_style(style)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if field == "room" else WD_ALIGN_PARAGRAPH.CENTER
    label_run = paragraph.add_run(f"{label}{suffix}")
    set_east_asia_font(label_run, size_pt=10)
    value_run = paragraph.add_run(str(value or ""))
    apply_field_run_style(value_run, normalized)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


UNDERLINE_PATTERN = re.compile(r"[_＿]{2,}|[—－-]{4,}")


def target_kind(target: Dict[str, Any]) -> str:
    return str(target.get("kind") or "cell")


def run_has_underline(run) -> bool:
    underline = run.font.underline
    if underline:
        return True
    rpr = run._element.rPr
    underline_element = rpr.u if rpr is not None else None
    if underline_element is None:
        return False
    value = underline_element.val
    return value is None or str(value).lower() not in {"none", "false", "0"}


def is_blank_underline_text(text: str) -> bool:
    cleaned = (text or "").replace("\t", "").replace("\r", "").replace("\n", "")
    cleaned = cleaned.replace(" ", "").replace("\u3000", "").replace("\xa0", "")
    cleaned = cleaned.replace("_", "").replace("＿", "").replace("-", "").replace("－", "").replace("—", "")
    return bool(text) and not cleaned.strip()


def formatted_underline_spans(paragraph, base_target: Dict[str, Any]) -> List[Dict[str, Any]]:
    spans: List[Dict[str, Any]] = []
    offset = 0
    underline_index = 0
    for run_index, run in enumerate(paragraph.runs):
        text = run.text or ""
        if run_has_underline(run) and text:
            target = dict(base_target)
            target["kind"] = "cellRunUnderline" if target.get("kind") == "cellUnderline" else "runUnderline"
            target["run"] = run_index
            target["underline"] = underline_index
            spans.append({"target": target, "start": offset, "end": offset + len(text), "text": text})
            underline_index += 1
        offset += len(text)
    return spans


def paragraph_underline_targets(paragraph, base_target: Dict[str, Any]) -> List[Dict[str, Any]]:
    items: List[Dict[str, Any]] = []
    for index, match in enumerate(UNDERLINE_PATTERN.finditer(paragraph.text)):
        target = dict(base_target)
        target["kind"] = "cellUnderline" if base_target.get("kind") == "cellUnderline" else "underline"
        target["underline"] = index
        items.append({"start": match.start(), "end": match.end(), "target": target})
    for span in formatted_underline_spans(paragraph, base_target):
        items.append({"start": span["start"], "end": span["end"], "target": span["target"]})
    return sorted(items, key=lambda item: int(item["start"]))


def cell_underline_targets(cell, table_index: int, row_index: int, col_index: int) -> List[Dict[str, Any]]:
    items: List[Dict[str, Any]] = []
    paragraph_offset = 0
    for paragraph_index, paragraph in enumerate(cell.paragraphs):
        base_target = {
            "kind": "cellUnderline",
            "table": table_index,
            "row": row_index,
            "col": col_index,
            "paragraph": paragraph_index,
        }
        for item in paragraph_underline_targets(paragraph, base_target):
            copied = dict(item)
            copied["start"] = int(copied["start"]) + paragraph_offset
            copied["end"] = int(copied["end"]) + paragraph_offset
            items.append(copied)
        paragraph_offset += len(paragraph.text) + 1
    return sorted(items, key=lambda item: int(item["start"]))


def address_component_targets_in_paragraph(
    paragraph,
    base_target: Dict[str, Any],
) -> Dict[str, Tuple[Dict[str, int], bool, str]]:
    text = paragraph.text or ""
    underline_items = paragraph_underline_targets(paragraph, base_target)
    if len(underline_items) < 2:
        return {}
    building_positions = [text.find(unit) for unit in ("栋", "幢") if text.find(unit) >= 0]
    if not building_positions:
        return {}
    building_unit = min(building_positions)
    room_unit = text.find("室", building_unit + 1)
    if room_unit < 0:
        return {}
    building_item = next(
        (item for item in reversed(underline_items) if int(item["end"]) <= building_unit + 1),
        None,
    )
    room_item = next(
        (
            item
            for item in underline_items
            if int(item["start"]) >= building_unit and int(item["start"]) < room_unit
        ),
        None,
    )
    if not building_item or not room_item or building_item["target"] == room_item["target"]:
        return {}
    return {
        "building": (dict(building_item["target"]), False, "楼栋"),
        "roomNo": (dict(room_item["target"]), False, "房号"),
    }


def infer_address_component_targets(document) -> Dict[str, Tuple[Dict[str, int], bool, str]]:
    result: Dict[str, Tuple[Dict[str, int], bool, str]] = {}
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        base = {"kind": "underline", "paragraph": paragraph_index}
        result.update({key: value for key, value in address_component_targets_in_paragraph(paragraph, base).items() if key not in result})
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    base = {
                        "kind": "cellUnderline",
                        "table": table_index,
                        "row": row_index,
                        "col": col_index,
                        "paragraph": paragraph_index,
                    }
                    found = address_component_targets_in_paragraph(paragraph, base)
                    result.update({key: value for key, value in found.items() if key not in result})
    return result


def get_body_paragraph(document, index: int):
    return document.paragraphs[int(index)]


def get_target_paragraph(document, target: Dict[str, Any]):
    if target_kind(target) in {"cellUnderline", "cellRunUnderline"}:
        cell = get_cell(document, target)
        paragraph_index = int(target.get("paragraph", 0))
        if not cell.paragraphs:
            return cell.add_paragraph()
        paragraph_index = max(0, min(paragraph_index, len(cell.paragraphs) - 1))
        return cell.paragraphs[paragraph_index]
    return get_body_paragraph(document, int(target.get("paragraph", 0)))


def replace_underline_in_paragraph(paragraph, underline_index: int, value: str) -> None:
    text = paragraph.text
    matches = list(UNDERLINE_PATTERN.finditer(text))
    if not matches:
        paragraph.add_run(value)
        return
    index = max(0, min(int(underline_index), len(matches) - 1))
    match = matches[index]
    new_text = text[:match.start()] + value + text[match.end():]
    first_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    for run in paragraph.runs:
        run.text = ""
    first_run.text = new_text
    set_generated_run_visible(first_run)


def display_width(text: str) -> int:
    width = 0
    for char in str(text):
        width += 1 if ord(char) < 128 else 2
    return width


def center_value_for_underline(
    value: str,
    template_text: str,
    min_slots: int = 0,
    style: Optional[Dict[str, Any]] = None,
) -> str:
    value = "" if value is None else str(value).strip()
    original = template_text or ""
    # Keep the replacement text length aligned to the original template slot.
    # Counting Chinese unit characters as double-width made fields visibly longer.
    slots = max(len(original), len(value), min_slots)
    padding = max(0, slots - len(value))
    normalized = normalized_field_style(style)
    approximate_space_width = max(2.5, float(normalized["fontSize"]) * 0.5)
    shift_slots = int(round(float(normalized["offsetX"]) / approximate_space_width))
    left = max(0, min(padding, padding // 2 + shift_slots))
    right = padding - left
    return f"{' ' * left}{value}{' ' * right}"


def replace_run_underline_in_paragraph(
    paragraph,
    run_index: int,
    value: str,
    centered: bool = False,
    style: Optional[Dict[str, Any]] = None,
) -> None:
    runs = paragraph.runs
    if not runs:
        run = paragraph.add_run(center_value_for_underline(value, "", style=style) if centered else value)
        if centered:
            run.font.underline = True
        if style is not None:
            apply_field_run_style(run, style)
        return
    index = max(0, min(int(run_index), len(runs) - 1))
    run = runs[index]
    original = run.text or ""
    run.text = center_value_for_underline(value, original, style=style) if centered else value
    if centered:
        run.font.underline = True
    if style is not None:
        apply_field_run_style(run, style)
    else:
        set_generated_run_visible(run)


def replace_plain_underline_with_run(
    paragraph,
    underline_index: int,
    value: str,
    centered: bool = False,
    style: Optional[Dict[str, Any]] = None,
) -> None:
    text = paragraph.text
    matches = list(UNDERLINE_PATTERN.finditer(text))
    if not matches:
        run = paragraph.add_run(center_value_for_underline(value, "", style=style) if centered else value)
        if centered:
            run.font.underline = True
        if style is not None:
            apply_field_run_style(run, style)
        return
    index = max(0, min(int(underline_index), len(matches) - 1))
    match = matches[index]
    offset = 0
    for run in list(paragraph.runs):
        run_text = run.text or ""
        run_end = offset + len(run_text)
        if match.start() >= offset and match.end() <= run_end:
            local_start = match.start() - offset
            local_end = match.end() - offset
            before_part = run_text[:local_start]
            matched = run_text[local_start:local_end]
            after_part = run_text[local_end:]
            run.text = before_part
            value_run = paragraph.add_run(center_value_for_underline(value, matched, style=style) if centered else value)
            if style is not None:
                apply_field_run_style(value_run, style)
            else:
                set_generated_run_visible(value_run)
            if centered:
                value_run.font.underline = True
                if style is None:
                    set_east_asia_font(value_run, size_pt=12)
            run._r.addnext(value_run._r)
            if after_part:
                after_run = paragraph.add_run(after_part)
                copy_run_color(run, after_run)
                value_run._r.addnext(after_run._r)
            return
        offset = run_end
    before = text[: match.start()]
    matched = text[match.start() : match.end()]
    after = text[match.end() :]
    first_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    for run in paragraph.runs:
        run.text = ""
    first_run.text = before
    value_run = paragraph.add_run(center_value_for_underline(value, matched, style=style) if centered else value)
    if style is not None:
        apply_field_run_style(value_run, style)
    else:
        set_generated_run_visible(value_run)
    if centered:
        value_run.font.underline = True
        if style is None:
            set_east_asia_font(value_run, size_pt=12)
    after_run = paragraph.add_run(after)
    copy_run_color(first_run, after_run)


def underline_targets_for_same_paragraph(document, target: Dict[str, Any]) -> List[Dict[str, Any]]:
    paragraph = get_target_paragraph(document, target)
    if target_kind(target) in {"cellUnderline", "cellRunUnderline"}:
        base = {
            "kind": "cellUnderline",
            "table": int(target["table"]),
            "row": int(target["row"]),
            "col": int(target["col"]),
            "paragraph": int(target.get("paragraph", 0)),
        }
    else:
        base = {"kind": "underline", "paragraph": int(target.get("paragraph", 0))}
    return [item["target"] for item in paragraph_underline_targets(paragraph, base)]


def replace_field_target(
    document,
    target: Dict[str, Any],
    value: str,
    style: Optional[Dict[str, Any]] = None,
) -> None:
    if target_kind(target) in {"runUnderline", "cellRunUnderline"}:
        paragraph = get_target_paragraph(document, target)
        replace_run_underline_in_paragraph(paragraph, int(target.get("run", 0)), value, centered=True, style=style)
        return
    if target_kind(target) in {"underline", "cellUnderline"}:
        paragraph = get_target_paragraph(document, target)
        replace_plain_underline_with_run(paragraph, int(target.get("underline", 0)), value, centered=True, style=style)
        return
    cell = get_cell(document, target)
    set_field_cell_text(cell, value, style)


def find_first_unit(text: str, units: str) -> int:
    indexes = [text.find(unit) for unit in units if text.find(unit) >= 0]
    return min(indexes) if indexes else -1


def find_next_non_room_field_position(text: str, start: int = 0) -> int:
    return find_next_field_position(text, start, {"room"})


def find_next_field_position(text: str, start: int = 0, exclude_fields: Optional[set[str]] = None) -> int:
    exclude_fields = exclude_fields or set()
    best = -1
    for field, aliases in AUTO_FIELD_ALIASES.items():
        if field in exclude_fields:
            continue
        for alias in aliases:
            pos = find_alias_position(text[start:], alias)
            if pos < 0:
                continue
            absolute = start + pos
            if best < 0 or absolute < best:
                best = absolute
    return best


def replace_text_range_in_runs(paragraph, start: int, end: int, replacement: str = "") -> None:
    if end <= start:
        return
    runs = list(paragraph.runs)
    offset = 0
    replacement_written = False
    for run in runs:
        run_text = run.text or ""
        run_start = offset
        run_end = offset + len(run_text)
        offset = run_end
        if run_end <= start or run_start >= end:
            continue
        local_start = max(0, start - run_start)
        local_end = min(len(run_text), end - run_start)
        prefix = run_text[:local_start]
        suffix = run_text[local_end:]
        if not replacement_written:
            run.text = prefix + replacement + suffix
            if replacement:
                set_generated_run_visible(run)
            replacement_written = True
        else:
            run.text = suffix


def replace_text_range_with_underlined_run(paragraph, start: int, end: int, value: str) -> bool:
    if end <= start:
        return False
    text = paragraph.text
    template_text = text[start:end]
    runs = list(paragraph.runs)
    if not runs:
        value_run = paragraph.add_run(center_value_for_underline(value, template_text))
        value_run.font.underline = True
        set_east_asia_font(value_run, size_pt=12)
        return True

    start_run = None
    start_local = 0
    end_run = None
    end_local = 0
    offset = 0
    for run in runs:
        run_text = run.text or ""
        run_end = offset + len(run_text)
        if start_run is None and start >= offset and start <= run_end:
            start_run = run
            start_local = start - offset
        if end_run is None and end >= offset and end <= run_end:
            end_run = run
            end_local = end - offset
            break
        offset = run_end

    if start_run is None or end_run is None:
        return False

    start_index = runs.index(start_run)
    end_index = runs.index(end_run)
    start_text = start_run.text or ""
    end_text = end_run.text or ""
    before = start_text[:start_local]
    after = end_text[end_local:]

    start_run.text = before
    for run in runs[start_index + 1 : end_index]:
        run.text = ""
    if end_run is start_run:
        end_run.text = before
    else:
        end_run.text = after

    value_run = paragraph.add_run(center_value_for_underline(value, template_text))
    value_run.font.underline = True
    set_east_asia_font(value_run, size_pt=12)
    start_run._r.addnext(value_run._r)
    if end_run is start_run and after:
        after_run = paragraph.add_run(after)
        copy_run_color(start_run, after_run)
        value_run._r.addnext(after_run._r)
    return True


def target_range_in_paragraph(paragraph, target: Dict[str, Any]) -> Optional[Tuple[int, int]]:
    kind = target_kind(target)
    if kind in {"runUnderline", "cellRunUnderline"}:
        run_index = int(target.get("run", 0))
        runs = list(paragraph.runs)
        if not runs:
            return None
        run_index = max(0, min(run_index, len(runs) - 1))
        start = sum(len(run.text or "") for run in runs[:run_index])
        return start, start + len(runs[run_index].text or "")
    if kind in {"underline", "cellUnderline"}:
        matches = list(UNDERLINE_PATTERN.finditer(paragraph.text))
        if not matches:
            return None
        underline_index = max(0, min(int(target.get("underline", 0)), len(matches) - 1))
        match = matches[underline_index]
        return match.start(), match.end()
    return None


def find_field_label_position_before(text: str, field: str, before: int) -> Tuple[int, int]:
    best_pos = -1
    best_end = -1
    for alias in sorted(AUTO_FIELD_ALIASES.get(field, []), key=len, reverse=True):
        search_start = 0
        while search_start < len(text):
            pos = find_alias_position(text[search_start:], alias)
            if pos < 0:
                break
            absolute = search_start + pos
            if absolute <= before and absolute >= best_pos:
                best_pos = absolute
                best_end = absolute + len(alias)
            search_start = absolute + 1
    return best_pos, best_end


def find_room_label_position_before(text: str, before: int) -> Tuple[int, int]:
    return find_field_label_position_before(text, "room", before)


def labeled_field_range_around_target(paragraph, target: Dict[str, Any], field: str) -> Optional[Tuple[int, int]]:
    text = paragraph.text
    target_range = target_range_in_paragraph(paragraph, target)
    if not text or not target_range:
        return None
    target_start, target_end = target_range
    label_pos, label_end = find_field_label_position_before(text, field, target_start)
    if label_pos < 0:
        return None

    if target_kind(target) in {"cellUnderline", "cellRunUnderline"}:
        base_target = {
            "kind": "cellUnderline",
            "table": int(target["table"]),
            "row": int(target["row"]),
            "col": int(target["col"]),
            "paragraph": int(target.get("paragraph", 0)),
        }
    else:
        base_target = {"kind": "underline", "paragraph": int(target.get("paragraph", 0))}

    underline_items = paragraph_underline_targets(paragraph, base_target)
    field_start = -1
    for item in underline_items:
        start = int(item["start"])
        if start >= label_end and start <= target_end:
            field_start = start
            break
    if field_start < 0:
        field_start = target_start

    field_end = find_next_field_position(text, field_start, {field})
    if field_end < 0:
        field_end = len(text)
    if field_end <= field_start:
        return None
    return field_start, field_end


def room_field_range_around_target(paragraph, target: Dict[str, Any]) -> Optional[Tuple[int, int]]:
    return labeled_field_range_around_target(paragraph, target, "room")


def collapse_labeled_field_around_target(paragraph, target: Dict[str, Any], field: str, value: str) -> bool:
    field_range = labeled_field_range_around_target(paragraph, target, field)
    if not field_range:
        return False
    return replace_text_range_with_underlined_run(paragraph, field_range[0], field_range[1], value)


def collapse_room_field_around_target(paragraph, target: Dict[str, Any], room_value: str) -> bool:
    field_range = room_field_range_around_target(paragraph, target)
    if not field_range:
        return False
    return replace_text_range_with_underlined_run(paragraph, field_range[0], field_range[1], room_value)


def clear_room_tail_after_run(paragraph, run_index: int) -> None:
    runs = paragraph.runs
    if run_index < 0 or run_index + 1 >= len(runs):
        return
    offsets: List[int] = []
    offset = 0
    for run in runs:
        offsets.append(offset)
        offset += len(run.text or "")
    start = offsets[run_index] + len(runs[run_index].text or "")
    full_text = paragraph.text
    label_pos = find_next_non_room_field_position(full_text, start)
    if label_pos >= 0:
        replace_text_range_in_runs(paragraph, start, label_pos, "    ")
        return
    room_end_candidates = [full_text.find(unit, start) for unit in "\u5ba4" if full_text.find(unit, start) >= 0]
    if room_end_candidates:
        replace_text_range_in_runs(paragraph, start, min(room_end_candidates) + 1, "")


def collapse_room_field_in_plain_underline(paragraph, underline_index: int, room_value: str) -> bool:
    text = paragraph.text
    matches = list(UNDERLINE_PATTERN.finditer(text))
    if not matches:
        return False
    index = max(0, min(int(underline_index), len(matches) - 1))
    first = matches[index]
    label_pos = find_next_non_room_field_position(text, first.end())
    end_pos = label_pos if label_pos >= 0 else len(text)
    template_text = text[first.start() : end_pos]
    runs = list(paragraph.runs)
    if not runs:
        value_run = paragraph.add_run(center_value_for_underline(room_value, template_text))
        value_run.font.underline = True
        set_east_asia_font(value_run, size_pt=12)
        return True

    start_run = None
    start_local = 0
    end_run = None
    end_local = 0
    offset = 0
    for run in runs:
        run_text = run.text or ""
        run_end = offset + len(run_text)
        if start_run is None and first.start() >= offset and first.start() <= run_end:
            start_run = run
            start_local = first.start() - offset
        if end_run is None and end_pos >= offset and end_pos <= run_end:
            end_run = run
            end_local = end_pos - offset
            break
        offset = run_end

    if start_run is None or end_run is None:
        return False

    start_index = runs.index(start_run)
    end_index = runs.index(end_run)
    start_text = start_run.text or ""
    end_text = end_run.text or ""
    before = start_text[:start_local]
    after = end_text[end_local:]

    start_run.text = before
    for run in runs[start_index + 1 : end_index]:
        run.text = ""
    if end_run is start_run:
        after_run = None
    else:
        end_run.text = after

    value_run = paragraph.add_run(center_value_for_underline(room_value, template_text))
    value_run.font.underline = True
    set_east_asia_font(value_run, size_pt=12)
    start_run._r.addnext(value_run._r)
    if end_run is start_run and after:
        after_run = paragraph.add_run(after)
        copy_run_color(start_run, after_run)
        value_run._r.addnext(after_run._r)
    return True


def collapse_room_units_in_runs(paragraph, run_index: int, room_value: str) -> bool:
    runs = paragraph.runs
    if run_index < 0 or run_index >= len(runs):
        return False
    building_run_index = -1
    building_pos = -1
    room_underline_index = -1
    room_unit_run_index = -1
    room_unit_pos = -1

    for index in range(run_index + 1, min(len(runs), run_index + 7)):
        pos = find_first_unit(runs[index].text or "", "\u5e62\u680b")
        if pos >= 0:
            building_run_index = index
            building_pos = pos
            break
    if building_run_index < 0:
        return False

    for index in range(building_run_index + 1, min(len(runs), building_run_index + 5)):
        if run_has_underline(runs[index]) and runs[index].text:
            room_underline_index = index
            break
    if room_underline_index < 0:
        return False

    for index in range(room_underline_index + 1, min(len(runs), room_underline_index + 5)):
        pos = find_first_unit(runs[index].text or "", "\u5ba4")
        if pos >= 0:
            room_unit_run_index = index
            room_unit_pos = pos
            break
    if room_unit_run_index < 0:
        return False

    template_text = (
        (runs[run_index].text or "")
        + (runs[building_run_index].text or "")[building_pos : building_pos + 1]
        + (runs[room_underline_index].text or "")
        + (runs[room_unit_run_index].text or "")[room_unit_pos : room_unit_pos + 1]
    )
    runs[run_index].text = center_value_for_underline(room_value, template_text)
    runs[run_index].font.underline = True
    set_east_asia_font(runs[run_index], size_pt=12)
    building_text = runs[building_run_index].text or ""
    runs[building_run_index].text = building_text[:building_pos] + building_text[building_pos + 1 :]
    runs[room_underline_index].text = ""
    room_text = runs[room_unit_run_index].text or ""
    runs[room_unit_run_index].text = room_text[:room_unit_pos] + room_text[room_unit_pos + 1 :]
    return True


def collapse_room_units_in_plain_underline(paragraph, underline_index: int, room_value: str) -> bool:
    text = paragraph.text
    matches = list(UNDERLINE_PATTERN.finditer(text))
    if not matches:
        return False
    index = max(0, min(int(underline_index), len(matches) - 1))
    first = matches[index]
    building_pos_candidates = [text.find(unit, first.end()) for unit in "\u5e62\u680b"]
    building_positions = [pos for pos in building_pos_candidates if pos >= 0]
    if not building_positions:
        return False
    building_pos = min(building_positions)
    second = next((match for match in matches if match.start() > building_pos), None)
    if not second:
        return False
    room_pos = text.find("\u5ba4", second.end())
    if room_pos < 0:
        return False
    before = text[: first.start()]
    template_text = text[first.start() : room_pos + 1]
    after = text[room_pos + 1 :]
    first_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    for run in paragraph.runs:
        run.text = ""
    first_run.text = before
    value_run = paragraph.add_run(center_value_for_underline(room_value, template_text))
    value_run.font.underline = True
    set_east_asia_font(value_run, size_pt=12)
    after_run = paragraph.add_run(after)
    copy_run_color(first_run, after_run)
    return True


def replace_room_field_target(document, target: Dict[str, Any], room_value: str, label_text: str = "") -> None:
    room_text = "" if room_value is None else str(room_value).strip()
    if target_kind(target) in {"runUnderline", "cellRunUnderline"}:
        paragraph = get_target_paragraph(document, target)
        if collapse_room_field_around_target(paragraph, target, room_text):
            return
        run_index = int(target.get("run", 0))
        replace_run_underline_in_paragraph(paragraph, run_index, room_text, centered=True)
        clear_room_tail_after_run(paragraph, run_index)
        return
    if target_kind(target) in {"underline", "cellUnderline"}:
        paragraph = get_target_paragraph(document, target)
        if collapse_room_field_around_target(paragraph, target, room_text):
            return
        if collapse_room_field_in_plain_underline(paragraph, int(target.get("underline", 0)), room_text):
            return
        replace_field_target(document, target, room_text)
        return
    replace_field_target(document, target, room_text)


def replace_labeled_field_target(document, target: Dict[str, Any], field: str, value: str) -> None:
    field_text = "" if value is None else str(value).strip()
    if target_kind(target) in {"runUnderline", "cellRunUnderline", "underline", "cellUnderline"}:
        paragraph = get_target_paragraph(document, target)
        if collapse_labeled_field_around_target(paragraph, target, field, field_text):
            return
    replace_field_target(document, target, field_text)


def get_target_text(document, target: Dict[str, Any]) -> str:
    if target_kind(target) in {"underline", "runUnderline", "cellUnderline", "cellRunUnderline"}:
        paragraph = get_target_paragraph(document, target)
        return paragraph.text
    return get_cell(document, target).text


def get_cell(document, target: Dict[str, int]):
    table = document.tables[int(target["table"])]
    return table.rows[int(target["row"])].cells[int(target["col"])]


def target_label(target: Optional[Dict[str, int]]) -> str:
    if not target:
        return "未设置"
    if target_kind(target) == "underline":
        return f"下划线 / 段{int(target.get('paragraph', 0)) + 1} / 第{int(target.get('underline', 0)) + 1}处"
    if target_kind(target) == "cellUnderline":
        return (
            f"表{int(target['table']) + 1} / 行{int(target['row']) + 1} / "
            f"列{int(target['col']) + 1} / 下划线{int(target.get('underline', 0)) + 1}"
        )
    if target_kind(target) == "runUnderline":
        return f"Word下划线 / 段{int(target.get('paragraph', 0)) + 1} / 第{int(target.get('underline', 0)) + 1}处"
    if target_kind(target) == "cellRunUnderline":
        return (
            f"表{int(target['table']) + 1} / 行{int(target['row']) + 1} / "
            f"列{int(target['col']) + 1} / Word下划线{int(target.get('underline', 0)) + 1}"
        )
    return f"表{int(target['table']) + 1} / 行{int(target['row']) + 1} / 列{int(target['col']) + 1}"


def config_pairs(config: Dict[str, Any]) -> List[Dict[str, Any]]:
    pairs: List[Dict[str, Any]] = []
    for pair in config.get("pairs", []) or []:
        if isinstance(pair, dict):
            pairs.append(pair)
    if config.get("judgment") or config.get("mark"):
        legacy_pair = {
            "judgment": config.get("judgment"),
            "judgmentText": config.get("judgmentText", ""),
            "mark": config.get("mark"),
            "markStyle": config.get("markStyle"),
        }
        if not any(pair.get("judgment") == legacy_pair.get("judgment") and pair.get("mark") == legacy_pair.get("mark") for pair in pairs):
            pairs.insert(0, legacy_pair)
    return pairs


def has_configured_mark(config: Dict[str, Any]) -> bool:
    return any(pair.get("mark") for pair in config_pairs(config))


def document_tables(path: str | Path) -> List[List[List[str]]]:
    document = Document(path)
    tables: List[List[List[str]]] = []
    for table in document.tables:
        table_rows: List[List[str]] = []
        for row in table.rows:
            table_rows.append([cell.text.strip() for cell in row.cells])
        tables.append(table_rows)
    return tables


def blank_mapping() -> Dict[str, Any]:
    return {
        "version": 1,
        "templateName": "",
        "markText": CHECK_MARK,
        "markStyle": {
            "horizontal": "center",
            "vertical": "middle",
            "fontName": "Arial",
            "fontSize": 10,
            "bold": False,
            "offsetX": 0,
            "offsetY": 0,
            "offsetUnits": "pt",
        },
        "fieldTargets": {},
        "fieldStyles": {
            field: default_field_style()
            for field in ("building", "roomNo", "room", "name", "phone")
        },
        "options": {},
        "exportMode": "multi",
        "cleanMode": False,
        "validation": {
            "mode": "range",
            "min": 1,
            "max": 11,
            "exact": 1,
            "skipInvalid": True,
        },
    }


def safe_filename(text: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\r\n\t]+', "_", text).strip(" ._")
    return cleaned or "未命名"


def unique_output_path(directory: Path, stem: str, suffix: str) -> Path:
    path = directory / f"{stem}{suffix}"
    if not path.exists():
        return path
    for index in range(2, 10000):
        candidate = directory / f"{stem}_{index}{suffix}"
        if not candidate.exists():
            return candidate
    raise RuntimeError("输出文件重名过多，请清理输出目录。")


def timestamped_output_dir(directory: str | Path) -> Path:
    base = Path(directory)
    now = datetime.now()
    stem = f"{now.year}-{now.month}-{now.day}-{now.hour}-{now.minute}-{now.second}"
    path = base / stem
    if not path.exists():
        path.mkdir(parents=True, exist_ok=True)
        return path
    for index in range(2, 10000):
        candidate = base / f"{stem}_{index}"
        if not candidate.exists():
            candidate.mkdir(parents=True, exist_ok=True)
            return candidate
    raise RuntimeError("输出目录重名过多，请稍后再试或清理输出目录。")


def split_room_value(value: Any) -> Optional[Tuple[str, str]]:
    text = normalize_text(value)
    if not text:
        return None
    match = re.match(r"^(.+?)[\-－–—](.+)$", text)
    if match:
        return match.group(1), match.group(2)
    match = re.match(r"^(.+?)[幢栋](.+?)(?:室)?$", text)
    if match:
        return match.group(1), match.group(2)
    return None


def format_room_value(value: Any, template_text: str = "") -> str:
    text = "" if value is None else str(value).strip()
    parts = split_room_value(text)
    if not parts:
        return text
    building, room_no = parts
    template = normalize_text(template_text)
    if "-" in template or "－" in template:
        return f"{building}-{room_no}"
    return f"{building}幢{room_no}"


def option_sort_key(value: str) -> Tuple[int, Any]:
    match = re.search(r"(\d+)", value)
    if match:
        return 0, int(match.group(1))
    return 1, value


def configured_option_keys(mapping: Dict[str, Any]) -> List[str]:
    return sorted(
        [
            key
            for key, config in mapping.get("options", {}).items()
            if has_configured_mark(config)
        ],
        key=option_sort_key,
    )


def validate_vote_record(record: VoteRecord, mapping: Dict[str, Any]) -> List[str]:
    reasons: List[str] = []
    validation = mapping.get("validation", {})
    mode = validation.get("mode", "range")
    total_count = len(record.result_options) if record.result_options else len(record.options)
    if mode == "exact":
        exact = int(validation.get("exact") or 0)
        if exact > 0 and total_count != exact:
            reasons.append(f"票数过多或过少，视为废票或弃票：应为 {exact} 个，实际 {total_count} 个")
    else:
        min_count = int(validation.get("min") or 0)
        max_count = int(validation.get("max") or 0)
        if min_count > 0 and total_count < min_count:
            reasons.append(f"票数过多或过少，视为废票或弃票：至少 {min_count} 个，实际 {total_count} 个")
        if max_count > 0 and total_count > max_count:
            reasons.append(f"票数过多或过少，视为废票或弃票：最多 {max_count} 个，实际 {total_count} 个")

    return reasons


def export_exception_records(
    records: List[Tuple[VoteRecord, List[str]]],
    output_dir: str | Path,
    filename: str = "异常数据.xlsx",
) -> Optional[Path]:
    if not records:
        return None
    output_directory = Path(output_dir)
    output_directory.mkdir(parents=True, exist_ok=True)
    path = unique_output_path(output_directory, Path(filename).stem, ".xlsx")

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "异常数据"
    headers = ["数据行号", "房号/地址", "姓名", "电话号码", "投票选项", "异常原因"]
    sheet.append(headers)
    for record, reasons in records:
        sheet.append([
            record.row_no,
            record.room,
            record.name,
            record.phone,
            "、".join(display_option_key(option) for option in record.options),
            "；".join(reasons),
        ])

    widths = [12, 16, 14, 18, 38, 58]
    for index, width in enumerate(widths, start=1):
        sheet.column_dimensions[openpyxl.utils.get_column_letter(index)].width = width
    for cell in sheet[1]:
        cell.font = openpyxl.styles.Font(bold=True)
        cell.fill = openpyxl.styles.PatternFill("solid", fgColor="FCE4D6")
    sheet.freeze_panes = "A2"
    workbook.save(path)
    return path


def collect_unwritten_votes(record: VoteRecord, mapping: Dict[str, Any]) -> List[Tuple[VoteRecord, str, str, str]]:
    missing: List[Tuple[VoteRecord, str, str, str]] = []
    options = mapping.get("options", {})

    if record.result_options:
        for result_name, selected_options in record.result_options.items():
            for selected_option in selected_options:
                config_key, config = result_option_config(options, result_name, selected_option)
                if not config:
                    missing.append((record, result_name, selected_option, "该结果未配置判断区和标记区"))
                    continue
                matches = pairs_for_result_selection(config_key or "", result_name, config, selected_option)
                if not matches:
                    missing.append((record, result_name, selected_option, "没有匹配到判断区文本"))
                elif not any(pair.get("mark") for pair in matches):
                    missing.append((record, result_name, selected_option, "匹配到判断区，但没有配置对应标记区"))
        return missing

    for option in record.options:
        config = options.get(option)
        if not config:
            missing.append((record, "", option, "该选项未配置判断区和标记区"))
            continue
        pairs = config_pairs(config)
        if not pairs:
            missing.append((record, "", option, "该选项未配置判断区和标记区"))
        elif not any(pair.get("mark") for pair in pairs):
            missing.append((record, "", option, "没有配置标记区"))
    return missing


def export_vote_summary(
    failed_records: List[Tuple[VoteRecord, str]],
    normal_records: List[Tuple[VoteRecord, Path, str]],
    output_dir: str | Path,
) -> Optional[Path]:
    if not failed_records and not normal_records:
        return None
    output_directory = Path(output_dir)
    output_directory.mkdir(parents=True, exist_ok=True)
    path = output_directory / "投票结果汇总.xlsx"

    workbook = openpyxl.Workbook()
    success_sheet = workbook.active
    success_sheet.title = "成功导出的数据"
    failed_sheet = workbook.create_sheet("未导出的数据")

    raw_headers: List[str] = []
    for record, _output_path, _note in normal_records:
        for header in record.raw.keys():
            header_text = str(header)
            if header_text not in raw_headers:
                raw_headers.append(header_text)
    for record, _reason in failed_records:
        for header in record.raw.keys():
            header_text = str(header)
            if header_text not in raw_headers:
                raw_headers.append(header_text)

    headers = raw_headers + ["原因", "结果"]
    success_sheet.append(headers)
    for record, output_path, note in normal_records:
        raw_values = [record.raw.get(header, "") for header in raw_headers]
        success_sheet.append(raw_values + [note or "正常", "有效"])

    failed_sheet.append(headers)
    for record, reason in failed_records:
        raw_values = [record.raw.get(header, "") for header in raw_headers]
        failed_sheet.append(raw_values + [reason, "废票"])

    for sheet in (success_sheet, failed_sheet):
        for index, header in enumerate(headers, start=1):
            width = 16
            if str(header) in {"投票选项", "表决意见", "短信结果"}:
                width = 42
            elif str(header) == "原因":
                width = 22
            elif str(header) == "结果":
                width = 12
            elif str(header) in {"序号", "姓名"}:
                width = 12
            sheet.column_dimensions[openpyxl.utils.get_column_letter(index)].width = width
        for cell in sheet[1]:
            cell.font = openpyxl.styles.Font(bold=True)
            cell.fill = openpyxl.styles.PatternFill("solid", fgColor="E2F0D9")
        sheet.freeze_panes = "A2"
    workbook.save(path)
    return path


def apply_field_targets(document, mapping: Dict[str, Any], record: VoteRecord) -> None:
    room_parts = split_room_value(record.room)
    building, room_no = room_parts if room_parts else ("", "")
    values = {
        "building": building,
        "roomNo": room_no,
        "room": record.room,
        "name": record.name,
        "phone": record.phone,
    }
    clean_mode = bool(mapping.get("cleanMode"))
    inferred = infer_field_targets(document)
    configured = {
        field: (target, False, "")
        for field, target in mapping.get("fieldTargets", {}).items()
        if field in values
    }
    field_targets = {**inferred, **configured}
    if ("building" in field_targets or "roomNo" in field_targets) and "room" not in configured:
        field_targets.pop("room", None)
    field_styles = mapping.get("fieldStyles", {})
    # Fill the later address slot first. Plain underscore indexes remain valid
    # after the first replacement, and template units such as 栋/幢/室 stay put.
    for field in ("name", "phone", "room", "roomNo", "building"):
        target_info = field_targets.get(field)
        if not target_info:
            continue
        target, same_cell, label_text = target_info
        value = values.get(field, "")
        if not value:
            continue
        try:
            style = field_styles.get(field) or default_field_style()
            if target_kind(target) in {"runUnderline", "cellRunUnderline", "underline", "cellUnderline"}:
                # Replace only the exact selected underline. Never collapse the
                # surrounding label or address units.
                replace_field_target(document, target, value, style)
            else:
                cell = get_cell(document, target)
                if same_cell:
                    if clean_mode:
                        set_field_cell_text(cell, value, style)
                    else:
                        set_labeled_field_text(cell, field, value, label_text, style)
                else:
                    set_field_cell_text(cell, value, style)
        except Exception:
            continue


def set_mark_at_target(
    document,
    target: Dict[str, Any],
    text: str,
    style: Optional[Dict[str, Any]] = None,
    clean_mode: bool = False,
) -> None:
    if target_kind(target) in {"runUnderline", "cellRunUnderline"}:
        paragraph = get_target_paragraph(document, target)
        replace_run_underline_in_paragraph(paragraph, int(target.get("run", 0)), text)
        if clean_mode:
            run_index = int(target.get("run", 0))
            if paragraph.runs and 0 <= run_index < len(paragraph.runs):
                paragraph.runs[run_index].font.underline = False
        return
    if target_kind(target) in {"underline", "cellUnderline"}:
        paragraph = get_target_paragraph(document, target)
        if clean_mode:
            replace_plain_underline_with_run(paragraph, int(target.get("underline", 0)), text)
        else:
            replace_underline_in_paragraph(paragraph, int(target.get("underline", 0)), text)
        return
    cell = get_cell(document, target)
    set_mark_text(cell, text, style)


def apply_vote_marks(document, mapping: Dict[str, Any], record: VoteRecord) -> List[str]:
    warnings: List[str] = []
    mark_text = mapping.get("markText") or CHECK_MARK
    clean_mode = bool(mapping.get("cleanMode"))
    options = mapping.get("options", {})

    def apply_pairs(display_option: str, config: Dict[str, Any], pairs: List[Dict[str, Any]]) -> None:
        if not any(pair.get("mark") for pair in pairs):
            warnings.append(f"第{record.row_no}行 {record.name}：{display_option} 没有配置打勾区域")
            return
        for pair_index, pair in enumerate(pairs, start=1):
            judgment = pair.get("judgment")
            expected_text = normalize_text(pair.get("judgmentText", ""))
            if judgment and expected_text:
                try:
                    actual_text = normalize_text(get_target_text(document, judgment))
                    if expected_text not in actual_text:
                        warnings.append(
                            f"第{record.row_no}行 {record.name}：{display_option} 第{pair_index}组判断区文本变化，原为“{pair.get('judgmentText')}”，当前为“{get_target_text(document, judgment).strip()}”"
                        )
                except Exception:
                    warnings.append(f"第{record.row_no}行 {record.name}：{display_option} 第{pair_index}组判断区不存在")

            mark = pair.get("mark")
            if not mark:
                warnings.append(f"第{record.row_no}行 {record.name}：{display_option} 第{pair_index}组没有配置打勾区域")
                continue
            try:
                set_mark_at_target(
                    document,
                    mark,
                    mark_text,
                    pair.get("markStyle") or config.get("markStyle") or mapping.get("markStyle"),
                    clean_mode=clean_mode,
                )
            except Exception:
                warnings.append(f"第{record.row_no}行 {record.name}：{display_option} 第{pair_index}组打勾区域不存在")

    if record.result_options:
        for result_name, selected_options in record.result_options.items():
            for selected_option in selected_options:
                display_option = f"{result_name}：{selected_option}"
                config_key, config = result_option_config(options, result_name, selected_option)
                if not config:
                    warnings.append(f"第{record.row_no}行 {record.name}：{display_option} 没有配置标记区域")
                    continue
                pairs = pairs_for_result_selection(config_key or "", result_name, config, selected_option)
                if not pairs:
                    warnings.append(f"第{record.row_no}行 {record.name}：{display_option} 没有匹配的判断区")
                    continue
                apply_pairs(display_option, config, pairs)
        return warnings

    for option in record.options:
        display_option = display_option_key(option)
        config = options.get(option)
        if not config:
            warnings.append(f"第{record.row_no}行 {record.name}：{display_option} 没有配置标记区域")
            continue

        pairs = config_pairs(config)
        apply_pairs(display_option, config, pairs)
    return warnings


def build_document_for_record(
    template_path: str | Path,
    mapping: Dict[str, Any],
    record: VoteRecord,
) -> Tuple[Document, List[str]]:
    document = Document(template_path)
    clean_mode = bool(mapping.get("cleanMode"))
    if clean_mode:
        hide_template_content(document)
    replace_placeholders(
        document,
        {
            "room": record.room,
            "name": record.name,
            "phone": record.phone,
        },
        clean_mode=clean_mode,
    )
    apply_field_targets(document, mapping, record)
    warnings = apply_vote_marks(document, mapping, record)
    if clean_mode:
        remove_generated_underlines(document)
        # Do not flatten tables into tab-delimited paragraphs. Keeping the
        # original table tree makes pure and non-pure output share the exact
        # same page geometry and mark coordinates.
    return document, warnings


def generate_docx_for_record(
    template_path: str | Path,
    mapping: Dict[str, Any],
    record: VoteRecord,
    output_dir: str | Path,
) -> Tuple[Path, List[str]]:
    document, warnings = build_document_for_record(template_path, mapping, record)

    output_directory = Path(output_dir)
    output_directory.mkdir(parents=True, exist_ok=True)
    filename_prefix = str(mapping.get("filenamePrefix") or "").strip()
    stem = safe_filename(f"{filename_prefix}{record.name or f'第{record.row_no}行'}")
    output_path = unique_output_path(output_directory, stem, ".docx")
    document.save(output_path)
    return output_path, warnings


def append_document_body(target_document: Document, source_document: Document, add_page_break: bool = True) -> None:
    if add_page_break:
        target_document.add_page_break()
    target_body = target_document.element.body
    target_sect_pr = target_body.find(qn("w:sectPr"))
    if target_sect_pr is not None:
        target_body.remove(target_sect_pr)
    for element in source_document.element.body:
        if element.tag == qn("w:sectPr"):
            continue
        target_body.append(deepcopy(element))
    if target_sect_pr is not None:
        target_body.append(target_sect_pr)


def generate_all(
    template_path: str | Path,
    data_path: str | Path,
    mapping: Dict[str, Any],
    output_dir: str | Path,
) -> Tuple[List[Path], List[str], Optional[Path], int, Optional[Path], int, Path]:
    records = read_vote_records_for_mapping(data_path, mapping)
    run_output_dir = timestamped_output_dir(output_dir)
    export_mode = str(mapping.get("exportMode") or "multi")
    outputs: List[Path] = []
    warnings: List[str] = []
    failed_records: List[Tuple[VoteRecord, str]] = []
    normal_records: List[Tuple[VoteRecord, Path, str]] = []

    if export_mode == "single":
        filename_prefix = str(mapping.get("filenamePrefix") or "").strip()
        combined_stem = safe_filename(f"{filename_prefix}合并导出")
        combined_path = unique_output_path(run_output_dir, combined_stem, ".docx")
        combined_document: Optional[Document] = None
        for record in records:
            reasons = validate_vote_record(record, mapping)
            if reasons:
                failed_records.append((record, "；".join(reasons)))
                continue
            document, record_warnings = build_document_for_record(template_path, mapping, record)
            if combined_document is None:
                combined_document = document
            else:
                append_document_body(combined_document, document, add_page_break=True)
            normal_records.append((record, combined_path, "正常"))
            warnings.extend(record_warnings)
        if combined_document is not None:
            combined_document.save(combined_path)
            outputs.append(combined_path)
    else:
        for record in records:
            reasons = validate_vote_record(record, mapping)
            if reasons:
                failed_records.append((record, "；".join(reasons)))
                continue
            output_path, record_warnings = generate_docx_for_record(template_path, mapping, record, run_output_dir)
            outputs.append(output_path)
            normal_records.append((record, output_path, "正常"))
            warnings.extend(record_warnings)

    exception_path = None
    summary_path = export_vote_summary(failed_records, normal_records, run_output_dir)
    return outputs, warnings, exception_path, len(failed_records), summary_path, len(failed_records), run_output_dir


def generate_preview_docx(
    template_path: str | Path,
    data_path: str | Path,
    mapping: Dict[str, Any],
    output_dir: str | Path,
) -> Tuple[Path, List[str]]:
    records = read_vote_records_for_mapping(data_path, mapping)
    preview_record, preview_reasons = select_preview_record(records, mapping)

    if preview_record is None:
        if records:
            raise ValueError("没有可预览的正常数据。第一条异常原因：" + "；".join(preview_reasons))
        raise ValueError("数据源为空，无法生成预览。")

    preview_dir = Path(output_dir) / "预览"
    return generate_docx_for_record(template_path, mapping, preview_record, preview_dir)


def select_preview_record(
    records: Iterable[VoteRecord],
    mapping: Dict[str, Any],
) -> Tuple[Optional[VoteRecord], List[str]]:
    """Choose exactly the same first valid record for DOCX and UI preview."""
    first_reasons: List[str] = []
    for record in records:
        reasons = validate_vote_record(record, mapping)
        if not reasons:
            return record, []
        if not first_reasons:
            first_reasons = reasons
    return None, first_reasons


def convert_doc_to_docx(path: str | Path, output_dir: str | Path) -> Path:
    source = Path(path)
    if source.suffix.lower() == ".docx":
        return source
    if source.suffix.lower() != ".doc":
        raise ValueError("模板只支持 .doc 或 .docx。")

    output_directory = Path(output_dir)
    output_directory.mkdir(parents=True, exist_ok=True)

    possible_soffice = [
        shutil.which("soffice"),
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    soffice = next((item for item in possible_soffice if item and Path(item).exists()), None)
    if not soffice:
        raise RuntimeError("当前电脑未找到 LibreOffice，无法自动把 .doc 转成 .docx。请先另存为 .docx 后导入。")

    command = [
        str(soffice),
        "--headless",
        "--convert-to",
        "docx",
        "--outdir",
        str(output_directory),
        str(source),
    ]
    completed = subprocess.run(command, capture_output=True, text=True, check=False)
    converted = output_directory / f"{source.stem}.docx"
    if completed.returncode != 0 or not converted.exists():
        raise RuntimeError("模板转换失败：" + (completed.stderr or completed.stdout or "未知错误"))
    return converted


def create_figure2_template(path: str | Path) -> Path:
    target = Path(path)
    target.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(9.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("百家湖西花园小区 2026 年第一次业主大会临时会议")
    set_east_asia_font(run, "黑体", 15, True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("表  决  票")
    set_east_asia_font(run, "黑体", 17, True)

    meta = document.add_paragraph()
    meta.paragraph_format.space_after = Pt(4)
    run = meta.add_run("投票时间：2026 年 6 月 15 日至 2026 年 6 月 15 日")
    set_east_asia_font(run, size_pt=10, bold=True)
    run = meta.add_run(" " * 30 + "编号：0002316")
    set_east_asia_font(run, size_pt=10, bold=True)

    stamp = document.add_paragraph()
    stamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    stamp.paragraph_format.space_after = Pt(2)
    run = stamp.add_run("小程序投票")
    set_east_asia_font(run, "黑体", 12, True)

    heading = document.add_paragraph()
    heading.paragraph_format.space_after = Pt(2)
    run = heading.add_run("一、业主大会临时会议的议题：")
    set_east_asia_font(run, "黑体", 10.5, True)

    def add_issue(title_text: str, options: List[str]) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(title_text)
        set_east_asia_font(run, size_pt=9.3, bold=True)

        table = document.add_table(rows=3, cols=len(options) + 1)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.style = "Table Grid"
        set_table_grid(table)

        table.rows[0].cells[0].merge(table.rows[1].cells[0])
        set_cell_text(table.rows[0].cells[0], "选项", size_pt=9.5, bold=True)
        set_cell_shading(table.rows[0].cells[0], "F3F3F3")
        for col, option_text in enumerate(options, start=1):
            set_cell_text(table.rows[0].cells[col], f"选项{OPTION_DIGITS_REVERSE[col]}", size_pt=9.5, bold=True)
            set_cell_shading(table.rows[0].cells[col], "F3F3F3")
            set_cell_text(table.rows[1].cells[col], option_text, size_pt=9.5, bold=True)
        set_cell_text(table.rows[2].cells[0], "表决意见", size_pt=9.5, bold=True)
        set_cell_shading(table.rows[2].cells[0], "F3F3F3")
        for col in range(1, len(options) + 1):
            set_cell_text(table.rows[2].cells[col], "", size_pt=12)

    add_issue(
        "议题一：是否同意授权业主委员会对公共收益进行分配。（具体方案见公告附件一）",
        ["同意", "反对", "弃权"],
    )
    add_issue(
        "议题二：是否同意授权业主委员会根据所附的《物业服务合同》（草案）（见公告附件三）与世茂天成物业服务集团有限公司续签物业服务合同",
        ["同意", "反对", "弃权"],
    )
    add_issue(
        "议题三：同意授权业主委员会采取何种方案对小区南园主干道、北园主干道、东门和南门外车场路面进行道路维修（费用从公共收益支出），并根据相关方案通过公开招标的方式确定中标人（施工单位）（具体方案见公告附件四、公告附件四）",
        ["方案一", "方案二", "均反对", "弃权"],
    )
    add_issue(
        "议题四：同意授权业主委员会采取哪种方案对小区北园、南园水景进行改造或维修（费用从公共收益支出），并根据相关方案通过公开招标的方式确定中标人。（施工单位）（具体方案见公告附件五）",
        ["方案一", "方案二", "均反对", "弃权"],
    )

    instructions_title = document.add_paragraph()
    instructions_title.paragraph_format.space_before = Pt(3)
    instructions_title.paragraph_format.space_after = Pt(1)
    run = instructions_title.add_run("二、填票说明：")
    set_east_asia_font(run, "黑体", 10.5, True)

    instructions = [
        "1、请用黑色或者蓝色墨水笔填写，铅笔填写无效。",
        "2、填写错误，请重新领取表格，涂改无效。",
        "3、请分别在每个议题的选项下面对应的“表决意见”栏打“√”，每个议题都是单选，多选的为废票，计入参与票。",
        "4、未按照要求填写的表决票，在表决票上有明确的房号（或可以确认业主房号）及投票人签名时，为废票，计入参与票。",
        "5、表决票应由业主本人签字，同住家庭成员代签须征得业主本人同意，家庭成员代签后视为已经征得业主本人同意。",
        "6、本人在签字之前，已仔细阅读、理解公示（公告）之具体内容。",
    ]
    for item in instructions:
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.first_line_indent = Cm(-0.45)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(item)
        set_east_asia_font(run, size_pt=8.8)

    fields = document.add_table(rows=2, cols=4)
    fields.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field_data = [
        ["楼栋房号：", "", "姓    名：", ""],
        ["房屋面积：", "", "电话号码：", ""],
    ]
    for row_index, row in enumerate(fields.rows):
        for col_index, cell in enumerate(row.cells):
            value = field_data[row_index][col_index]
            align = WD_ALIGN_PARAGRAPH.LEFT if col_index in (0, 2) else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cell, value, size_pt=10, bold=col_index in (0, 2), align=align)
            set_cell_border(
                cell,
                top={"val": "nil"},
                left={"val": "nil"},
                right={"val": "nil"},
                bottom={"val": "single" if col_index in (1, 3) else "nil", "sz": "8", "color": "888888"},
            )

    committee = document.add_paragraph()
    committee.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    committee.paragraph_format.space_before = Pt(5)
    run = committee.add_run("南京市江宁区百家湖西花园业主委员会")
    set_east_asia_font(run, size_pt=9.5)

    date = document.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = date.add_run("2026 年 4 月 29 日")
    set_east_asia_font(run, size_pt=9.5)

    footer = document.add_paragraph()
    footer.paragraph_format.space_before = Pt(4)
    run = footer.add_run("一式三联：一联业主委员会、一联社区、一联业主留存、一联、第二联投入票箱。")
    set_east_asia_font(run, size_pt=8.5)

    document.save(target)
    return target


def auto_mapping_from_option_tables(template_path: str | Path) -> Dict[str, Any]:
    document = Document(template_path)
    mapping = blank_mapping()
    mapping["templateName"] = Path(template_path).stem
    option_number = 1

    for table_index, table in enumerate(document.tables):
        if not table.rows:
            continue
        rows = table.rows
        option_header_row: Optional[int] = None
        mark_row: Optional[int] = None

        for row_index, row in enumerate(rows):
            first_text = normalize_text(row.cells[0].text if row.cells else "")
            row_text = normalize_text("".join(cell.text for cell in row.cells))
            if option_header_row is None and ("选项" in first_text or row_text.startswith("选项")) and len(row.cells) >= 3:
                option_header_row = row_index
            if mark_row is None and "表决意见" in first_text:
                mark_row = row_index

        if option_header_row is None or mark_row is None:
            continue
        if mark_row <= option_header_row:
            continue

        label_row = option_header_row + 1 if option_header_row + 1 < mark_row else option_header_row
        max_cols = min(len(rows[option_header_row].cells), len(rows[mark_row].cells))
        for col in range(1, max_cols):
            label_text = rows[label_row].cells[col].text.strip() if label_row < len(rows) and col < len(rows[label_row].cells) else ""
            header_text = rows[option_header_row].cells[col].text.strip() if col < len(rows[option_header_row].cells) else ""
            if not label_text and not header_text:
                continue
            key = f"选项{option_number}"
            mapping["options"][key] = {
                "label": key,
                "judgment": {"table": table_index, "row": label_row, "col": col},
                "judgmentText": label_text or header_text,
                "mark": {"table": table_index, "row": mark_row, "col": col},
                "markStyle": {"horizontal": "center", "vertical": "middle", "fontSize": 10},
            }
            option_number += 1

    for table_index, table in enumerate(document.tables):
        flat = normalize_text("\n".join(cell.text for row in table.rows for cell in row.cells))
        if any(keyword in flat for keyword in ("楼栋房号", "房号", "地址", "电话号码", "电话", "姓名")):
            for row_index, row in enumerate(table.rows):
                texts = [normalize_text(cell.text) for cell in row.cells]
                for col_index, text in enumerate(texts):
                    if any(keyword in text for keyword in ("楼栋房号", "房号", "地址")) and col_index + 1 < len(texts):
                        mapping["fieldTargets"].setdefault("room", {"table": table_index, "row": row_index, "col": col_index + 1})
                    if "姓名" in text and col_index + 1 < len(texts):
                        mapping["fieldTargets"].setdefault("name", {"table": table_index, "row": row_index, "col": col_index + 1})
                    if any(keyword in text for keyword in ("电话号码", "电话", "手机号")) and col_index + 1 < len(texts):
                        mapping["fieldTargets"].setdefault("phone", {"table": table_index, "row": row_index, "col": col_index + 1})

    if not mapping["options"]:
        raise ValueError("没有识别到可自动标注的选项表格，请手动使用格式刷标注。")
    return mapping


def auto_mapping_for_template(template_path: str | Path) -> Dict[str, Any]:
    return auto_mapping_from_option_tables(template_path)
